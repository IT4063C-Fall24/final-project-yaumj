#!/usr/bin/env python
# coding: utf-8

# # {Breaking the Glass Firewall: Exploring Salary and Career Disparities Between Men and Women in Tech}📝
# 
# ![Banner](./assets/banner.jpeg)

# ## Topic
# *What problem are you (or your stakeholder) trying to address?*
# 📝 <!-- Answer Below -->
# 
# This project investigates the disparities in salary, career opportunities, and workplace treatment between men and women in the tech industry. Despite having similar qualifications and experience, women in technology face significant barriers to reaching senior positions and earning competitive salaries compared to their male counterparts. Understanding these differences in career trajectories is essential for promoting fairness in the workplace, reducing turnover costs, and improving organizational success by retaining diverse talent.

# ## Project Question
# *What specific question are you seeking to answer with this project?*
# *This is not the same as the questions you ask to limit the scope of the project.*
# 📝 <!-- Answer Below -->
# 
# How do salaries, workplace treatment, and job roles for women compare to those for men with similar educational backgrounds and in comparable fields in the tech industry?

# ## What would an answer look like?
# *What is your hypothesized answer to your question?*
# 📝 <!-- Answer Below -->
# 
# **Hypothesis:**  
# In the tech industry, women with similar qualifications to men earn lower salaries, face more challenges in workplace treatment, and are often relegated to less senior roles. These disparities are influenced by factors that disproportionately affect women, including gender-based discrimination and limited access to career advancement opportunities. As a result, women are less likely than men to achieve competitive salaries or reach senior positions, impacting their long-term career trajectories and earning potential in the industry.

# ## Data Sources
# *What data sources have you identified for this project?*
# *How are you going to relate these datasets?*
# 📝 <!-- Answer Below -->
# 
# **Data Sources**  
# - *HackerRank 2018 Developer Survey (Kaggle):* Published in 2018, based on 2017 survey responses. Primarily focuses on respondents' job roles and educational backgrounds, providing insights into positioning within the tech industry based on qualifications.  
# - *Pew Research Center 2017 STEM Survey (zip file):* Based on 2017 responses, this survey includes data on gender representation and experiences in STEM fields.  
# - *NSF's National Center for Science and Engineering Statistics (Web-Scraped Tables):* Covering data over multiple years, ending in 2019, with an emphasis on STEM workforce statistics.  
# - *NSF's National Center for Science and Engineering Statistics (zip folders with xlsx files):* A supplementary set of tables spanning several years, ending in 2019, offering additional STEM field statistics.  
# - *From College to Jobs American Community Survey 2019 (U.S. Census Bureau xls files):* Detailed data about occupations, salaries, gender, and educational backgrounds for a comprehensive view of career outcomes. 
# - *Stack Overflow Annual Developer Survey - 2017 (zip file):* Based on 2017 survey responses regarding salaries in the tech industry. Provides gender, salary, job title, job rank, years of experience in the industry, and years of experience with the specific employer.
# 
# **Relating the Data**  
# - The datasets can be linked based on the several criteria for consistency:
#     - All of the data is part of a shared timeframe of all of the data (2017-2019).  
#     - Gender will be limited to only male or female. Anyone who did not answer or chose non-binary will be excluded from analysis.
#     - Responses will be limited to the United States. 
#     - Only data where the individual works in a STEM career will be included.
#     - Ages will be limited to the 25 - 64 range to match the American Community Survey data, ensuring consistency across employment data.
#     - The analysis will focus on individuals who were employed at the time of each survey, and each dataset will be filtered accordingly.

# ## Approach and Analysis
# *What is your approach to answering your project question?*
# *How will you use the identified data to answer your project question?*
# 📝 <!-- Start Discussing the project here; you can add as many code cells as you need -->
# 
# After limiting each dataset to core demographics (e.g., country, age, gender, employment status, employment field), I will analyze trends across individual datasets and determine if the responses are similar for the other datasets. As an example, both the HackerRank and Pew Research Center surveys have data about job search criteria, so recreating the same visualization type on each will allow me to observe patterns related to gender disparities.
# 
# **Planned Visualizations to Support the Hypothesis:**
# 1. *Line Chart: Gender and Age Distribution in Technology*  
# A line chart will display the number of employees in the technology sector, separated by gender and age groups. The x-axis will represent the age groups (e.g., 20-25, 26-30, etc.), and the y-axis will show the number of employees. This chart will highlight the drop-off point where women begin to leave the industry earlier than men, helping visualize retention issues.
# 
# 2. *100% Stacked Column Chart: Men vs Women in Different Tech Roles*  
# A 100% stacked column chart will show the proportional representation of men and women across different tech roles (e.g., Junior Developers, Senior Developers, Managers, Executives). Each column will represent a different role, and the stacked columns will show the gender distribution within that role as a percentage. This will provide a clear visual of how underrepresented women are in higher-level positions.
# 
# 3. *Side-by-Side Column Chart: Workplace Concerns for Men vs Women*  
# A side-by-side column chart will compare the key workplace concerns between men and women, such as issues with career progression, work-life balance, pay disparity, and workplace discrimination. Each concern will have two columns—one representing men and one representing women. This will make it easy to see where concerns overlap and where significant differences exist between the genders.
# 
# ### UPDATE: Plan Revisions After Exploratory Data Analysis:
# - After the initial exploratory data analysis, the visualization choices were adjusted to better align with the patterns emerging in the raw data.

# #### Package Imports

# In[246]:


#import packages
import os # to create subfolder for data organization
from dotenv import load_dotenv
load_dotenv(override=True)

import opendatasets as od
import pandas as pd
import numpy as np
import pyreadstat
import requests
import re # for string manipulation

import matplotlib.pyplot as plt
import matplotlib.ticker as mtick
import seaborn as sns
import plotly.express as px
import plotly.graph_objects as go

import warnings

from zipfile import ZipFile
from urllib.request import urlretrieve
from bs4 import BeautifulSoup

from sklearn.model_selection import train_test_split, StratifiedShuffleSplit, cross_val_score
from sklearn.neighbors import KNeighborsClassifier
from sklearn.preprocessing import LabelEncoder, OneHotEncoder, OrdinalEncoder, MultiLabelBinarizer, StandardScaler, PolynomialFeatures
from sklearn.metrics import classification_report, confusion_matrix
from sklearn.linear_model import LinearRegression
from sklearn.metrics import r2_score, mean_squared_error, mean_absolute_error
from sklearn.compose import ColumnTransformer
from sklearn.impute import SimpleImputer
from sklearn.pipeline import Pipeline
from sklearn.compose import make_column_selector
from sklearn.tree import DecisionTreeRegressor

from mpl_toolkits.mplot3d import Axes3D

from pandas.plotting import scatter_matrix

from docx import Document


# #### Import Dataset 1: HackerRank Developer Survey Published in 2018 that covered 2017 Questionnaire Responses

# In[2]:


# import dataset from Kaggle using URL 
dataset_url = "https://www.kaggle.com/datasets/hackerrank/developer-survey-2018/data"
od.download(dataset_url, data_dir="./data")


# #### Convert dataset to a pandas dataframe and inspect data for Exploratory Data Analysis (EDA)

# In[3]:


# Define the data directory
data_dir = './data/developer-survey-2018'

# Read each CSV file into its own DataFrame with meaningful names
country_code_df = pd.read_csv(os.path.join(data_dir, 'Country-Code-Mapping.csv'))
dev_survey_codebook_df = pd.read_csv(os.path.join(data_dir, 'HackerRank-Developer-Survey-2018-Codebook.csv'))
dev_survey_numeric_mapping_df = pd.read_csv(os.path.join(data_dir, 'HackerRank-Developer-Survey-2018-Numeric-Mapping.csv'))
dev_survey_numeric_df = pd.read_csv(os.path.join(data_dir, 'HackerRank-Developer-Survey-2018-Numeric.csv'))
dev_survey_values_df = pd.read_csv(os.path.join(data_dir, 'HackerRank-Developer-Survey-2018-Values.csv'))

# Display the first 5 records from each DataFrame
display(country_code_df.head(5))
display(dev_survey_codebook_df.head(5))
display(dev_survey_numeric_mapping_df.head(5))
display(dev_survey_numeric_df.head(5))
display(dev_survey_values_df.head(5))


# #### Exploratory Data Analysis (EDA) Insights:
# - The dataset is global, not local, meaning it does not precisely align with my other datasets.
#     - Information from the country_code_df can be used to identify the country code and country name format so I can filter the other dataframes to only include the United States
# - The datset includes NaN values in several fields including the country. 
#     - Data will need to be explored further to determine the handling of NaN or #NULL! values
#         - For example, if the NaN or NULL value is in gender or country, it cannot be determined using other methods and needs to be excluded from analysis. However,
#         if the NaN or NULL is in the Job Level or Current Rank field, I may be able to populate those fields based on data from similar records.
# - The Numeric and Values datasets have identical infomation, just in different formats.
#     - In the Numeric dataset, all of the values are numerically coded.
#     - In the Values dataset, all of the responses are in plain English.
# - There are outliers in this data as far as what can be directly related to my other datasets and my hypothesis for the purpose of analysis.
#     - Age ranges will need to be limited to match what is available in the datasets from other sources to make "apples to apples" comparisons.

# In[4]:


# Find United States in the country_code_df to use for filtering purposes

# Rename columns
country_code_df.rename(columns={'Value': 'country_code', 'Label': 'country'}, inplace=True)

# Filter for the United States
us_country_code_df = country_code_df[country_code_df['country'] == 'United States']

# Display the resulting DataFrame
display(us_country_code_df)


# #### Filtered dataframe to include only respondents in the United States

# In[5]:


# Filter the DataFrame for United States questionnaire responses
us_dev_survey_numeric_df = dev_survey_numeric_df[dev_survey_numeric_df['CountryNumeric2'] == 167]

# Display the number of records
num_records_numeric = us_dev_survey_numeric_df.shape[0]
display(f"Number of records: {num_records_numeric}")

# Display the resulting DataFrame
display(us_dev_survey_numeric_df.head(5))


# #### Reduce dataframe columns to only those relevant to supporting or disproving my hypothesis
# - fields such as date survey was completed and questions about the HackerRank survey were removed from dataframe for simplification

# In[6]:


# List of relevant columns to keep
columns_to_keep = [
    'RespondentID', 'q2Age', 'q3Gender', 'q10Industry',
    'q8JobLevel', 'q9CurrentRole',
    'q12JobCritPrefTechStack', 'q12JobCritCompMission',
    'q12JobCritCompCulture', 'q12JobCritWorkLifeBal',
    'q12JobCritCompensation', 'q12JobCritProximity',
    'q12JobCritPerks', 'q12JobCritSmartPeopleTeam',
    'q12JobCritImpactwithProduct', 'q12JobCritInterestProblems',
    'q12JobCritFundingandValuation', 'q12JobCritStability',
    'q12JobCritProfGrowth', 'q16HiringManager',
    'q17HirChaNoDiversCandidates', 'q20CandYearExp',
    'q20CandCompScienceDegree', 'q20CandCodingBootcamp',
    'q20CandSkillCert', 'q20CandHackerRankActivity',
    'q20CandOtherCodingCommAct', 'q20CandGithubPersProj',
    'q20CandOpenSourceContrib', 'q20CandHackathonPart',
    'q20CandPrevWorkExp', 'q20CandPrestigeDegree',
    'q20CandLinkInSkills', 'q20CandGithubPersProj2'
]

# Create a new DataFrame with only the selected columns
filtered_us_dev_survey_numeric_df = us_dev_survey_numeric_df[columns_to_keep]

# Display information about the DataFrame 
filtered_us_dev_survey_numeric_df.info()

# Display the resulting DataFrame
display(filtered_us_dev_survey_numeric_df.head(5))


# #### Check for Duplicate Records
# - Some records will be similar, but all records should have a unique RespondentID

# In[7]:


# Check for duplicates in the RespondentID field
duplicates = filtered_us_dev_survey_numeric_df[filtered_us_dev_survey_numeric_df.duplicated('RespondentID', keep=False)]

# Count the number of duplicate RespondentIDs
num_duplicates = duplicates.shape[0]

# Display the duplicates if any
if num_duplicates > 0:
    print(f"Number of duplicate records found: {num_duplicates}")
    display(duplicates[['RespondentID']])
else:
    print("No duplicate records found in RespondentID.")


# #### Filter DataFrame to be Consistent with American Community Survey Parameters for Data Matching
# - Remove ages under 25 years old (coded as q2Age: 1, 2, or 3)
# - Remove ages over 64 years old (coded as q2Age: 8 or 9)
# - Remove non-binary respondents (coded as q3Gender: 3)

# In[8]:


# Summary of counts for each value in q2Age
age_summary = filtered_us_dev_survey_numeric_df['q2Age'].value_counts(dropna=False)
print("Summary of counts for each value in q2Age:")
print(age_summary)

# Summary of counts for each value in q3Gender
gender_summary = filtered_us_dev_survey_numeric_df['q3Gender'].value_counts(dropna=False)
print("\nSummary of counts for each value in q3Gender:")
print(gender_summary)


# In[9]:


# Remove records where q2Age is #NULL!, 1, 2, 3, 8, or 9
filtered_us_dev_survey_numeric_df = filtered_us_dev_survey_numeric_df[
    ~filtered_us_dev_survey_numeric_df['q2Age'].isin(['#NULL!', '1', '2', '3', '8', '9'])
]

# Remove records where q3Gender is #NULL! or 3
filtered_us_dev_survey_numeric_df = filtered_us_dev_survey_numeric_df[
    ~filtered_us_dev_survey_numeric_df['q3Gender'].isin(['#NULL!', '3'])
]

# Display the row count of the refined DataFrame
row_count_after_filtering = filtered_us_dev_survey_numeric_df.shape[0]
print(f"Row count after filtering: {row_count_after_filtering}")

# New summary of counts for each value in q2Age
age_summary = filtered_us_dev_survey_numeric_df['q2Age'].value_counts(dropna=False)
print("New summary of counts for each value in q2Age:")
print(age_summary)

# Summary of counts for each value in q3Gender
gender_summary = filtered_us_dev_survey_numeric_df['q3Gender'].value_counts(dropna=False)
print("\nNew summary of counts for each value in q3Gender:")
print(gender_summary)


# #### Exploratory Data Analysis (EDA)
# - Begin renaming columns
# - The numeric dataframe should consist entirely of int64 data types, yet the majority have an "object" data type instead.
#     - These datatypes will need to be converted for certain types of analysis like a correlation matrix.

# In[10]:


# Rename columns
filtered_us_dev_survey_numeric_df.rename(columns={
    'q2Age': 'Age',
    'q3Gender': 'Gender',
    'q10Industry': 'Industry',
    'q8JobLevel': 'Job Level',
    'q9CurrentRole': 'Current Role',
    'q10Industry': 'Industry',
    'q12JobCritPrefTechStack': 'Job Search Priority - Preferred Tech Stack',
    'q12JobCritCompMission': 'Job Search Priority - Company Mission',
    'q12JobCritCompCulture': 'Job Search Priority - Company Culture',
    'q12JobCritWorkLifeBal': 'Job Search Priority - Work-Life Balance',
    'q12JobCritCompensation': 'Job Search Priority - Compensation',
    'q12JobCritProximity': 'Job Search Priority - Proximity to Home',
    'q12JobCritPerks': 'Job Search Priority - Perks',
    'q12JobCritSmartPeopleTeam': 'Job Search Priority - Smart Workers',
    'q12JobCritImpactwithProduct': 'Job Search Priority - Product Impact',
    'q12JobCritInterestProblems': 'Job Search Priority - Interesting Problems to Solve',
    'q12JobCritFundingandValuation': 'Job Search Priority - Funding and Valuation',
    'q12JobCritStability': 'Job Search Priority - Stability of Large Company',
    'q12JobCritProfGrowth': 'Job Search Priority - Professional Growth and Learning',
    'q16HiringManager': 'Hiring Manager (Y/N)'
}, inplace=True)

# Convert all columns to Int64 
filtered_us_dev_survey_numeric_df = filtered_us_dev_survey_numeric_df.apply(pd.to_numeric, errors='coerce').astype('Int64')

# Display the resulting Numeric DataFrame and info
print("Revised Numeric DataFrame (No Splitting, Renamed Columns):")
display(filtered_us_dev_survey_numeric_df.info())
display(filtered_us_dev_survey_numeric_df.head(5))


# #### Matplotlib Stacked Bar Chart: Job Criteria Priorities by Gender (Normalized)
# **Purpose:** Compare the importance of various job criteria between men and women in the technology sector, using a stacked bar chart that has been normalized to account for the gender imbalance of respondents in the dataset.
# 
# **Insights:**
# - **Work-Life Balance as Top Priority:** *Work-Life Balance* ranks as the highest priority for women (60.1%) and is also highly valued by men (56%). This could reflect broader societal expectations, where women are often primary caretakers, balancing professional responsibilities with family obligations. While this dataset doesn’t confirm this directly, it aligns with common trends observed in gender studies.
# - **Professional Growth and Learning Priority:** Women place a significantly higher value on *Professional Growth and Learning* (54.6%) compared to men (44.9%). This aligns with findings from the Pew Research Center, suggesting that women, who often face challenges with gender discrimination, prioritize opportunities for advancement and skill development more heavily.
# - **Company Culture Importance:** *Company Culture* ranks higher for women (44.2%) than for men (39.0%), possibly reflecting women’s desire for supportive and inclusive workplace environments in light of reported experiences of exclusion and discrimination.
# - **Compensation Comparison:** For women, *Compensation* is the fourth most important job criterion (41.2%), following *Work-Life Balance* (60.1%), *Professional Growth and Learning* (54.6%), and *Company Culture* (44.2%). This suggests that while pay disparities are significant, women may prioritize job factors that foster long-term satisfaction and career advancement over salary alone when choosing an employer.
# - **Low-Value Criteria for Both Genders:** *Funding and Valuation*, *Stability of a Large Company*, and *Perks* all show low prioritization across both genders, with less than 10% of respondents marking them as essential. This indicates that while job security and benefits matter, they are not primary factors in job selection for either group.
# 
# **Summary:** The findings reinforce the hypothesis that women in tech value career advancement opportunities, inclusive workplace culture, and work-life balance over compensation alone. These insights highlight the importance of creating growth-oriented, flexible, and supportive environments to retain female talent in the technology industry.

# In[11]:


# Define a dictionary mapping the actual DataFrame column names to the desired display names
column_display_names = {
    'Job Search Priority - Preferred Tech Stack': 'Preferred Tech Stack',
    'Job Search Priority - Company Mission': 'Company Mission',
    'Job Search Priority - Company Culture': 'Company Culture',
    'Job Search Priority - Work-Life Balance': 'Work-Life Balance',
    'Job Search Priority - Compensation': 'Compensation',
    'Job Search Priority - Proximity to Home': 'Proximity to Home',
    'Job Search Priority - Perks': 'Perks',
    'Job Search Priority - Smart Workers': 'Smart Workers',
    'Job Search Priority - Product Impact': 'Product Impact',
    'Job Search Priority - Interesting Problems to Solve': 'Interesting Problems to Solve',
    'Job Search Priority - Funding and Valuation': 'Funding and Valuation',
    'Job Search Priority - Stability of Large Company': 'Stability of Large Company',
    'Job Search Priority - Professional Growth and Learning': 'Professional Growth and Learning'
}

# Select only the job criteria columns, excluding the "Other" column
job_criteria_columns = list(column_display_names.keys())

# Group data by gender and sum up the counts for each job criteria
criteria_counts = filtered_us_dev_survey_numeric_df.groupby('Gender')[job_criteria_columns].sum()

# Calculate the total number of male and female respondents
total_counts_by_gender = filtered_us_dev_survey_numeric_df['Gender'].value_counts()

# Normalize the data by calculating the percentage of each job criterion per gender
# This step addresses the disparity between male and female respondents
criteria_percentages = criteria_counts.div(total_counts_by_gender, axis=0) * 100

# Rename gender values for readability if needed
criteria_percentages.index = criteria_percentages.index.map({1: 'Male', 2: 'Female'})

# Rename columns in the DataFrame for plotting
criteria_percentages = criteria_percentages.rename(columns=column_display_names)

# Plotting the normalized stacked bar chart
ax = criteria_percentages.T.plot(kind='bar', stacked=True, figsize=(14, 8), color=['slategray', 'lightseagreen'], edgecolor='none')
plt.suptitle('Job Choice Priorities by Gender (Normalized)')
plt.title('HackerRank 2018 Developer Survey', fontsize=12, color='gray')
plt.ylabel('Percentage of Respondents (%)')
plt.legend(title='Gender')
plt.xticks(rotation=45, ha='right')
plt.tight_layout()

# Adding percentage labels to each section of the bars
for container in ax.containers:
    ax.bar_label(container, fmt='%.1f%%', label_type='center')

plt.show()


# ## Repeat filtering logic from numeric version of data on the values version of the data
# - Both datasets contain the same records, but one has numeric codes for all responses and the other has plain language values for all responses 
# so the same logic can be used for both dataframes.

# In[12]:


# Filter the DataFrame for CountryNumeric2 = "United States"
us_dev_survey_values_df = dev_survey_values_df[dev_survey_values_df['CountryNumeric2'] == "United States"]

# Display the number of records
num_records_values = us_dev_survey_values_df.shape[0]
display(f"Number of records: {num_records_values}")

# Display the resulting DataFrame
display(us_dev_survey_values_df.head(5))


# In[13]:


# Reduce dataframe columns to only those relevant to supporting or disproving my hypothesis

# List of relevant columns to keep
columns_to_keep = [
    'RespondentID', 'q2Age', 'q3Gender', 'q10Industry',
    'q8JobLevel', 'q9CurrentRole',
    'q12JobCritPrefTechStack', 'q12JobCritCompMission',
    'q12JobCritCompCulture', 'q12JobCritWorkLifeBal',
    'q12JobCritCompensation', 'q12JobCritProximity',
    'q12JobCritPerks', 'q12JobCritSmartPeopleTeam',
    'q12JobCritImpactwithProduct', 'q12JobCritInterestProblems',
    'q12JobCritFundingandValuation', 'q12JobCritStability',
    'q12JobCritProfGrowth', 'q16HiringManager',
    'q17HirChaNoDiversCandidates', 'q20CandYearExp',
    'q20CandCompScienceDegree', 'q20CandCodingBootcamp',
    'q20CandSkillCert', 'q20CandHackerRankActivity',
    'q20CandOtherCodingCommAct', 'q20CandGithubPersProj',
    'q20CandOpenSourceContrib', 'q20CandHackathonPart',
    'q20CandPrevWorkExp', 'q20CandPrestigeDegree',
    'q20CandLinkInSkills', 'q20CandGithubPersProj2'
]

# Create a new DataFrame with only the selected columns
filtered_us_dev_survey_values_df = us_dev_survey_values_df[columns_to_keep]

# Display information about the DataFrame 
filtered_us_dev_survey_values_df.info()

# Display the resulting DataFrame
display(filtered_us_dev_survey_values_df.head(5))


# #### Clean the values dataframe
# - Removed records where gender is missing because they cannot be used to prove/disprove the hypothesis.
# - Filtered out records where gender is non-binary to match the parameters of the US Census Burea's American Community Survey.
# - Filtered out records where respondents were under 25 or over 64 to match the parameters of the US Census Bureau's American Community Survey.
# - Filtered out records where the age is null because both age and gender are necessary to determine job level comparisons.
# 

# In[14]:


# Create a copy to work on and avoid SettingWithCopyWarning
filtered_us_dev_survey_values_df = filtered_us_dev_survey_values_df.copy()

# Drop records where Gender is null
filtered_us_dev_survey_values_df.dropna(subset=['q3Gender'], inplace=True)

# Drop records where Gender is '#NULL!' or 'Non-Binary'
filtered_us_dev_survey_values_df.drop(
    filtered_us_dev_survey_values_df[filtered_us_dev_survey_values_df['q3Gender'].isin(['#NULL!', 'Non-Binary'])].index,
    inplace=True
)

# Filter out respondents who are under 25 or over 64 
filtered_us_dev_survey_values_df = filtered_us_dev_survey_values_df[
    ~filtered_us_dev_survey_values_df['q2Age'].isin(["Under 12 years old", "12 - 18 years old", "18 - 24 years old", "65 - 74 years old", "75 years or older"])
]

# Filter out rows where Age is '#NULL!'
filtered_us_dev_survey_values_df = filtered_us_dev_survey_values_df[
    filtered_us_dev_survey_values_df['q2Age'] != '#NULL!'
]

# Display how many rows remain
print(f"Remaining records after cleaning (SHOULD BE 2942 IF CORRECT): {filtered_us_dev_survey_values_df.shape[0]}")

# New summary of counts for each value in q2Age
age_summary = filtered_us_dev_survey_values_df['q2Age'].value_counts(dropna=False)
print("New summary of counts for each value in Age:")
print(age_summary)

# Summary of counts for each value in q3Gender
gender_summary = filtered_us_dev_survey_values_df['q3Gender'].value_counts(dropna=False)
print("\nNew summary of counts for each value in Gender:")
print(gender_summary)


# #### Clean the values dataframe
# - Rename columns
# - Filtered out records where current role or job level is student because they are not relevant to my hypothesis.
# - Filtered out records where both the Job Level and Current Role were NaN because there is no way to determine values for the field if both are blank.
# 

# In[15]:


# Rename columns
filtered_us_dev_survey_values_df.rename(columns={
    'q2Age': 'Age',
    'q3Gender': 'Gender',
    'q10Industry': 'Industry',
    'q8JobLevel': 'Job Level',
    'q9CurrentRole': 'Current Role',
    'q10Industry': 'Industry',
    'q12JobCritPrefTechStack': 'Job Search Priority - Preferred Tech Stack',
    'q12JobCritCompMission': 'Job Search Priority - Company Mission',
    'q12JobCritCompCulture': 'Job Search Priority - Company Culture',
    'q12JobCritWorkLifeBal': 'Job Search Priority - Work-Life Balance',
    'q12JobCritCompensation': 'Job Search Priority - Compensation',
    'q12JobCritProximity': 'Job Search Priority - Proximity to Home',
    'q12JobCritPerks': 'Job Search Priority - Perks',
    'q12JobCritSmartPeopleTeam': 'Job Search Priority - Smart Workers',
    'q12JobCritImpactwithProduct': 'Job Search Priority - Product Impact',
    'q12JobCritInterestProblems': 'Job Search Priority - Interesting Problems to Solve',
    'q12JobCritFundingandValuation': 'Job Search Priority - Funding and Valuation',
    'q12JobCritStability': 'Job Search Priority - Stability of Large Company',
    'q12JobCritProfGrowth': 'Job Search Priority - Professional Growth and Learning',
    'q16HiringManager': 'Hiring Manager (Y/N)'
}, inplace=True)

# Filter out rows where Current Role or Job Level is "Student"
filtered_us_dev_survey_values_df = filtered_us_dev_survey_values_df[
    (filtered_us_dev_survey_values_df['Current Role'] != 'Student') & 
    (filtered_us_dev_survey_values_df['Job Level'] != 'Student')
]

# Filter out rows where both Job Level and Current Role are NaN
filtered_us_dev_survey_values_df = filtered_us_dev_survey_values_df[
    ~ (filtered_us_dev_survey_values_df['Job Level'].isna() & filtered_us_dev_survey_values_df['Current Role'].isna())
]

# Display how many rows remain
print(f"Remaining responses after cleaning: {filtered_us_dev_survey_values_df.shape[0]}")


# #### Matplotlib Bar Graph: Count of Workers by Age Group and Gender
# **Purpose:** Provide an initial demographic breakdown of the HackerRank dataset by age and gender to establish a baseline understanding of workforce composition, which informs the need for further analysis on gender representation across job levels.
# 
# **Insights:**
# - **Age Group Representation**: The distribution of workers by age group reveals a strong representation in the 25–34 and 35–44 age groups, which are critical career-building stages.
# - **Gender Disparity in Age Groups**: Male respondents dominate all age groups, particularly in the 25–34 range, indicating a potential gender imbalance that could influence job levels and advancement opportunities.
# - **Next Steps in Analysis:** Given the evident gender imbalance, further analysis will explore how this demographic distribution correlates with job levels. Examining job levels across genders and age groups can identify whether disproportionate career progression patterns exist. Additional visualizations, such as a stacked bar chart by job level and gender or a heatmap for job level concentration, will deepen the understanding of gender-based trends in career advancement.

# In[16]:


# Group the data by Age Group and Gender, and count occurrences
age_gender_counts = filtered_us_dev_survey_values_df.groupby(['Age', 'Gender'], observed=False).size().unstack(fill_value=0)

# Plotting
plt.figure(figsize=(10, 6))

# Plot each gender as a separate bar
age_gender_counts.plot(kind='bar', width=0.8, ax=plt.gca(), color=['slategray', 'lightseagreen'])

# Customize the plot
plt.title('Count of Workers by Age Group and Gender')
plt.xlabel('Age Group')
plt.ylabel('Count')
plt.xticks(rotation=45)
plt.legend(title='Gender', loc='upper right')
plt.grid(axis='y')  # Add horizontal grid lines for better readability

plt.show()


# #### Additional Exploratory Data Analyis (EDA) for Data Relevance
# - Determine if records where Job Level = New grad have current roles other than "unemployed" to decide if they should be included in analysis.
# - Look for records where the Job Level is NaN but the Current Role is not NaN.
#     - These records can be used with machine learning classification to populate missing values.

# In[17]:


# Review dataset to determine what data is relevant

# Check Current Role for Job Level "New grad"
new_grad_roles = filtered_us_dev_survey_values_df[filtered_us_dev_survey_values_df['Job Level'] == 'New grad'] \
    .groupby('Current Role').size().reset_index(name='Counts')

# Filter for rows where Job Level is NaN and Current Role is not NaN
nan_job_level_current_role_df = filtered_us_dev_survey_values_df[filtered_us_dev_survey_values_df['Job Level'].isna() & 
    filtered_us_dev_survey_values_df['Current Role'].notna()]

# Display results
print("Current Roles for Job Level 'New grad':")
display(new_grad_roles)

print("Rows where Job Level is NaN and Current Role is not NaN:")
display(nan_job_level_current_role_df[['Job Level', 'Current Role']])


# #### More Exploratory Data Analysis (EDA)
# - Check if there is a dominant job level associated with the Current Role field that could be used to populate empty fields.

# In[18]:


# Group by Current Role and Job Level, and count occurrences
role_level_counts = filtered_us_dev_survey_values_df.groupby(['Current Role', 'Job Level']).size().reset_index(name='Counts')

# Set the option to display all rows for this code cell only
with pd.option_context('display.max_rows', None):
    display(role_level_counts)


# #### Seaborn Box Plot: Job Levels for a Specific Career Role, Divided by Gender
# **Purpose:** Demonstrate how gender may impact job levels within specific career roles, in this case, *Development Operations Engineer*. This is to support the use of a machine learning approach to populate empty fields. Machine learning can adapt to demographic differences, rather than relying on single-point estimates like the mean or median for classification.
# 
# **Insights:**
# - **Gender-based Distribution Differences:** 
#     - The box plot shows a significant disparity in job level distribution by gender within the *Development Operations Engineer* role. 
#     - **Male Distribution:** Males exhibit a broader range of job levels, with a median at the *Senior Developer* level. This distribution includes representation across higher job levels, such as *Principal Engineer* and *Engineering Manager*, suggesting greater advancement opportunities or a wider presence in senior roles for men within this career path. 
#     - **Female Distribution:** Females, on the other hand, are primarily clustered at the *Level 1 Developer (junior)* level, with no representation above the *Freelancer* level in this role. This limited range could indicate barriers to progression or underrepresentation in senior positions, reflecting broader industry challenges women face in advancing within technical roles.
# 
# - **Rationale for Machine Learning Approach:** 
#     - Given these clear disparities in job level distribution by gender, a machine learning model like K-Nearest Neighbors (KNN) is well-suited for predicting job levels. KNN can capture complex relationships between demographic features (such as gender) and job levels, providing a more nuanced prediction than simple averages or medians.
#     - By considering demographic factors, this approach enables the model to better reflect real-world patterns of representation and career advancement, offering a data-driven perspective on disparities within roles across the dataset.

# In[19]:


# Define the desired order for job levels from junior to senior based on numeric dataset mapping
job_level_order = [
    "New grad",
    "Freelancer",
    "Level 1 developer (junior)",
    "Senior developer",
    "Principal engineer",
    "Architect",
    "Engineering manager",
    "Founder / CEO / CTO"
]

# Filter for a specific current role 
specific_role = "Development Operations Engineer"
role_data = filtered_us_dev_survey_values_df[filtered_us_dev_survey_values_df['Current Role'] == specific_role].copy()
role_data['Job Level'] = pd.Categorical(role_data['Job Level'], categories=job_level_order, ordered=True)

# Define a custom palette to match matplotlib graphs
custom_palette = ['slategray', 'lightseagreen']

# Create a horizontal box plot with gaps between the boxes
plt.figure(figsize=(6, 6))
sns.boxplot(data=role_data, y='Job Level', x='Gender', hue='Gender', palette=custom_palette, width=0.9, dodge=True)
plt.suptitle(f'Distribution of Job Levels for {specific_role} by Gender')
plt.title('HackerRank 2018 Developer Survey', fontsize=12, color='gray')
plt.ylabel('Job Level')
plt.xlabel('')
plt.xticks(rotation=0)
plt.grid(axis='y')  # Add horizontal grid lines for better readability
plt.show()


# #### Use Machine Learning to Populate NaN Job Level Records

# In[20]:


# Use KNearestNeighbors to determine most likely Job Level based on age, gender, and current role

age_mapping = {
    '25 - 34 years old': 4,
    '35 - 44 years old': 5,
    '45 - 54 years old': 6,
    '55 - 64 years old': 7,
}

gender_mapping = {
    'Male': 1,
    'Female': 2,
}

# Prepare the DataFrame for KNN training
train_df = filtered_us_dev_survey_values_df[filtered_us_dev_survey_values_df['Job Level'].notna()].copy()

# Create numeric mappings for Age and Gender for training
train_df.loc[:, 'Age Numeric'] = train_df['Age'].map(age_mapping)
train_df.loc[:, 'Gender Numeric'] = train_df['Gender'].map(gender_mapping)

# Encode Current Role using Label Encoding
le_role = LabelEncoder()
train_df.loc[:, 'Current Role Encoded'] = le_role.fit_transform(train_df['Current Role'])

# Prepare the feature set and target variable for training
X_train = train_df[['Age Numeric', 'Gender Numeric', 'Current Role Encoded']]
y_train = train_df['Job Level']

# Filter the rows with NaN Job Level for predictions
predict_df = filtered_us_dev_survey_values_df[filtered_us_dev_survey_values_df['Job Level'].isna()].copy()
predict_df.loc[:, 'Age Numeric'] = predict_df['Age'].map(age_mapping)
predict_df.loc[:, 'Gender Numeric'] = predict_df['Gender'].map(gender_mapping)
predict_df.loc[:, 'Current Role Encoded'] = le_role.transform(predict_df['Current Role'])

# Create the feature set for predictions
X_predict = predict_df[['Age Numeric', 'Gender Numeric', 'Current Role Encoded']]

# Initialize and fit KNN model
knn = KNeighborsClassifier(n_neighbors=5)
knn.fit(X_train, y_train)

# Predict NaN Job Levels
predicted_levels = knn.predict(X_predict)

# Display the predictions before overwriting the original DataFrame
predict_df['Predicted Job Level'] = predicted_levels
print("Predicted Job Levels for records with NaN Job Level:")
display(predict_df[['Age', 'Gender', 'Current Role', 'Predicted Job Level']])


# In[21]:


#Update Job Level field with predictions and verify changes

# Update the Job Level in the original DataFrame where it was NaN
filtered_us_dev_survey_values_df.loc[filtered_us_dev_survey_values_df['Job Level'].isna(), 'Job Level'] = predicted_levels

# Verify that no records remain with NaN Job Level
remaining_nan_job_levels = filtered_us_dev_survey_values_df['Job Level'].isna().sum()
print(f"Number of records with NaN Job Level after updating: {remaining_nan_job_levels}")

# Display the cleaned DataFrame
display(filtered_us_dev_survey_values_df.head(5))


# #### Seaborn Strip Plot: Distribution of Job Levels by Gender within Each Age Group
# **Purpose:** Visualize the distribution of job levels across genders within each age group in the technology sector, providing insights into potential patterns of gender representation and career progression at different stages of professional life.
# 
# **Insights:**
# - **25–34 Age Group:** Men and women are represented at all job levels from New grad to Engineering Manager. However, Founder/CEO/CTO roles are exclusively occupied by men, and multiple data points at this level indicate that this is not an isolated outlier. Additionally, men appear to have continuous representation across all job levels, while higher levels for women, such as Principal Engineer, Architect, and Engineering Manager, display distinct individual points. This suggests lower representation for women in higher roles at this career stage.
# - **35–44 Age Group:** Both genders are represented across all job levels, including Founder/CEO/CTO, showing relatively balanced representation in this age range. This distribution may indicate that women who reach this age group achieve a broader range of job levels than in earlier career stages.
# - **45–54 Age Group:** There is a noticeable reduction in representation, particularly for women. In this group, women are clustered around Level 1 Developer (junior) and Senior Developer, with only a single outlier at Engineering Manager. Men continue to show representation across more levels, indicating a possible disparity in advancement opportunities as careers progress.
# - **55–64 Age Group:** Men are represented across several senior roles, from Senior Developer to Engineering Manager. Women, however, are predominantly clustered at Senior Developer, with sparse representation at Level 1 Developer (junior), Architect, and Engineering Manager. This pattern may reflect limited opportunities for women to progress to higher roles at later career stages.
# 
# **Summary:** The visualization highlights potential gender-based differences in career progression. Women in the 25–34 age group appear to be underrepresented in higher roles, with the disparity becoming more pronounced in later career stages (45–54 and 55–64 age groups). This analysis suggests that women in tech may face challenges in advancing to senior positions as they progress in their careers, potentially reflecting systemic barriers to higher-level roles.

# In[22]:


# Define the desired order for job levels and age groups
job_level_order = [
    "Student", "New grad", "Freelancer", "Level 1 developer (junior)", 
    "Senior developer", "Principal engineer", "Architect", 
    "Engineering manager", "Founder / CEO / CTO"
]
age_order = ["25 - 34 years old", "35 - 44 years old", "45 - 54 years old", "55 - 64 years old"]

# Ensure Job Level column is ordered categorically
filtered_us_dev_survey_values_df['Job Level'] = pd.Categorical(
    filtered_us_dev_survey_values_df['Job Level'], categories=job_level_order, ordered=True
)
filtered_us_dev_survey_values_df['Gender'] = pd.Categorical(
    filtered_us_dev_survey_values_df['Gender'], categories=["Male", "Female"], ordered=True
)

# Define a custom palette to match matplotlib graphs
custom_palette = ['slategray', 'lightseagreen']

# Create the catplot
g = sns.catplot(
    data=filtered_us_dev_survey_values_df,
    x="Gender", y="Job Level", col="Age", hue="Gender", 
    kind="strip", col_order=age_order, hue_order=["Male", "Female"],
    palette=custom_palette, dodge=True, jitter=True, height=6, aspect=0.8
)

# Set title for each plot
g.set_titles("Age Group: {col_name}")

# Remove x-axis label
g.set(xlabel="")  

# Overall title for the plot
plt.subplots_adjust(top=0.9)
g.figure.suptitle("Distribution of Job Levels by Gender within Each Age Group", fontsize=16)

plt.show()


# #### Import Dataset 2: 2017 Pew Research Center STEM Survey

# In[23]:


# import zip file from Pew Research
file_handle, _ = urlretrieve("https://www.pewresearch.org/wp-content/uploads/sites/20/2019/04/2017-Pew-Research-Center-STEM-survey.zip")
zipfile = ZipFile(file_handle, "r")
zipfile.extractall("./data")
zipfile.close()


# #### Examine contents of .sav file

# In[24]:


file_path = 'data/materials for public release/2017 Pew Research Center STEM survey.sav'

# Read the .sav file into a DataFrame
prc_stem_df, meta = pyreadstat.read_sav(file_path)

# Display basic information about the DataFrame
print(prc_stem_df.info())

# Display the first few rows of the DataFrame
print(prc_stem_df.head())

print(prc_stem_df.tail())


# #### Read the .docx file
# - Read the Pew Research Center files associated with the .sav file and convert them into .txt files to understand the codes used.

# In[25]:


# Load the Questionnaire document
questionnaire_path = './data/materials for public release/Questionnaire - 2017 Pew Research Center STEM Survey.docx'
questionnaire_doc = Document(questionnaire_path)

# Extract text from the Questionnaire document
prc_questions = []  
for paragraph in questionnaire_doc.paragraphs:
    prc_questions.append(paragraph.text)

# Join all paragraphs into a single string for the questionnaire
prc_questions_text = "\n".join(prc_questions)

# Define the path for saving the questionnaire .txt file
questionnaire_txt_file_path = './data/materials for public release/prc_questionnaire.txt'

# Save the full text to a .txt file
with open(questionnaire_txt_file_path, 'w', encoding='utf-8') as f:
    f.write(prc_questions_text)

print(f"The questionnaire has been extracted and saved to '{questionnaire_txt_file_path}'.")

# Preview the first few lines of the extracted questionnaire
print("Preview of the extracted questionnaire:")
print(prc_questions_text[:400])

# Load the Codebook document
codebook_path = './data/materials for public release/Codebook - 2017 Pew Research Center STEM Survey.docx'
codebook_doc = Document(codebook_path)

# Extract text from the Codebook document
prc_codebook = []  
for paragraph in codebook_doc.paragraphs:
    prc_codebook.append(paragraph.text)

# Join all paragraphs into a single string for the codebook
prc_codebook_text = "\n".join(prc_codebook)

# Define the path for saving the codebook .txt file
codebook_txt_file_path = './data/materials for public release/prc_codebook.txt'

# Save the full text to a .txt file
with open(codebook_txt_file_path, 'w', encoding='utf-8') as f:
    f.write(prc_codebook_text)

print(f"The codebook has been extracted and saved to '{codebook_txt_file_path}'.")

# Preview the first few lines of the extracted codebook
print("Preview of the extracted codebook:")
print(prc_codebook_text[:400])


# #### Display All Column Names in Dataframe
# - Reading the column names with the new context of the Questionnaire and Codebook file will help to determine which columns are needed for analysis

# In[26]:


# Display all column names in the DataFrame
print("Column names in the Pew Research dataset:")
for col in prc_stem_df.columns:
    print(col)


# #### Convert CaseID from float to int64

# In[27]:


# Convert 'CaseID' to int64
prc_stem_df['CaseID'] = prc_stem_df['CaseID'].astype('int64')

# Confirm the change
print("\nData types after conversion:")
print(prc_stem_df.dtypes)


# #### Exploratory Data Analysis (EDA)
# - Search for null values or refused responses in fields required for analysis.

# In[28]:


# Summary of counts for each value in WORK_1
employed_full_time = prc_stem_df['WORK_1'].value_counts(dropna=False)
print("Summary of counts for each value in WORK_1:")
print(employed_full_time)

# Summary of counts for each value in WORK_2
employed_part_time = prc_stem_df['WORK_2'].value_counts(dropna=False)
print("\nSummary of counts for each value in WORK_2:")
print(employed_part_time)

# Summary of counts for each value in WORK_3
self_employed_full_time = prc_stem_df['WORK_3'].value_counts(dropna=False)
print("\nSummary of counts for each value in WORK_3:")
print(self_employed_full_time)

# Summary of counts for each value in WORK_4
self_employed_part_time = prc_stem_df['WORK_4'].value_counts(dropna=False)
print("\nSummary of counts for each value in WORK_4:")
print(self_employed_part_time)

# Summary of counts for each value in EMPLOYED
employment_status = prc_stem_df['EMPLOYED'].value_counts(dropna=False)
print("\nSummary of counts for each value in EMPLOYED:")
print(employment_status)

# Summary of counts for each value in FULLPART
employment_full_part = prc_stem_df['FULLPART'].value_counts(dropna=False)
print("\nSummary of counts for each value in FULLPART:")
print(employment_full_part)

# Summary of counts for each value in SELFEMPLOYED
self_employment = prc_stem_df['SELFEMPLOYED'].value_counts(dropna=False)
print("\nSummary of counts for each value in SELFEMPLOYED:")
print(self_employment)

# Summary of counts for each value in ppagecat
prc_age = prc_stem_df['ppagecat'].value_counts(dropna=False)
print("\nSummary of counts for each value in ppagecat:")
print(prc_age)

# Summary of counts for each value in PPGENDER
prc_gender = prc_stem_df['PPGENDER'].value_counts(dropna=False)
print("\nSummary of counts for each value in PPGENDER:")
print(prc_gender)

# Summary of counts for each value in PPGENDER
prc_stem_degree = prc_stem_df['STEM_DEGREE'].value_counts(dropna=False)
print("\nSummary of counts for each value in STEM_DEGREE:")
print(prc_stem_degree)


# In[29]:


# Convert specified columns to int64
columns_to_convert = [
    'WORK_1', 'WORK_2', 'WORK_3', 'WORK_4', 
    'EMPLOYED', 'FULLPART', 'SELFEMPLOYED'
]

# Convert to int64
for column in columns_to_convert:
    prc_stem_df[column] = prc_stem_df[column].astype('int64')

# Drop records with 9 for WORK_1, WORK_2, WORK_3, WORK_4, EMPLOYED, FULLPART, SELFEMPLOYED
prc_stem_df = prc_stem_df[~prc_stem_df[columns_to_convert].isin([9]).any(axis=1)]

# For ppagecat, drop 1, 6, and 7
prc_stem_df = prc_stem_df[~prc_stem_df['ppagecat'].isin([1, 6, 7])]

# Confirming the changes
print(f"Remaining records after filtering: {prc_stem_df.shape[0]}")
print("\nUpdated DataFrame info:")
print(prc_stem_df.info())


# #### Refine DataFrame for Pew Research Center STEM Survey
# - Eliminate fields not needed for hypothesis
# - Rename columns to more meaningful names

# In[30]:


# List of columns to exclude
columns_to_exclude = [
    'SCH2a', 'SCH2b', 'SCH3a', 'SCH3b', 'SCH3c', 'SCH3d',
    'SCH4', 'SCH5a', 'SCH5b', 'SCH5c', 'SCH6a', 'SCH6b',
    'SCH6c', 'SCH6d', 'SCH6e', 'SCH6f', 'SCH6g', 'SCH6h',
    'SCH7', 'SCH8a', 'SCH8b', 'SCH9a', 'SCH9b', 'SCH10_flag',
    'SCH10A_1', 'SCH10A_2', 'SCH10A_3', 'SCH10A_4', 'SCH10A_5',
    'SCH10A_6', 'SCH10A_Refused', 'SCH10B_1', 'SCH10B_2',
    'SCH10B_3', 'SCH10B_4', 'SCH10B_5', 'SCH10B_6', 'SCH10B_Refused',
    'STEMJOBa', 'STEMJOBb', 'STEMJOBc', 'STEMJOBd', 'STEMJOBe',
    'STEMJOBf', 'STEMJOBg', 'STEMJOBh', 'REASON2a', 'REASON2b',
    'REASON2c', 'REASON2d', 'REASON2e', 'REASON2f', 'REASON2g',
    'ETHNJOB2', 'ETHNJOB2_OE1_col', 'ETHNJOB2_OE2_col', 
    'ETHNJOB2_OE3_col', 'RACE_col', 'SCICOUR2_t',
    'MATHCOUR2_t', 'PPT017_t', 'PPT18OV_t', 'PPHHSIZE_t', 
    'ETHN1', 'ETHN2a', 'ETHN2b', 'ETHN2c', 'ETHN2d', 
    'ETHN3a', 'ETHN3b', 'ETHN3c', 'ETHN3d', 'ETHN4', 
    'ETHN5', 'ETHN6_a', 'ETHN6_b', 'ETHN6_c', 'ETHN6_d',
    'ETHN6_Refused', 'ETHNDISC_a', 'ETHNDISC_b', 'ETHNDISC_c',
    'ETHNDISC_d', 'ETHNDISC_e', 'ETHNDISC_f', 'ETHNDISC_g',
    'ETHNDISC_h', 'ETHNDISC_i', 'ETHNDISC_Refused', 'FAMSTEM1',
    'FAMSTEM2_1', 'FAMSTEM2_2', 'FAMSTEM2_Refused', 'INTEREST1',
    'TECH4', 'TECH5', 'TECH6', 'ETHNJOB1', 'PARTY', 'PARTYLN',
    'IDEO', 'PUBLIC', 'DOV_FORM', 'PPHHHEAD', 'HH_INCOME_col',
    'PPMARIT', 'PPMSACAT', 'SCH1_OE1_col', 'SCH1_OE2_col',
    'SCH1_OE3_col', 'SCH7_OE1_col', 'INTEREST2_OE', 'INTEREST3_OE1_col',
    'INTEREST3_OE2_col', 'INTEREST3_OE3_col', 'XSPANISH', 'Xparent',
    'PPETHM', 'FIRSTCOLL', 'SCICOUR1', 'MATHCOUR1'
]

# Create a new DataFrame excluding the specified columns
pew_research_numeric = prc_stem_df.loc[:, ~prc_stem_df.columns.isin(columns_to_exclude)].copy()

# Create a dictionary for renaming specified columns
columns_to_rename = {
    'WORK_1': 'Employed Full-Time by Company',
    'WORK_2': 'Employed Part-Time by Company',
    'WORK_3': 'Self-Employed Full-Time',
    'WORK_4': 'Self-Employed Part-Time',
    'EMPLOYED': 'Employment Status',
    'FULLPART': 'Employee Type',
    'SELFEMPLOYED': 'Self-Employment Status',
    'OCCUPATION_col': 'Occupation',
    'INDUSTRY_col': 'Industry',
    'TEACHSTEM': 'STEM Teacher Y/N',
    'WORKTYPE_FINAL': 'STEM Worker Y/N',
    'EDUC4CAT': 'Education Level Categorical',
    'RECONA_col': 'Computer Work Y/N',
    'RECONB_col': 'Engineer Y/N',
    'RECONC_col': 'Science Worker Type',
    'STEM_DEGREE': 'STEM Degree Y/N',
    'JOBVALU1_1': 'Job Choice - High Pay',
    'JOBVALU1_2': 'Job Choice - Work-Life Balance',
    'JOBVALU1_3': 'Job Choice - Advancement Opportunities',
    'JOBVALU1_4': 'Job Choice - Contribution to Society',
    'JOBVALU1_5': 'Job Choice - Respect of Others',
    'JOBVALU1_6': 'Job Choice - Helping Others',
    'JOBVALU1_7': 'Job Choice - Welcoming Environment',
    'JOBVALU1_8': 'Job Choice - None of the Above',
    'JOBVALU2': 'Job Choice - Most Important',
    'AHEADa': 'Employment Advancement - Assertiveness',
    'AHEADb': 'Employment Advancement - Socializing with Co-Workers',
    'AHEADc': 'Employment Advancement - Point Out Workplace Problems',
    'AHEADd': 'Employment Advancement - Workplace Mentor',
    'AHEADe': 'Employment Advancement - Sharing Personal Life',
    'AHEADf': 'Employment Advancement - Working Harder',
    'AHEADg': 'Employment Advancement - Point Out Personal Accomplishments',
    'TALENT': 'Employment Advancement - Natural Ability',
    'PROVE': 'Workplace Respect - Need to Prove Oneself',
    'RESPECTA': 'Workplace Respect - Valued by Supervisor',
    'RESPECTB': 'Workplace Respect - Valued by Co-Workers',
    'JOBVALU1_Refused': 'Job Choice - Refused to Answer',
    'REASON1a': 'Gender Exclusion - Lack of Early Encouragement',
    'REASON1b': 'Gender Exclusion - Lack of Belief in Success',
    'REASON1c': 'Gender Exclusion - Lack of Female Role Models',
    'REASON1d': 'Gender Exclusion - Discrimination in Recruitment Hiring and Promotion',
    'REASON1e': 'Gender Exclusion - Slow Increase in Female Representation',
    'REASON1f': 'Gender Exclusion - Lower Interest in STEM',
    'REASON1g': 'Gender Exclusion - Work-Family Balance Challenges',
    'TECH1': 'Descrimination Perception - Awareness',
    'TECH2': 'Descrimination Perception - Industry Specific',
    'TECH3': 'Descrimination Perception - Problem Severity',
    'GEND1': 'Respondent Workplace - Gender Balance',
    'GEND2': 'Respondent Workplace - Recruitment and Hiring',
    'GEND3': 'Respondent Workplace - Advancement Opportunities',
    'GEND4': 'Respondent Workplace - Gender Diversity Initiatives',
    'GEND5': 'Respondent Workplace - Gender Diversity Importance',
    'GEND6_a': 'Gender Diversity Reason - Equal Opportunity',
    'GEND6_b': 'Gender Diversity Reason - Contributes to Success',
    'GEND6_c': 'Gender Diversity Reason - Expands Workforce Supply',
    'GEND6_d': 'Gender Diversity Reason - None of the Above',
    'GEND6_Refused': 'Gender Diversity Reason - Refused to Answer',
    'GENDJOB1': 'Gender Effect on Job Success',
    'GENDDISC_a': 'Personal Experience - Denied Promotion',
    'GENDDISC_b': 'Personal Experience - Earned Less than Opposite Gender Counterpart',
    'GENDDISC_c': 'Personal Experience - Denied Job',
    'GENDDISC_d': 'Personal Experience - Felt Isolated',
    'GENDDISC_e': 'Personal Experience - Received Less Senior Support than Opposite Gender Counterpart', 
    'GENDDISC_f': 'Personal Experience - Treated as Incompetent',
    'GENDDISC_g': 'Personal Experience - Denied Important Assignments',
    'GENDDISC_h': 'Personal Experience - Microaggressions',
    'GENDDISC_i': 'Personal Experience - None of These',
    'GENDDISC_Refused': 'Personal Experience - Refused to Answer',
    'PPCM1301_col': 'Employer Type',
    'PPGENDER': 'Gender',
    'PPREG4': 'US Region',
    'ppagecat': 'Age',
    'ppagect4': 'Broad Age Range',
    'DEGREE1_computer': 'College Major - Computer and Information Sciences', 
    'DEGREE1_math': 'College Major - Mathematics and Statistics',
    'DEGREE1_life': 'College Major - Life Sciences or Agriculture',
    'DEGREE1_physical': 'College Major - Physical or Earth Sciences',
    'DEGREE1_engineering': 'College Major - Engineering or Architecture',
    'DEGREE1_health': 'College Major - Health-Related',
    'DEGREE1_otherstem': 'College Major - Other STEM-related Degree',
    'DEGREE1_othernonstem': 'College Major - Non-STEM Degree',
    'DEGREE2_computer': 'Graduate Degree - Computer and Information Sciences',
    'DEGREE2_math': 'Graduate Degree - Mathematics and Statistics',
    'DEGREE2_life': 'Graduate Degree - Life Sciences or Agriculture',
    'DEGREE2_physical': 'Graduate Degree - Physical or Earth Sciences',
    'DEGREE2_engineering': 'Graduate Degree - Engineering or Architecture',
    'DEGREE2_health': 'Graduate Degree - Health-Related',
    'DEGREE2_otherstem': 'Graduate Degree - Other STEM-related Degree',
    'DEGREE2_othernonstem': 'Graduate Degree - Non-STEM Degree',
    'HARASS1': 'Harassment - Personal Workplace Awareness',
    'HARASS2': 'Harassment - Industry Awareness',
    'HARASS3': 'Harassment - Personal Experience',
    'VOTECH': 'Background - Vocational or Technical Training',
    'RELATE1': 'Job Related to College Major Y/N',
    'RELATE2': 'Use of College Major Skills Y/N',
    'PPCM0166': 'Company Size',
    'GENDJOB2_OE1_col': 'Gender Impact - Barrier to Success (First Mention)',
    'GENDJOB2_OE2_col': 'Gender Impact - Barrier to Success (Second Mention)',
    'GENDJOB2_OE3_col': 'Gender Impact - Barrier to Success (Third Mention)'
}

# Rename the columns as specified
pew_research_numeric.rename(columns=columns_to_rename, inplace=True)


# Display the new DataFrame info to confirm changes
print("\nUpdated DataFrame info:")
print(pew_research_numeric.info())

print("Column names in the Pew Research Numeric dataset:")
for col in pew_research_numeric.columns:
    print(col)


# **Exploratory Data Analyis on Pew Research Center dataframe**
# - Check for employees who marked that they are self-employed and also employed by a company
#     - Respondents who are only self-employed are not useful to my hypothesis so they can be filtered out from the dataset
# 

# In[31]:


# Count of respondents who are both self-employed part-time and employed by a company
self_employed_and_company_employed_count = pew_research_numeric[
    (pew_research_numeric['Self-Employed Part-Time'] == 1) &
    ((pew_research_numeric['Employed Full-Time by Company'] == 1) |
     (pew_research_numeric['Employed Part-Time by Company'] == 1))
].shape[0]

# Count of respondents who are self-employed part-time and not employed by a company
self_employed_full_time_only_count = pew_research_numeric[
    (pew_research_numeric['Self-Employed Part-Time'] == 1) &
    ((pew_research_numeric['Employed Full-Time by Company'].isin([2, 9])) &
     (pew_research_numeric['Employed Part-Time by Company'].isin([2, 9])))
].shape[0]

print("Count of respondents who are both self-employed part-time and employed by a company:", self_employed_and_company_employed_count)
print("Count of respondents who are self-employed part-time and not employed by a company:", self_employed_full_time_only_count)


# In[32]:


# Count of respondents who are both self-employed full-time and employed by a company
self_employed_and_company_employed_count = pew_research_numeric[
    (pew_research_numeric['Self-Employed Full-Time'] == 1) &
    ((pew_research_numeric['Employed Full-Time by Company'] == 1) |
     (pew_research_numeric['Employed Part-Time by Company'] == 1))
].shape[0]

# Count of respondents who are self-employed full-time and not employed by a company
self_employed_full_time_only_count = pew_research_numeric[
    (pew_research_numeric['Self-Employed Full-Time'] == 1) &
    ((pew_research_numeric['Employed Full-Time by Company'].isin([2, 9])) &
     (pew_research_numeric['Employed Part-Time by Company'].isin([2, 9])))
].shape[0]

print("Count of respondents who are both self-employed full-time and employed by a company:", self_employed_and_company_employed_count)
print("Count of respondents who are self-employed full-time and not employed by a company:", self_employed_full_time_only_count)


# In[33]:


# Filter to retain only records where the respondent is employed by a company, has Employment Status = 1, and is a STEM Worker
pew_research_stem_employed_filtered = pew_research_numeric[
    ((pew_research_numeric['Employed Full-Time by Company'] == 1) |
     (pew_research_numeric['Employed Part-Time by Company'] == 1)) &
    (pew_research_numeric['Employment Status'] == 1) &
    (pew_research_numeric['STEM Worker Y/N'] == 1)
]

# Display the count of remaining records after filtering
remaining_stem_employed_count = pew_research_stem_employed_filtered.shape[0]
remaining_stem_employed_count


# **Find values for "Job Choice" columns**
# - The Job Choice columns were coded with numbers in the original survey as a single column. I need to know if the split columns have 0 and 1 or 1 and 2 as Yes/No representation to create visualizations.

# In[34]:


# List of job choice columns to check value counts
job_choice_columns = [
    'Job Choice - High Pay',
    'Job Choice - Work-Life Balance',
    'Job Choice - Advancement Opportunities',
    'Job Choice - Contribution to Society',
    'Job Choice - Respect of Others',
    'Job Choice - Helping Others',
    'Job Choice - Welcoming Environment',
]

# Display value counts for each column in the job choice list
for column in job_choice_columns:
    print(f"Value counts for '{column}':")
    print(pew_research_stem_employed_filtered[column].value_counts(dropna=False))
    print("\n")


# #### Matplotlib Stacked Bar Chart: Job Choice Priorities by Gender (Normalized)
# **Purpose:** Recreate the HackerRank survey analysis with Pew Research Center data to validate findings on job choice priorities by gender in the technology and STEM fields. This chart has been both weighted and normalized to reflect the gender imbalance in respondents and the demographic adjustments applied in the Pew Research Center dataset.
# 
# **Insights:**
# - **Work-Life Balance as Top Priority:** Both women (76.6%) and men (72.5%) consider *Work-Life Balance* crucial when choosing a job, consistent with the HackerRank survey. This emphasizes the importance of flexible work environments in addressing the demands of professional and personal life balance for both genders.
# - **Higher Emphasis on Helping Others and Contribution to Society by Women:** Women in the Pew Research Center data place significantly more value on *Helping Others* (57.9% for women vs. 28.5% for men) and *Contribution to Society* (60.4% for women vs. 49.5% for men). In contrast, the HackerRank survey reported lower ratings for *Company Mission* and *Product Impact* for both genders, suggesting that while women in the Pew sample prioritize purpose-driven work, this emphasis was not as pronounced in the HackerRank survey. This could reflect differing values or motivations across these two groups of respondents.
# - **Company Culture as a Key Factor:** *Company Culture* remains a higher priority for women (53.2%) compared to men (48.3%)in the Pew Research Center data, echoing findings from the HackerRank survey (44.2% for women vs. 39.0% for men). This preference underlines the importance of inclusive, supportive work environments for women in STEM fields.
# - **Compensation as a Higher Priority for Men:** *Compensation* is more highly prioritized by men than women in both surveys, with Pew Research Center data showing 60.3% for men vs. 47.6% for women, and HackerRank showing 56.9% for men vs. 41.2% for women. This consistent trend suggests that, although compensation is important, it may not be the foremost factor for women in STEM.
# - **Contrasting Trends in Professional Growth Priorities:** The *Professional Growth* findings between the two surveys reveal nearly opposite trends. In the HackerRank survey, women rated *Professional Growth and Learning* higher than men (54.6% for women vs. 44.9% for men). However, in the Pew Research Center data, men prioritized *Professional Growth* more (57.5% for men vs. 46.3% for women). This discrepancy may indicate varied career motivations between the samples, with HackerRank respondents possibly focusing more on career advancement opportunities for women than their counterparts in the Pew Research Center survey.
# 
# **Methodology Consideration:** It is important to note a key difference in methodology: the HackerRank survey required respondents to select exactly three job priorities, whereas the Pew Research Center survey allowed respondents to choose any number of priorities or none at all. This flexibility in the Pew Research Center survey may have influenced the distribution of choices, allowing for a broader expression of priorities compared to the forced-choice format of the HackerRank survey. Additionally, the Pew data incorporates a weighting factor to adjust for demographic representation, adding rigor to the analysis but potentially influencing comparisons with the HackerRank data.
# 
# **Summary:** This analysis of the Pew Research Center data generally supports the trends observed in the HackerRank survey, particularly around work-life balance and company culture. However, the contrasting findings for professional growth and the emphasis on helping others and societal contribution underscore that motivations can vary widely across STEM fields. These insights emphasize the importance of creating flexible, purpose-driven, and growth-oriented workplaces to effectively retain a diverse workforce in technology and STEM.

# In[35]:


# Define a dictionary mapping the Pew Research column names to display names
column_display_names_pew = {
    'Job Choice - Contribution to Society': 'Contribution to Society',
    'Job Choice - Helping Others': 'Helping Others',
    'Job Choice - Welcoming Environment': 'Company Culture',
    'Job Choice - Work-Life Balance': 'Work-Life Balance',
    'Job Choice - High Pay': 'Compensation',
    'Job Choice - Respect of Others': 'Respected Job',
    'Job Choice - Advancement Opportunities': 'Professional Growth'
}

# Select only the job choice columns to analyze
job_choice_columns = list(column_display_names_pew.keys())

# Multiply each job choice column by the weight column to apply weighting
weighted_data = pew_research_stem_employed_filtered.copy()
for col in job_choice_columns:
    weighted_data[col] = weighted_data[col] * weighted_data['weight']

# Group data by gender and sum up the weighted counts for each job choice criterion
weighted_criteria_counts_pew = weighted_data.groupby('Gender')[job_choice_columns].sum()

# Calculate the total weighted count by gender to normalize the data
total_weighted_counts_by_gender_pew = weighted_data.groupby('Gender')['weight'].sum()

# Normalize the data by calculating the percentage of each job criterion per gender
criteria_percentages_pew_weighted = weighted_criteria_counts_pew.div(total_weighted_counts_by_gender_pew, axis=0) * 100

# Rename gender values for readability
criteria_percentages_pew_weighted.index = criteria_percentages_pew_weighted.index.map({1: 'Male', 2: 'Female'})

# Rename columns in the DataFrame for plotting
criteria_percentages_pew_weighted = criteria_percentages_pew_weighted.rename(columns=column_display_names_pew)

# Plotting the weighted normalized stacked bar chart
ax = criteria_percentages_pew_weighted.T.plot(kind='bar', stacked=True, figsize=(14, 8), color=['slategray', 'lightseagreen'], edgecolor='none')
plt.suptitle('Job Choice Priorities by Gender (Weighted and Normalized)')
plt.title('Pew Research Center Data with Applied Weighting', fontsize=12, color='gray')
plt.ylabel('Percentage of Respondents (%)')
plt.legend(title='Gender')
plt.xticks(rotation=45, ha='right')
plt.tight_layout()

# Adding percentage labels to each section of the bars
for container in ax.containers:
    ax.bar_label(container, fmt='%.1f%%', label_type='center')

plt.show()


# #### Seaborn Heatmap: Concentration of Discrimination Experiences by Age Group and Gender (Weighted Counts)
# **Purpose:** This heatmap visualizes the concentration of reported discrimination experiences across age groups and gender, using weighted counts from the Pew Research Center data. By applying weighting factors, this chart provides a more representative view of how these experiences vary across demographics, especially by age and gender.
# 
# **Insights:**
# - **High Concentration of Negative Experiences for Younger Women:** Women aged 25–34 report notably high weighted counts in experiences such as *Earning Less than Opposite Gender Counterpart* (21.5), *Experiencing Microaggressions* (18.2), and *Being Treated as Incompetent* (22.8). These elevated counts suggest significant barriers for younger women in the workforce, particularly in terms of respect and equitable treatment. 
#   
# - **Persisting Challenges for Women Aged 35–44:** In the 35–44 age group, women continue to report substantial discrimination experiences, including *Earning Less* (14.8), *Microaggressions* (10.4), *Receiving Less Senior Support* (11.5), and *Being Treated as Incompetent* (13.1). These experiences highlight ongoing barriers to career advancement and the importance of mentorship and support in the workplace. The need for a supportive company culture may be particularly strong in this group as they strive for professional growth and equity.
# 
# - **Decreasing but Consistent Reports in Later Career Stages:** For women aged 45–54, notable experiences include *Earning Less* (11.3) and *Being Treated as Incompetent* (11.3), while women aged 55–64 still report *Earning Less* (14.0) and *Being Treated as Incompetent* (10.2). Although these counts are slightly lower than in younger age groups, they indicate that discriminatory treatment persists throughout many women's careers, possibly impacting their satisfaction and retention in STEM fields.
# 
# - **Comparative Low Reports for Men Across All Age Groups:** Men report much lower weighted counts for these discrimination experiences across all age groups, suggesting they encounter these issues less frequently than women. This difference further highlights why women, compared to men, may prioritize company culture as a significant factor when evaluating job opportunities.
# 
# **Methodology Consideration:** This visualization uses weighted counts, which adjust each reported experience according to the demographic representation in the Pew Research Center dataset. These adjustments help reflect the actual prevalence of discrimination experiences in the workforce more accurately.
# 
# **Summary:** The analysis reveals that younger and mid-career women (ages 25–44) encounter the highest concentration of discrimination experiences, especially in terms of earning disparities, microaggressions, and perceived incompetence. These trends emphasize the need for supportive and equitable workplace cultures, as these factors likely influence why women place high value on company culture when considering job opportunities in STEM fields.

# In[36]:


# Define the age group labels only for ages 25-34, 35-44, 45-54, and 55-64
age_group_labels = {
    2: "Age 25-34", 
    3: "Age 35-44", 
    4: "Age 45-54", 
    5: "Age 55-64"
}

# Define the mapping for renaming experience labels by removing "Personal Experience - "
experience_columns = [
    'Personal Experience - Denied Important Assignments',
    'Personal Experience - Denied Job',
    'Personal Experience - Denied Promotion',
    'Personal Experience - Earned Less than Opposite Gender Counterpart',
    'Personal Experience - Felt Isolated',
    'Personal Experience - Microaggressions',
    'Personal Experience - Received Less Senior Support than Opposite Gender Counterpart',
    'Personal Experience - Treated as Incompetent'
]
experience_labels = {col: col.replace("Personal Experience - ", "") for col in experience_columns}

# Select relevant columns and apply weighting
weighted_data = pew_research_stem_employed_filtered.copy()
for col in experience_columns:
    weighted_data[col] = weighted_data[col] * weighted_data['weight']

# Group by age and gender, summing weighted experiences
weighted_experience_counts = weighted_data.groupby(['Age', 'Gender'])[experience_columns].sum().reset_index()

# Map age groups and rename experience columns
weighted_experience_counts['Age'] = weighted_experience_counts['Age'].map(age_group_labels)
weighted_experience_counts = weighted_experience_counts.rename(columns=experience_labels)

# Melt the DataFrame for heatmap plotting
melted_data = weighted_experience_counts.melt(id_vars=['Age', 'Gender'], var_name='Discrimination Experience', value_name='Weighted Count')

# Pivot for heatmap
heatmap_data = melted_data.pivot_table(
    index=['Age', 'Discrimination Experience'], 
    columns='Gender', 
    values='Weighted Count', 
    fill_value=0
)

# Rename the Gender values for clarity
heatmap_data.columns = heatmap_data.columns.map({1: 'Male', 2: 'Female'})

# Plotting the heatmap
plt.figure(figsize=(16, 14))
sns.heatmap(heatmap_data, cmap='coolwarm', annot=True, fmt='.3f', cbar_kws={'label': 'Weighted Count of Experiences'})
plt.title('Concentration of Discrimination Experiences by Age Group and Gender: Pew Research Center Data with Applied Weighting')
plt.xlabel('')
plt.ylabel('Age Group and Discrimination Experience')
plt.show()


# #### Import Dataset 3: National Center for Science and Engineering Statistics (NCSES)
# - Includes demographic breakdown of STEM participation in the workforce from 1993 - 2019
# - Data was compiled by the NCSES from the U.S. Census Bureau, American Community Survey, National Center for Science and Engineering Statistics, and more
# - For the full list of compiled sources: https://ncses.nsf.gov/pubs/nsb20212/data#source-block 

# In[37]:


# scrape HTML file to extract tables
os.makedirs("data/ncses", exist_ok=True) # create new subfolder for tables scraped from website

url = "https://ncses.nsf.gov/pubs/nsb20212/participation-of-demographic-groups-in-stem"
page = requests.get(url)
soup = BeautifulSoup(page.content, "html.parser") # parse the html as a string

tables = soup.find_all("table")


# loop through each table to extract the title and data
for i in range(len(tables)):
    table = tables[i]

    # Extract the table's title attribute and clean it up for a filename
    title = table.get('title')  # Find the title attribute in the table

    if title:
        title = title.strip()  # Use the title if it exists
    else:
        title = f'table_{i}'  # Generic name if the table does not have a title
    
    # Clean title: replace spaces with hyphens and remove special characters except apostrophes
    title = title.replace("'", "")  # Remove apostrophes
    title = re.sub(r"[^\w\s]", "-", title)  # Replace special characters with hyphens
    title = title.replace(" ", "-")  # Replace spaces with hyphens

    # Truncate the title if it exceeds 50 characters for filename compatibility
    if len(title) > 50:
        title = title[:50]

    file_name = f"data/ncses/{title}.csv" # generate an empty csv for each table

    # convert each table to a pandas df
    df = pd.read_html(str(table))[0]

    # save the pandas df as a csv
    df.to_csv(file_name, index=False)


# **Import Additional Resources From National Center for Science and Engineering Statistics (NCSES)**
# - The Report titled *"The STEM Labor Force of Today: Scientists, Engineers, and Skilled Technical Workers"* spans several pages and has supplemental tables that are not included on any of the pages. 

# In[38]:


# import data-tables zip file from NCSES
file_handle, _ = urlretrieve("https://ncses.nsf.gov/pubs/nsb20212/assets/nsb20212-report-tables-excels.zip")
zipfile = ZipFile(file_handle, "r")
zipfile.extractall("./data/ncses")
zipfile.close()

# import data-figures zip file from NCSES
file_handle, _ = urlretrieve("https://ncses.nsf.gov/pubs/nsb20212/assets/nsb20212-report-figures-excels.zip")
zipfile = ZipFile(file_handle, "r")
zipfile.extractall("./data/ncses")
zipfile.close()

# import supplemental tables zip file from NCSES
file_handle, _ = urlretrieve("https://ncses.nsf.gov/pubs/nsb20212/assets/supplemental-tables/nsb20212-supplemental-tables-figures-tables-excels.zip")
zipfile = ZipFile(file_handle, "r")
zipfile.extractall("./data/ncses")
zipfile.close()


# **Convert relevant xlsx files into pandas dataframes**

# In[39]:


# Define the path to the file: Table LBR-7 - Women with a bachelor's degree or above, by broad occupational group and highest degree: 1993, 2003, 2019
file_path = './data/ncses/nsb20212-tablbr-007.xlsx'

# Define start rows for each section based on the Excel file structure
start_row_degrees = 9  
start_row_occupations = 6  
num_rows_degrees = 2  # Number of rows for the degree section
num_rows_occupations = 2  # Number of rows for the occupation section

# Load the Degree Focus data
degree_focus_df = pd.read_excel(
    file_path,
    skiprows=start_row_degrees,
    nrows=num_rows_degrees,
    names=["Category", "1993 - Count", "2003 - Count", "2019 - Count", "1993 - Percent", "2003 - Percent", "2019 - Percent"]
)

# Load the Occupational Group data
occupational_group_df = pd.read_excel(
    file_path,
    skiprows=start_row_occupations,
    nrows=num_rows_occupations,
    names=["Category", "1993 - Count", "2003 - Count", "2019 - Count", "1993 - Percent", "2003 - Percent", "2019 - Percent"]
)

# Add a column to indicate whether the data is for degrees or occupations
degree_focus_df['Type'] = 'Degree'
occupational_group_df['Type'] = 'Occupation'

# Convert "Thousands" columns to actual numbers by multiplying by 1,000
for col in ["1993 - Count", "2003 - Count", "2019 - Count"]:
    degree_focus_df[col] = degree_focus_df[col] * 1000
    occupational_group_df[col] = occupational_group_df[col] * 1000

# Concatenate both DataFrames to have one unified DataFrame for comparison
ncses_women_science_and_engineering_ed_vs_employment_df = pd.concat([degree_focus_df, occupational_group_df], ignore_index=True)

# Display the resulting DataFrame
print("Combined DataFrame for Degrees and Occupations:")
display(ncses_women_science_and_engineering_ed_vs_employment_df)


# #### Matplotlib Area Chart: Comparison of Women in S&E Degrees vs. S&E Occupations (1993, 2003, 2019)
# **Purpose:** This area chart provides a comparison of the growth in the number of women earning degrees in science and engineering (S&E) fields versus the number of women employed in S&E occupations. By overlaying these two trends, this visualization highlights potential disparities between education and workforce participation in Science and Engineering for women.
# 
# **Insights:**
# - **Significant Growth in S&E Degree Attainment:** From 1993 to 2019, the number of women earning degrees in S&E fields increased by approximately 202.67%, reflecting substantial progress in encouraging women to pursue studies in these fields. This trend suggests that more women are being equipped with the qualifications necessary for careers in science and engineering.
# 
# - **Lagging Workforce Representation Despite Educational Gains:** Over the same period, the increase in the number of women in S&E occupations was approximately 190.4%, a slower growth rate than the increase in degree attainment, suggesting that barriers may still prevent women from entering or advancing in S&E careers.
# 
# - **Indications of Structural Barriers:** The gap between educational achievement and workforce representation suggests that women in S&E fields may face structural or cultural barriers limiting their entry and advancement in these careers. These barriers likely impact not only individual career progression but also overall industry diversity and innovation potential.
# 
# **Summary:** This area chart emphasizes the challenge of translating educational advancements for women in S&E into equivalent representation in S&E occupations. Although women are earning degrees in S&E fields at increasing rates, this academic progress does not appear to be fully mirrored in workforce participation. This demonstrates that many women are encountering barriers to even entering into S&E fields, impacting their long-term career trajectories in the Science and Engineering Industry. These findings underscore the need for supportive, equitable workplace practices to bridge the gap between education and career progression for women in S&E.

# In[40]:


# Filter data for S&E category only and pivot it for plotting
degree_data = ncses_women_science_and_engineering_ed_vs_employment_df[
    (ncses_women_science_and_engineering_ed_vs_employment_df['Category'] == 'S&E') &
    (ncses_women_science_and_engineering_ed_vs_employment_df['Type'] == 'Degree')
]

occupation_data = ncses_women_science_and_engineering_ed_vs_employment_df[
    (ncses_women_science_and_engineering_ed_vs_employment_df['Category'] == 'S&E') &
    (ncses_women_science_and_engineering_ed_vs_employment_df['Type'] == 'Occupation')
]

# Extract counts and convert years to labels for plotting
years = ['1993', '2003', '2019']
degree_counts = degree_data[['1993 - Count', '2003 - Count', '2019 - Count']].values.flatten()
occupation_counts = occupation_data[['1993 - Count', '2003 - Count', '2019 - Count']].values.flatten()

# Plotting the line chart with area fill
plt.figure(figsize=(12, 6))
plt.plot(years, degree_counts, color='lightseagreen', marker='o', label='S&E Degrees')
plt.fill_between(years, degree_counts, color='lightseagreen', alpha=0.3)

plt.plot(years, occupation_counts, color='slategray', marker='o', label='S&E Occupations')
plt.fill_between(years, occupation_counts, color='slategray', alpha=0.5)

# Adding data point labels
for i, txt in enumerate(degree_counts):
    plt.text(years[i], degree_counts[i], f'{int(txt):,}', ha='center', va='bottom', color='black')

for i, txt in enumerate(occupation_counts):
    plt.text(years[i], occupation_counts[i], f'{int(txt):,}', ha='center', va='bottom', color='black')

# Labeling
plt.title("Comparison of S&E Degrees vs. S&E Occupations for Women (1993-2019)")
plt.xlabel("Year")
plt.ylabel("Count of Women")

# Adding commas to y-axis labels for readability
plt.gca().yaxis.set_major_formatter(mtick.FuncFormatter(lambda x, _: f'{int(x):,}'))


plt.legend(loc="upper left", title="Legend")
plt.show()


# In[41]:


# Define the path to the file: Figure LBR-21 - Women with a bachelor's degree or higher in S&E and S&E-related occupations: Selected years, 1993–2019
file_path = './data/ncses/nsb20212-figlbr-021.xlsx'

# Load the data, skipping rows based on the Excel file structure
start_row = 3  

# Load the data with descriptive column names
women_s_e_degree_trends_df = pd.read_excel(
    file_path,
    skiprows=start_row,
    names=[
        "Year", "Computer and Mathematical Scientists (%)", "Biological, Agricultural, and Environmental Life Scientists (%)",
        "Physical Scientists (%)", "Social Scientists (%)", "Engineers (%)", "All S&E-related Workers (%)"
    ]
)

# Display the resulting DataFrame to ensure it loaded correctly
print("S&E Degree Trends for Women DataFrame (with % values):")
display(women_s_e_degree_trends_df)


# In[42]:


# Define the path to the file: Figure LBR-27 - Median annual salaries of full-time workers with highest degrees in S&E or S&E-related fields, by sex: Selected years, 1995, 2003, and 2019
file_path = './data/ncses/nsb20212-figlbr-027.xlsx'

# Load the data, skipping rows based on the Excel file structure
start_row = 3  

# Load the data into a DataFrame
median_salary_by_gender_df = pd.read_excel(
    file_path,
    skiprows=start_row,
    names=["Degree Field", "Gender", "1995 Salary", "2003 Salary", "2019 Salary"]
)

# Forward-fill the "Degree Field" column to handle merged cells properly
median_salary_by_gender_df["Degree Field"] = median_salary_by_gender_df["Degree Field"].ffill()

# Display the final structured DataFrame
print("Restructured Salary DataFrame:")
display(median_salary_by_gender_df)


# #### Interactive Line Graph using Plotly Graph Objects: Median Salary by Gender with Percentage Increase Over Time
# 
# **Purpose:** To analyze and demonstrate the disparity in median salaries between genders over time for workers with S&E degrees, highlighting how salary increases have influenced the wage gap.
# 
# **Insights:**
# - **Initial Salary Disparity:**  
#     - In 1995, the median salary for males was $69,000, while females earned $47,000, creating an initial wage gap of 31.88%.
# - **Unequal Salary Increases:**  
#     - Over the next 8 years (first measured interval), the male median salary increased by 20.3%, whereas the female median salary increased by only 17.0%. This disproportionate increase further widened the gap between male and female earnings.
# - **Persistent Disparity by 2019:**  
#     - At the next interval, 16 years later, both genders saw an identical salary increase of 3.6%. However, the male median salary reached $86,000, while the female median salary only rose to $57,000. This resulted in an even larger wage gap, with a final disparity of 33.72%.
# 
# This interactive visualization highlights the long-standing and widening disparity in median salaries between genders. The unequal salary increases at crucial intervals have exacerbated the wage gap, emphasizing how systemic disparities in salary growth prevent women from closing the gap in career fields that require science and engineering degrees.

# In[43]:


# Filter only rows where Degree Field is "S&E"
se_salary_df = median_salary_by_gender_df[median_salary_by_gender_df["Degree Field"] == "S&E"].copy()

# Reshape the DataFrame from wide to long format for plotting
salary_melted_df = se_salary_df.melt(
    id_vars=["Gender"],
    value_vars=["1995 Salary", "2003 Salary", "2019 Salary"],
    var_name="Year",
    value_name="Median Salary"
)

# Convert 'Year' to numeric format
salary_melted_df["Year"] = salary_melted_df["Year"].str.extract('(\d{4})').astype(int)

# Sort values to ensure correct plotting order
salary_melted_df = salary_melted_df.sort_values(["Gender", "Year"])

# Calculate the percentage increase for each gender
salary_melted_df["Percentage Increase"] = salary_melted_df.groupby("Gender")["Median Salary"].pct_change() * 100

# Create the line chart with percentage increase annotations
fig = go.Figure()

# Add Male Salary Line
fig.add_trace(go.Scatter(
    x=salary_melted_df[salary_melted_df["Gender"] == "Male"]["Year"],
    y=salary_melted_df[salary_melted_df["Gender"] == "Male"]["Median Salary"],
    mode="lines+markers+text",
    name="Male Salary",
    line=dict(color="slategray"),
    marker=dict(size=8),
    text=[f"{perc:.1f}% Increase" if not pd.isna(perc) else ""
          for perc in salary_melted_df[salary_melted_df["Gender"] == "Male"]["Percentage Increase"]],
    textposition="top left",
    hovertemplate="<b>Year:</b> %{x}<br><b>Salary:</b> %{y:$,}<br>%{text}<extra></extra>"
))

# Add Female Salary Line
fig.add_trace(go.Scatter(
    x=salary_melted_df[salary_melted_df["Gender"] == "Female"]["Year"],
    y=salary_melted_df[salary_melted_df["Gender"] == "Female"]["Median Salary"],
    mode="lines+markers+text",
    name="Female Salary",
    line=dict(color="lightseagreen"),
    marker=dict(size=8),
    text=[f"{perc:.1f}% Increase" if not pd.isna(perc) else ""
          for perc in salary_melted_df[salary_melted_df["Gender"] == "Female"]["Percentage Increase"]],
    textposition="top left",
    hovertemplate="<b>Year:</b> %{x}<br><b>Salary:</b> %{y:$,}<br>%{text}<extra></extra>"
))

# Customize layout
fig.update_layout(
    title="Median Salary by Gender with Percentage Increase Over Time (S&E Degrees)",
    xaxis_title="Year",
    yaxis_title="Median Salary ($)",
    yaxis_tickprefix="$",
    yaxis_range=[40000, 105000],  # Adjusted y-axis range to give more room for top labels
    legend=dict(x=0.1, y=1.1, orientation="h"),
)

# Show the figure
fig.show()


# In[44]:


# Define the path to the file: Table SLBR-30 - Number and median salary of full-time workers with highest degree in S&E field, by sex and occupation: 2019
file_path = './data/ncses/nsb20212-tabslbr-030.xlsx'

# Define specific row indices for the occupations 
selected_rows = [7, 28, 40, 53, 66, 85, 112] 

# Load the entire file first, then filter for the selected rows
data = pd.read_excel(file_path, header=None)

# Select the specified rows and reset the index for male and female data
slbr30_male_df = data.iloc[selected_rows, [0, 5, 6]].copy()
slbr30_female_df = data.iloc[selected_rows, [0, 3, 4]].copy()

# Rename columns for both DataFrames
slbr30_male_df.columns = ["Occupation", "Total Workers", "Median Salary"]
slbr30_female_df.columns = ["Occupation", "Total Workers", "Median Salary"]

# Convert Thousands to actual counts
slbr30_male_df["Total Workers"] = slbr30_male_df["Total Workers"] * 1000
slbr30_female_df["Total Workers"] = slbr30_female_df["Total Workers"] * 1000

# Add Gender columns
slbr30_male_df["Gender"] = "Male"
slbr30_female_df["Gender"] = "Female"

# Combine the DataFrames
employment_count_and_salary_by_occupation_and_gender_df = pd.concat([slbr30_male_df, slbr30_female_df], ignore_index=True)

# Display the resulting DataFrame
print("Combined DataFrame for Selected Occupations and Salaries:")
display(employment_count_and_salary_by_occupation_and_gender_df)


# In[45]:


# Define the path to the file: Table SLBR-32 - Employed S&E highest degree holders, by sex, race or ethnicity, field of highest degree, and broad occupational category: 2019
file_path = './data/ncses/nsb20212-tabslbr-032.xlsx'

# Define start rows and number of rows for each gender based on the Excel file structure
start_row_female = 7  
start_row_male = 14   
num_rows_female = 5   
num_rows_male = 5     

# Load the Female data
female_df = pd.read_excel(
    file_path,
    skiprows=start_row_female,
    nrows=num_rows_female,
    names=[
        "Degree Type", "Total S&E Occupations (%)", 
        "S&E Occupations - Degree Related (%)", "S&E Occupations - Not Related to Degree (%)",
        "S&E-Related Occupations (%)", "Non-S&E Occupations (%)"
    ]
)

# Add Gender column for Female
female_df['Gender'] = 'Female'

# Load the Male data
male_df = pd.read_excel(
    file_path,
    skiprows=start_row_male,
    nrows=num_rows_male,
    names=[
        "Degree Type", "Total S&E Occupations (%)", 
        "S&E Occupations - Degree Related (%)", "S&E Occupations - Not Related to Degree (%)",
        "S&E-Related Occupations (%)", "Non-S&E Occupations (%)"
    ]
)

# Add Gender column for Male
male_df['Gender'] = 'Male'

# Combine the two DataFrames
se_degree_vs_occupation_by_gender_df = pd.concat([female_df, male_df], ignore_index=True)

# Display the final structured DataFrame
print("Occupation of S&E Degree Holders DataFrame with Gender:")
display(se_degree_vs_occupation_by_gender_df)


# #### Import Dataset 4: United States Census Bureau
# - From College to Jobs: American Community Survey 2019

# In[46]:


# Define the directory to store the downloaded files
download_dir = './data/from-college-to-jobs-acs-2019'
os.makedirs(download_dir, exist_ok=True)

# Dictionary mapping descriptive names to URLs
urls = {
    "job-by-bach-degree-field-all-ed-levels.xlsx": "https://www2.census.gov/programs-surveys/demo/tables/industry-occupation/2019/table1.xlsx",
    "job-by-bach-degree-field-bach-degree.xlsx": "https://www2.census.gov/programs-surveys/demo/tables/industry-occupation/2019/table2.xlsx",
    "job-by-bach-degree-field-grad-degree.xlsx": "https://www2.census.gov/programs-surveys/demo/tables/industry-occupation/2019/table3.xlsx",
    "med-earn-by-degree-level-field-and-occupation.xlsx": "https://www2.census.gov/programs-surveys/demo/tables/industry-occupation/2019/table4.xlsx"
}

# Download each file with the specified name
for file_name, url in urls.items():
    file_path = os.path.join(download_dir, file_name)
    response = requests.get(url)
    
    # Save the file
    with open(file_path, 'wb') as file:
        file.write(response.content)
    
    # Print confirmation message
    print(f"Downloaded {file_name} to {file_path}")


# #### Exploratory Data Analyis on Census Tables
# - The excel spreadsheet includes explanations of the data that spans several rows at the top and bottom of the data
#     - This means that rows need to be specifically extracted or excluded 
# - The census tables have headings and subheadings in the columns
#     - This means the columns need to be renamed to reflect the discriptive header rather than the repetitive subheading
# - The table contents are divided by row with a heading indicating the categorization in the rows following the row's subheading (e.g. male, female, white, hispanic)
#     - This means specific rows need to be extracted to ensure the data is only what is needed for analysis

# **Extract the data for the men from the first Excel file to test processing**

# In[47]:


# Define the path to the file
file_path = './data/from-college-to-jobs-acs-2019/job-by-bach-degree-field-all-ed-levels.xlsx'

# Specify only the required columns by their indices
selected_columns = [0, 1, 3, 5, 7, 9, 11, 13, 15, 17, 19, 21, 23, 25, 27, 29]

# Define the starting row and the number of rows to read
start_row = 27
num_rows = 20  # Number of rows to read after the starting row

# Load the data with only the selected columns and limit the rows
df_men_all_ed_levels = pd.read_excel(file_path, skiprows=start_row, nrows=num_rows, usecols=selected_columns)

# Define descriptive column names for the selected columns
columns = [
    "Occupation", "FoD-Computers_Math_Stats", "FoD-Engineering", "FoD-Physical_Sciences",
    "FoD-Biological_Environmental_Agricultural_Sciences", "FoD-Psychology",
    "FoD-Social_Sciences", "FoD-Multidiscipline", "FoD-Science_and_Engineering_Related",
    "FoD-Business", "FoD-Education", "FoD-Literature_and_Languages",
    "FoD-Liberal_Arts_and_History", "FoD-Visual_and_Performing_Arts",
    "FoD-Communications", "FoD-Other_EG_Criminal_Justice_or_Social_Work"
]

# Apply the custom column names
df_men_all_ed_levels.columns = columns

# Remove any leading dots or whitespace from the 'Occupation' column
df_men_all_ed_levels['Occupation'] = df_men_all_ed_levels['Occupation'].str.lstrip(". ")

# Add a gender column
df_men_all_ed_levels['Gender'] = 'Male'

# Preview to check the data
print(df_men_all_ed_levels.head())
print(df_men_all_ed_levels.tail())


# #### Exploratory Data Analyis (EDA): Check the Data Types

# In[48]:


df_men_all_ed_levels.info()


# **Convert columns to correct data types (float64 to int64)**

# In[49]:


# Select numeric columns that need conversion to int64
numeric_columns = [
    "FoD-Computers_Math_Stats", "FoD-Engineering", "FoD-Physical_Sciences", 
    "FoD-Biological_Environmental_Agricultural_Sciences", "FoD-Psychology",
    "FoD-Social_Sciences", "FoD-Multidiscipline", "FoD-Science_and_Engineering_Related",
    "FoD-Business", "FoD-Education", "FoD-Liberal_Arts_and_History", 
    "FoD-Visual_and_Performing_Arts", "FoD-Communications", 
    "FoD-Other_EG_Criminal_Justice_or_Social_Work"
]

# Convert selected numeric columns to int64
df_men_all_ed_levels[numeric_columns] = df_men_all_ed_levels[numeric_columns].astype('int64')

# Display updated data types to confirm changes
print("\nData types after conversion:")
print(df_men_all_ed_levels.info())


# **Repeat the process for the women's data in the same file**

# In[50]:


# Define the path to the file
file_path = './data/from-college-to-jobs-acs-2019/job-by-bach-degree-field-all-ed-levels.xlsx'

# Specify the selected columns by their indices (same as for men)
selected_columns = [0, 1, 3, 5, 7, 9, 11, 13, 15, 17, 19, 21, 23, 25, 27, 29]

# Define the header row and the number of rows for women
header_row = 49
num_rows = 20  # The same as for men

# Load the data for women with the specified columns, skipping initial rows and using the right number of rows
df_women_all_ed_levels = pd.read_excel(
    file_path, skiprows=header_row, nrows=num_rows, usecols=selected_columns
)

# Define descriptive column names
columns = [
    "Occupation", "FoD-Computers_Math_Stats", "FoD-Engineering", "FoD-Physical_Sciences",
    "FoD-Biological_Environmental_Agricultural_Sciences", "FoD-Psychology",
    "FoD-Social_Sciences", "FoD-Multidiscipline", "FoD-Science_and_Engineering_Related",
    "FoD-Business", "FoD-Education", "FoD-Literature_and_Languages",
    "FoD-Liberal_Arts_and_History", "FoD-Visual_and_Performing_Arts",
    "FoD-Communications", "FoD-Other_EG_Criminal_Justice_or_Social_Work"
]

# Apply the column names
df_women_all_ed_levels.columns = columns

# Remove any leading dots or whitespace from the 'Occupation' column
df_women_all_ed_levels['Occupation'] = df_women_all_ed_levels['Occupation'].str.lstrip(". ")

# Add a gender column with "Female" as the value
df_women_all_ed_levels['Gender'] = 'Female'

# Convert numeric columns to int64 to ensure consistency with the men’s DataFrame
numeric_columns = [
    "FoD-Computers_Math_Stats", "FoD-Engineering", "FoD-Physical_Sciences",
    "FoD-Biological_Environmental_Agricultural_Sciences", "FoD-Psychology",
    "FoD-Social_Sciences", "FoD-Multidiscipline", "FoD-Science_and_Engineering_Related",
    "FoD-Business", "FoD-Education", "FoD-Literature_and_Languages",
    "FoD-Liberal_Arts_and_History", "FoD-Visual_and_Performing_Arts",
    "FoD-Communications", "FoD-Other_EG_Criminal_Justice_or_Social_Work"
]
df_women_all_ed_levels[numeric_columns] = df_women_all_ed_levels[numeric_columns].astype('int64')

# Preview to confirm correct data import and types
print(df_women_all_ed_levels.head())
print(df_women_all_ed_levels.tail())
print(df_women_all_ed_levels.info())


# #### Matplotlib Stacked Bar Chart Analysis: Field of Degree by Occupation and Gender
# **Purpose:** Examine the distribution of fields of degree across occupations for men and women, highlighting patterns in career concentration and the alignment between degree backgrounds and career paths.
# 
# **Insights:**
# - **Gendered Career Concentrations:**
#   - *Men* show the highest representation in *Managerial (Non-STEM)* roles, suggesting a strong presence in leadership positions outside traditional science and engineering fields.
#   - *Women* are most concentrated in *Education*, with substantial representation in *Healthcare* as well. This indicates these areas as primary career paths for women, aligning with broader trends in occupational distribution by gender.
# 
# - **STEM Occupation Disparities:** 
#   - *Men* are much more likely to be employed as *Computer Workers* and *Engineers*, even when they hold degrees seemingly unrelated to these fields, such as *Liberal Arts and History*. This pattern may reflect broader hiring practices favoring men in technical roles, regardless of their degree background.
#   - *Women* show a much lower representation in technical fields like *Computer Work* and *Engineering*. This may indicate potential barriers to entry or differing career choices despite educational background.
# 
# - **Healthcare Sector Dominance by Women:** Women overwhelmingly dominate roles in *Healthcare*, aligning with the degree distributions in fields like *Biological, Environmental, and Agricultural Sciences*. This reflects an ongoing trend where women with science-related degrees often pursue healthcare-related careers.
# 
# - **Cross-Disciplinary Employment Trends:** While men more commonly cross into technical roles with non-STEM degrees, women tend to remain in roles closely related to their field of study, particularly in Education and Social Services.
# 

# In[51]:


# Define fields of degree columns
fields = df_men_all_ed_levels.columns[1:-1]  # Excluding 'Occupation' and 'Gender'

# Set up the figure with two subplots
fig, (ax_men, ax_women) = plt.subplots(1, 2, figsize=(18, 10), sharey=True)

# Define a custom color palette
custom_palette = sns.color_palette("deep", len(fields))

# Remove scientific notation by setting a scalar formatter
for ax in [ax_men, ax_women]:
    ax.get_yaxis().get_major_formatter().set_scientific(False)

# Plot for Men
bottom = None
for idx, field in enumerate(fields):
    ax_men.bar(
        df_men_all_ed_levels['Occupation'], df_men_all_ed_levels[field],
        label=field, bottom=bottom, color=custom_palette[idx]
    )
    bottom = df_men_all_ed_levels[fields[:list(fields).index(field) + 1]].sum(axis=1)
ax_men.set_title("Field of Degree Distribution by Occupation (Men)")
ax_men.set_xlabel("")
ax_men.set_ylabel("Count")
ax_men.legend(title="Field of Degree", bbox_to_anchor=(1.05, 1), loc='upper left')
ax_men.set_xticks(range(len(df_men_all_ed_levels['Occupation'])))
ax_men.set_xticklabels(df_men_all_ed_levels['Occupation'], rotation=45, ha='right')

# Plot for Women
bottom = None
for idx, field in enumerate(fields):
    ax_women.bar(
        df_women_all_ed_levels['Occupation'], df_women_all_ed_levels[field],
        label=field, bottom=bottom, color=custom_palette[idx]
    )
    bottom = df_women_all_ed_levels[fields[:list(fields).index(field) + 1]].sum(axis=1)
ax_women.set_title("Field of Degree Distribution by Occupation (Women)")
ax_women.set_xlabel("")
ax_women.set_xticks(range(len(df_women_all_ed_levels['Occupation'])))
ax_women.set_xticklabels(df_women_all_ed_levels['Occupation'], rotation=45, ha='right')

# Adjust layout and show plot
plt.tight_layout()

# Removing edge color from bars to eliminate white lines
for ax in [ax_men, ax_women]:
    for bar in ax.containers:
        for patch in bar:
            patch.set_edgecolor('none')

plt.show()


# #### Process the second xlsx file from the American Community Survey
# - Recreate the steps used on the first file from the dataset

# In[52]:


# Define the path to the file
file_path = './data/from-college-to-jobs-acs-2019/job-by-bach-degree-field-bach-degree.xlsx'

# Specify only the required columns by their indices
selected_columns = [0, 1, 3, 5, 7, 9, 11, 13, 15, 17, 19, 21, 23, 25, 27, 29]

# Define the starting row and the number of rows to read
start_row = 27
num_rows = 20  # Number of rows to read after the starting row

# Load the data with only the selected columns and limit the rows
df_men_bach_degree = pd.read_excel(file_path, skiprows=start_row, nrows=num_rows, usecols=selected_columns)

# Define descriptive column names for the selected columns
columns = [
    "Occupation", "FoD-Computers_Math_Stats", "FoD-Engineering", "FoD-Physical_Sciences",
    "FoD-Biological_Environmental_Agricultural_Sciences", "FoD-Psychology",
    "FoD-Social_Sciences", "FoD-Multidiscipline", "FoD-Science_and_Engineering_Related",
    "FoD-Business", "FoD-Education", "FoD-Literature_and_Languages",
    "FoD-Liberal_Arts_and_History", "FoD-Visual_and_Performing_Arts",
    "FoD-Communications", "FoD-Other_EG_Criminal_Justice_or_Social_Work"
]

# Apply the custom column names
df_men_bach_degree.columns = columns

# Remove any leading dots or whitespace from the 'Occupation' column
df_men_bach_degree['Occupation'] = df_men_bach_degree['Occupation'].str.lstrip(". ")

# Add a gender column
df_men_bach_degree['Gender'] = 'Male'

# Select numeric columns that need conversion to int64
numeric_columns = [
    "FoD-Computers_Math_Stats", "FoD-Engineering", "FoD-Physical_Sciences",
    "FoD-Biological_Environmental_Agricultural_Sciences", "FoD-Psychology",
    "FoD-Social_Sciences", "FoD-Multidiscipline", "FoD-Science_and_Engineering_Related",
    "FoD-Business", "FoD-Education", "FoD-Literature_and_Languages",
    "FoD-Liberal_Arts_and_History", "FoD-Visual_and_Performing_Arts",
    "FoD-Communications", "FoD-Other_EG_Criminal_Justice_or_Social_Work"
]

# Convert selected numeric columns to int64
df_men_bach_degree[numeric_columns] = df_men_bach_degree[numeric_columns].astype('int64')

# Preview to confirm correct data import and types
print(df_men_bach_degree.head())
print(df_men_bach_degree.tail())
print(df_men_bach_degree.info())


# In[53]:


# Define the path to the file
file_path = './data/from-college-to-jobs-acs-2019/job-by-bach-degree-field-bach-degree.xlsx'

# Specify the selected columns by their indices (same as for men)
selected_columns = [0, 1, 3, 5, 7, 9, 11, 13, 15, 17, 19, 21, 23, 25, 27, 29]

# Define the header row and the number of rows for women
header_row = 49
num_rows = 20  # The same as for men

# Load the data for women with the specified columns, skipping initial rows and using the right number of rows
df_women_bach_degree = pd.read_excel(
    file_path, skiprows=header_row, nrows=num_rows, usecols=selected_columns
)

# Define descriptive column names
columns = [
    "Occupation", "FoD-Computers_Math_Stats", "FoD-Engineering", "FoD-Physical_Sciences",
    "FoD-Biological_Environmental_Agricultural_Sciences", "FoD-Psychology",
    "FoD-Social_Sciences", "FoD-Multidiscipline", "FoD-Science_and_Engineering_Related",
    "FoD-Business", "FoD-Education", "FoD-Literature_and_Languages",
    "FoD-Liberal_Arts_and_History", "FoD-Visual_and_Performing_Arts",
    "FoD-Communications", "FoD-Other_EG_Criminal_Justice_or_Social_Work"
]

# Apply the column names
df_women_bach_degree.columns = columns

# Remove any leading dots or whitespace from the 'Occupation' column
df_women_bach_degree['Occupation'] = df_women_bach_degree['Occupation'].str.lstrip(". ")

# Add a gender column with "Female" as the value
df_women_bach_degree['Gender'] = 'Female'

# Convert numeric columns to int64 to ensure consistency with the men’s DataFrame
numeric_columns = [
    "FoD-Computers_Math_Stats", "FoD-Engineering", "FoD-Physical_Sciences",
    "FoD-Biological_Environmental_Agricultural_Sciences", "FoD-Psychology",
    "FoD-Social_Sciences", "FoD-Multidiscipline", "FoD-Science_and_Engineering_Related",
    "FoD-Business", "FoD-Education", "FoD-Literature_and_Languages",
    "FoD-Liberal_Arts_and_History", "FoD-Visual_and_Performing_Arts",
    "FoD-Communications", "FoD-Other_EG_Criminal_Justice_or_Social_Work"
]
df_women_bach_degree[numeric_columns] = df_women_bach_degree[numeric_columns].astype('int64')

# Preview to confirm correct data import and types
print(df_women_bach_degree.head())
print(df_women_bach_degree.tail())
print(df_women_bach_degree.info())


# #### Process the third xlsx file from the American Community Survey
# - Recreate the steps used on the first and second files from the dataset

# In[54]:


# Define the path to the file
file_path = './data/from-college-to-jobs-acs-2019/job-by-bach-degree-field-grad-degree.xlsx'

# Specify only the required columns by their indices
selected_columns = [0, 1, 3, 5, 7, 9, 11, 13, 15, 17, 19, 21, 23, 25, 27, 29]

# Define the starting row and the number of rows to read
start_row = 27
num_rows = 20  # Number of rows to read after the starting row

# Load the data with only the selected columns and limit the rows
df_men_grad_degree = pd.read_excel(file_path, skiprows=start_row, nrows=num_rows, usecols=selected_columns)

# Define descriptive column names for the selected columns
columns = [
    "Occupation", "FoD-Computers_Math_Stats", "FoD-Engineering", "FoD-Physical_Sciences",
    "FoD-Biological_Environmental_Agricultural_Sciences", "FoD-Psychology",
    "FoD-Social_Sciences", "FoD-Multidiscipline", "FoD-Science_and_Engineering_Related",
    "FoD-Business", "FoD-Education", "FoD-Literature_and_Languages",
    "FoD-Liberal_Arts_and_History", "FoD-Visual_and_Performing_Arts",
    "FoD-Communications", "FoD-Other_EG_Criminal_Justice_or_Social_Work"
]

# Apply the custom column names
df_men_grad_degree.columns = columns

# Remove any leading dots or whitespace from the 'Occupation' column
df_men_grad_degree['Occupation'] = df_men_grad_degree['Occupation'].str.lstrip(". ")

# Add a gender column
df_men_grad_degree['Gender'] = 'Male'

# Select numeric columns that need conversion to int64
numeric_columns = [
    "FoD-Computers_Math_Stats", "FoD-Engineering", "FoD-Physical_Sciences",
    "FoD-Biological_Environmental_Agricultural_Sciences", "FoD-Psychology",
    "FoD-Social_Sciences", "FoD-Multidiscipline", "FoD-Science_and_Engineering_Related",
    "FoD-Business", "FoD-Education", "FoD-Literature_and_Languages",
    "FoD-Liberal_Arts_and_History", "FoD-Visual_and_Performing_Arts",
    "FoD-Communications", "FoD-Other_EG_Criminal_Justice_or_Social_Work"
]

# Convert selected numeric columns to int64
df_men_grad_degree[numeric_columns] = df_men_grad_degree[numeric_columns].astype('int64')

# Preview to confirm correct data import and types
print(df_men_grad_degree.head())
print(df_men_grad_degree.tail())
print(df_men_grad_degree.info())


# In[55]:


# Define the path to the file
file_path = './data/from-college-to-jobs-acs-2019/job-by-bach-degree-field-grad-degree.xlsx'

# Specify the selected columns by their indices (same as for men)
selected_columns = [0, 1, 3, 5, 7, 9, 11, 13, 15, 17, 19, 21, 23, 25, 27, 29]

# Define the header row and the number of rows for women
header_row = 49
num_rows = 20  # The same as for men

# Load the data for women with the specified columns, skipping initial rows and using the right number of rows
df_women_grad_degree = pd.read_excel(
    file_path, skiprows=header_row, nrows=num_rows, usecols=selected_columns
)

# Define descriptive column names
columns = [
    "Occupation", "FoD-Computers_Math_Stats", "FoD-Engineering", "FoD-Physical_Sciences",
    "FoD-Biological_Environmental_Agricultural_Sciences", "FoD-Psychology",
    "FoD-Social_Sciences", "FoD-Multidiscipline", "FoD-Science_and_Engineering_Related",
    "FoD-Business", "FoD-Education", "FoD-Literature_and_Languages",
    "FoD-Liberal_Arts_and_History", "FoD-Visual_and_Performing_Arts",
    "FoD-Communications", "FoD-Other_EG_Criminal_Justice_or_Social_Work"
]

# Apply the column names
df_women_grad_degree.columns = columns

# Remove any leading dots or whitespace from the 'Occupation' column
df_women_grad_degree['Occupation'] = df_women_grad_degree['Occupation'].str.lstrip(". ")

# Add a gender column with "Female" as the value
df_women_grad_degree['Gender'] = 'Female'

# Convert numeric columns to int64 to ensure consistency with the men’s DataFrame
numeric_columns = [
    "FoD-Computers_Math_Stats", "FoD-Engineering", "FoD-Physical_Sciences",
    "FoD-Biological_Environmental_Agricultural_Sciences", "FoD-Psychology",
    "FoD-Social_Sciences", "FoD-Multidiscipline", "FoD-Science_and_Engineering_Related",
    "FoD-Business", "FoD-Education", "FoD-Literature_and_Languages",
    "FoD-Liberal_Arts_and_History", "FoD-Visual_and_Performing_Arts",
    "FoD-Communications", "FoD-Other_EG_Criminal_Justice_or_Social_Work"
]
df_women_grad_degree[numeric_columns] = df_women_grad_degree[numeric_columns].astype('int64')

# Preview to confirm correct data import and types
print(df_women_grad_degree.head())
print(df_women_grad_degree.tail())
print(df_women_grad_degree.info())


# #### Process the 4th xlsx file from the American Community Survey

# In[56]:


# Define the path to the file
file_path = './data/from-college-to-jobs-acs-2019/med-earn-by-degree-level-field-and-occupation.xlsx'

# Define the columns to read (based on the specified indices)
selected_columns = [0, 1, 3, 5, 7, 9, 11]

# Define column names for the data
columns = [
    "Occupation", "STEM Major All Degrees", "non-STEM Major All Degrees",
    "STEM Major Bachelors", "non-STEM Major Bachelors",
    "STEM Major Graduate Degree", "non-STEM Major Graduate Degree"
]

# Load the data for men
start_row_men = 16  
num_rows_men = 6  # Number of rows for the men's data section
male_median_earnings = pd.read_excel(file_path, skiprows=start_row_men, nrows=num_rows_men, usecols=selected_columns)
male_median_earnings.columns = columns

# Remove any leading dots or whitespace from the 'Occupation' column
male_median_earnings['Occupation'] = male_median_earnings['Occupation'].str.lstrip(". ")

male_median_earnings['Gender'] = 'Male'

# Load the data for women
start_row_women = start_row_men + num_rows_men + 3  
num_rows_women = 6  # Number of rows for the women's data section
female_median_earnings = pd.read_excel(file_path, skiprows=start_row_women, nrows=num_rows_women, usecols=selected_columns)
female_median_earnings.columns = columns

# Remove any leading dots or whitespace from the 'Occupation' column
female_median_earnings['Occupation'] = female_median_earnings['Occupation'].str.lstrip(". ")

female_median_earnings['Gender'] = 'Female'

# Preview the results to confirm
print("Men's Median Earnings DataFrame:")
display(male_median_earnings)
print("Women's Median Earnings DataFrame:")
display(female_median_earnings)


# In[57]:


# Remove STEM Major All Degrees and non-STEM Major All Degrees from both dataframes since it is unneeded for analysis
male_median_earnings = male_median_earnings.drop(columns=['STEM Major All Degrees', 'non-STEM Major All Degrees'])
female_median_earnings = female_median_earnings.drop(columns=['STEM Major All Degrees', 'non-STEM Major All Degrees'])

# Combine male and female earnings data into a single DataFrame for comparison
combined_median_earnings_by_degree_and_occupation = pd.concat([male_median_earnings, female_median_earnings], ignore_index=True)

# Display the resulting DataFrame to verify its structure 
print("Combined DataFrame for Median Earnings by Degree and Occupation")
display(combined_median_earnings_by_degree_and_occupation)


# #### Plotly Interactive Scatter Plot: Median Earnings by Occupation, Degree Type, and Gender
# **Purpose:** This interactive scatter plot visualizes median earnings by occupation, degree type, and gender within the tech industry. The objective is to demonstrate that, despite similar educational backgrounds and occupational fields, men frequently earn higher median salaries than women. By allowing users to hover over data points to view specific salary values, this plot highlights areas where men out-earn women with equivalent or even higher qualifications, underscoring the persistent gender disparities in earnings.
# 
# **Insights:**
# - **Engineers: Salary Disparity Despite Equal or Higher Qualifications for Women:** 
#     - Among *engineers*, *women with STEM graduate degrees* earn $91,300, whereas *men with STEM graduate degrees* earn a significantly higher $121,300, resulting in a wage gap of approximately 25% even with identical educational qualifications.
#     - This disparity is compounded by the fact that men with only a STEM bachelor’s degree earn $100,000—about 10% more than women with a higher-level STEM graduate degree. Additionally, men with non-STEM graduate degrees earn $106,400, further illustrating that women must attain advanced qualifications simply to approach salaries earned by men with lesser degrees.
# 
# - **Computer Workers: Higher Qualifications Don’t Close the Salary Gap:** 
#     - For *computer-related roles*, *women with STEM graduate degrees* earn $101,800, while *men with STEM bachelor’s degrees* earn nearly the same at $101,300—despite the men holding lower-level qualifications.
#     - *Men with non-STEM graduate degrees* make $107,100, outpacing *women with STEM graduate degrees* by over 5%. In a field where relevant STEM qualifications would seem essential, the data shows that women’s higher education still doesn’t yield comparable earnings to men’s lesser degrees.
# 
# - **Social Scientists: Persistent Gaps at Both Bachelor’s and Graduate Levels:** 
#     - In *social science roles*, *women with STEM bachelor’s degrees* earn $69,560, whereas men with equivalent STEM bachelor’s degrees earn $80,070, creating a wage gap of around 13% despite equal educational backgrounds.
#     - At the graduate level, the pattern persists. *Women with STEM graduate degrees* earn $80,070, while *men with non-STEM graduate degrees* earn $85,070. The top salary here goes to *men with STEM graduate degrees*, earning $100,200--25% more than their female counterparts with the same qualifications.
# 
# - **Physical Scientists: Women’s Earnings Outpaced by Men’s Lesser Degrees:** 
#     - In *physical science roles*, *women with STEM graduate degrees* earn $86,430, while *men with non-STEM graduate degrees* earn $89,050—a difference of about 3% in favor of men with less relevant qualifications.
#     - The gap becomes even more pronounced when comparing *women with STEM graduate degrees* ($86,430) to *men with STEM graduate degrees*, who earn $102,100. This equates to a wage disparity of approximately 18%, showing that even with equivalent qualifications, women’s earnings lag significantly behind their male counterparts.
# 
# **Summary:** This interactive scatter plot brings attention to the pervasive wage disparities between men and women in the tech industry. The visualization highlights that even when women attain higher qualifications—such as a *STEM graduate degree*—they are often out-earned by men with lower or unrelated qualifications. This pattern emerges across multiple fields, from *engineering* to *social science*, showing that women’s qualifications yield a lesser financial return than men’s. This data underscores a systemic issue within tech fields, pointing to potential structural barriers that prevent equitable compensation for women, even when they achieve the same or higher qualifications as their male counterparts.

# In[58]:


# Reshape the DataFrame to long format for easier plotting
df_long = combined_median_earnings_by_degree_and_occupation.melt(
    id_vars=["Occupation", "Gender"],
    value_vars=["STEM Major Bachelors", "non-STEM Major Bachelors", "STEM Major Graduate Degree", "non-STEM Major Graduate Degree"],
    var_name="Degree Type",
    value_name="Median Salary"
)

# Map degree types to more readable labels
df_long['Degree Type'] = df_long['Degree Type'].replace({
    "STEM Major Bachelors": "Bachelor's: STEM",
    "non-STEM Major Bachelors": "Bachelor's: Non-STEM",
    "STEM Major Graduate Degree": "Graduate Degree: STEM",
    "non-STEM Major Graduate Degree": "Graduate Degree: Non-STEM"
})

# Create an interactive scatter plot with Plotly
fig = px.scatter(
    df_long,
    x="Median Salary",
    y="Occupation",
    color="Gender",
    symbol="Degree Type",
    hover_data={
        "Median Salary": True,  # Show Median Salary on hover
        "Occupation": False,  # Hide Occupation as it's already on the y-axis
        "Gender": False,  # Hide Gender as it's shown in the color
        "Degree Type": False  # Hide Degree Type as it's shown in the symbol
    },
    color_discrete_map={"Male": "slategray", "Female": "lightseagreen"},
    title="Median Earnings by Occupation, Degree Type, and Gender"
)

# Update marker settings for point size and transparency
fig.update_traces(marker=dict(size=16, opacity=0.7))  # Increase size and set alpha to 0.7

# Customize the layout
fig.update_layout(
    xaxis_title="Median Salary",
    yaxis_title=None,
    legend_title_text="Gender & Degree Type",
    legend=dict(x=1, y=1, xanchor="left", yanchor="top")
)

# Show plot
fig.show()


# ## Machine Learning Plan for Checkpoint 3
# 
# #### What type of machine learning model are you planning to use? 📝
# 
# Linear regression is appropriate for this analysis because the target variable (salary) is continuous. The model will quantify the impact of gender, education, occupation, and years of experience on salaries while highlighting disparities.
# 
# If the relationships between variables appear non-linear, I may explore polynomial regression to better capture more complex patterns in the data.

# #### What are the challenges have you identified/are you anticipating in building your machine learning model? 📝
# #### How are you planning to address these challenges? 📝
#  
# **Challenges and Solutions:**  
# 1. **Ambiguous or Conflicting Data:**  
#    - Some columns, such as *DeveloperType*, *MobileDeveloperType*, and *NonDeveloperType*, contain multiple values per record, or entries that conflict with other responses (e.g., someone identifying as a *"Mobile Developer"* but having *NA* in *MobileDeveloperType*).  
#    - **Mitigation:** I will clean up these columns by separating multiple values into their own features and identifying any inconsistencies. These issues will be handled by either imputing the data with logical defaults or removing problematic records where necessary.
# 
# 2. **Excessive and Irrelevant Features:**  
#    - The dataset contains far more columns than necessary for the analysis, which could introduce noise and make the model more complex than required.  
#    - **Mitigation:** I will focus on features most relevant to salary prediction (e.g., gender, education, occupation type) and remove extraneous columns that are unrelated to the target variable.  
# 
# 3. **Filtering for Relevant Responses:**  
#    - The dataset includes responses that may not align with the analysis, such as students, unemployed individuals, or genders other than male and female. These records are not relevant to the analysis question.  
#    - **Mitigation:** I will filter out records that do not meet the analysis criteria, ensuring only relevant data is included.  
# 
# 4. **Imbalanced Dataset:**  
#    - The dataset may have significantly more male respondents than female respondents, which could bias predictions or skew results.  
#    - **Mitigation:** I will normalize the data to ensure gender groups are proportionally represented. If necessary, stratified sampling will be applied when splitting the dataset into training and testing sets. Regression evaluation metrics like mean squared error (MSE) and R-squared (R²) will be used to assess model performance, focusing on minimizing bias in predictions.
# 
# 5. **Outliers in Salary Data:**  
#    - Salary data may contain extreme values due to self-reporting, which could disproportionately influence the regression model. Outliers may result from errors 
#    (e.g., misplaced decimal points) or actual extremes in senior roles or low-paying positions.  
#    - **Mitigation:** I will identify outliers using statistical methods such as interquartile range (IQR) or z-scores. Depending on the findings, I will handle these outliers by:
#       - For genuine outliers: Retain them in both the training and testing sets but apply scaling or transformations (e.g., logarithmic scaling) to minimize their impact on the model.
#       - For data entry errors or invalid outliers: Correct obvious errors where feasible. If correction is not possible, remove these records from both the training and testing sets to avoid distortion.
#          
# 6. **Ordinal vs. Nominal Categorical Data:**  
#    - Certain categorical features, such as *Professional*, *DeveloperType*, *WebDeveloperType*, and *NonDeveloperType*, may not have a clear ordinal relationship. While some roles (e.g., *Full-Stack Developer* vs. *Front-End Developer*) may imply a difference in salary, others (e.g., *Application Developer* vs. *Full-Stack Developer*) may lack a consistent ranking.  
#    - **Mitigation:** I will analyze salary trends within these columns to identify any patterns or relationships. For roles without clear ordinal relationships, one-hot encoding will be applied to prevent imposing an artificial hierarchy. For roles with observable ordinal trends, custom mappings will be created based on the data. 

# ## Machine Learning Implementation Process  
# 

# #### Import Dataset: Stack Overflow Annual Developer Survey 2017

# In[59]:


# import zip file from Stack Overflow
file_handle, _ = urlretrieve("https://survey.stackoverflow.co/datasets/stack-overflow-developer-survey-2017.zip")
zipfile = ZipFile(file_handle, "r")
zipfile.extractall("./data/Stack_Overflow_2017")
zipfile.close()


# In[60]:


# Define the data directory
stack_data_dir = './data/Stack_Overflow_2017'

# Read each CSV file into its own DataFrame with meaningful names
stack_overflow_results_df = pd.read_csv(os.path.join(stack_data_dir, 'survey_results_public.csv'))
stack_overflow_schema_df = pd.read_csv(os.path.join(stack_data_dir, 'survey_results_schema.csv'))

# Display the first 5 records from each DataFrame
print("Stack Overflow Annual Developer Survey 2017 Results:")
display(stack_overflow_results_df.head(5))
display(stack_overflow_results_df.shape)
print("Stack Overflow Annual Developer Survey 2017 Schema:")
display(stack_overflow_schema_df.head(5))
display(stack_overflow_schema_df.shape)


# #### Exploratory Data Analysis (EDA): Perform initial dataset filtering for records and columns relevant to the scope of the project

# In[61]:


# Filter for the United States, respondents who chose male or female as the gender, and exclude anyone who is unemployed or a student
stack_overflow_results_df = stack_overflow_results_df[
    (stack_overflow_results_df['Country'] == 'United States') &
    (stack_overflow_results_df['Professional'] != 'Student') &
    (~stack_overflow_results_df['EmploymentStatus'].str.contains('Not employed', na=False)) &
    (stack_overflow_results_df['Gender'].isin(['Male', 'Female']))
]

display(stack_overflow_results_df.head())
display(stack_overflow_results_df.shape)


# In[62]:


# Keep only the specified columns
stack_overflow_results_df = stack_overflow_results_df[['Respondent', 'Gender', 'Professional', 'University', 'EmploymentStatus', 'FormalEducation',
                                                       'MajorUndergrad', 'HomeRemote', 'CompanySize', 'CompanyType', 'YearsCodedJob',
                                                       'YearsCodedJobPast', 'DeveloperType', 'WebDeveloperType', 'MobileDeveloperType', 'NonDeveloperType',
                                                       'Currency', 'Overpaid', 'Salary']]

# Check the dataframe structure
display(stack_overflow_results_df.head())
display(stack_overflow_results_df.shape)


# In[63]:


# Count unique values in each column
unique_counts = stack_overflow_results_df.nunique().sort_values(ascending=False)
print("Unique Values per Column:")
print(unique_counts)


# #### Exploratory Data Analysis: determine currencies and handle any that are not US Dollars
# - Non-US Currencies:
#     - Only four respondents who reported working in the US selected non-US currencies.
#     - These records were removed from the dataset due to their insignificance to the overall analysis.
# 
# - Missing Currency Values:
#     - 2916 respondents did not specify a currency type.
#     - Since these are American respondents, it was assumed their salary amounts are in US dollars.
#     - Planned outlier detection and handling will address any inconsistencies or errors in these assumptions.

# In[64]:


# Count occurrences of each currency type (including NaN values)
currency_counts = stack_overflow_results_df['Currency'].value_counts(dropna=False)

# Display the results
print("Counts of each currency type:")
print(currency_counts)


# In[65]:


# Replace NaN values in the Currency column with 'U.S. dollars ($)'
stack_overflow_results_df['Currency'] = stack_overflow_results_df['Currency'].fillna('U.S. dollars ($)')

# Filter the dataset to keep only rows with 'U.S. dollars ($)'
stack_overflow_results_df = stack_overflow_results_df[stack_overflow_results_df['Currency'] == 'U.S. dollars ($)']

# Verify the results
currency_counts = stack_overflow_results_df['Currency'].value_counts()
print("Updated counts of each currency type:")
print(currency_counts)


# #### Exploratory Data Analysis: find what values are in EmploymentStatus column and determine how to handle them
# - Objective: Ensure consistency with the scope of the project and alignment with other datasets used in the analysis.
# - Action Taken: All values in the EmploymentStatus column other than "Employed full-time" were removed.

# In[66]:


# Count occurrences of each EmploymentStatus (including NaN values)
employment_status_counts = stack_overflow_results_df['EmploymentStatus'].value_counts(dropna=False)

# Display the results
print("Counts of each employment status value:")
print(employment_status_counts)


# In[67]:


# Filter the dataset to keep only rows with 'Employed full-time' to be consistent with other datasets in the project
stack_overflow_results_df = stack_overflow_results_df[stack_overflow_results_df['EmploymentStatus'] == 'Employed full-time']

# Verify the results
employment_status_counts = stack_overflow_results_df['EmploymentStatus'].value_counts()
print("Updated counts of each employment status:")
print(employment_status_counts)


# #### Exploratory Data Analysis: find what values are in Professional column and determine how to handle them
# - Objective: Refine the dataset by excluding irrelevant or inconsistent records in the Professional column.
# 
# - Insights:
#     - The Professional column contains four distinct values:
#         - Professional developer (5699 records)
#         - Professional non-developer who sometimes writes code (678 records)
#         - Used to be a professional developer (106 records)
#         - None of these (33 records)
# 
# - Action Taken:
#     - Rows where Professional is "Used to be a professional developer" or "None of these" were removed only if all related developer/employment type fields (DeveloperType, WebDeveloperType, MobileDeveloperType, NonDeveloperType) were also NaN.
#     - This ensured that only meaningful and relevant professional data remained in the dataset.

# In[68]:


# Count occurrences of each value in Professional (including NaN values)
professional_counts = stack_overflow_results_df['Professional'].value_counts(dropna=False)

# Display the results
print("Counts of each 'Professional' value:")
print(professional_counts)


# In[69]:


# Filter the dataset to remove rows where Professional is 'None of these' or 'Used to be a professional developer' 
# and all fields that indicate developer/employment type are NA

# Define the filter condition
condition = (
    (stack_overflow_results_df['Professional'].isin(["Used to be a professional developer", "None of these"])) &
    (stack_overflow_results_df[['DeveloperType', 'WebDeveloperType', 'MobileDeveloperType', 'NonDeveloperType']].isna().all(axis=1))
)

# Apply the filter to remove rows matching the condition
stack_overflow_results_df = stack_overflow_results_df[~condition]

# Verify the filtering
print("Remaining records after filtering:")
print(stack_overflow_results_df['Professional'].value_counts())


# #### Exploratory Data Analysis: examine values/missing values in remaining fields

# In[70]:


missing_salary_count = stack_overflow_results_df['Salary'].isna().sum()
print(f"Number of missing Salary values: {missing_salary_count}")


# In[71]:


# List of columns to analyze
columns_to_check = [
    'DeveloperType', 'WebDeveloperType', 'MobileDeveloperType', 
    'NonDeveloperType', 'YearsCodedJob', 'YearsCodedJobPast',
    'MajorUndergrad', 'CompanyType', 'CompanySize',
    'FormalEducation', 'HomeRemote', 'University'
]

# Iterate through each column and display unique values and their counts
for column in columns_to_check:
    print(f"Unique values in {column} and their counts:")
    print(stack_overflow_results_df[column].value_counts(dropna=False))
    print("\n")


# #### Exploratory Data Analysis: Mapping and Analyzing CompanySize
# - Objective: Determine whether the size of a respondent's company correlates with their reported salary.
# 
# - Steps Taken: 
#     - Mapped CompanySize values into numerical equivalents based on the approximate median employee count in each category. For example:
#         - "Fewer than 10 employees" mapped to 5.
#         - "10 to 19 employees" mapped to 15.
#         - "10,000 or more employees" mapped to 15000.
#     - Excluded records where CompanySize or Salary was NaN
#     - Calculated the correlation between CompanySizeNumeric and Salary.
# 
# - Results:
#     - The positive correlation between CompanySizeNumeric and Salary is weak (correlation coefficient = 0.158), suggesting that larger companies may pay slightly higher salaries.

# In[72]:


# Create a numerical mapping for CompanySize
company_size_mapping = {
    "Fewer than 10 employees": 5,
    "10 to 19 employees": 15,
    "20 to 99 employees": 50,
    "100 to 499 employees": 300,
    "500 to 999 employees": 750,
    "1,000 to 4,999 employees": 3000,
    "5,000 to 9,999 employees": 7500,
    "10,000 or more employees": 15000,
    "I don't know": None,
    "I prefer not to answer": None,
    None: None
}

# Map the values to numerical equivalents
stack_overflow_results_df['CompanySizeNumeric'] = stack_overflow_results_df['CompanySize'].map(company_size_mapping)

# Filter the dataset for rows where both Salary and CompanySizeNumeric are not NaN
valid_data = stack_overflow_results_df[
    stack_overflow_results_df['Salary'].notna() & stack_overflow_results_df['CompanySizeNumeric'].notna()
]

# Calculate the correlation matrix on the filtered data
correlation_matrix = valid_data[['CompanySizeNumeric', 'Salary']].corr()

# Print the correlation matrix
print("Correlation matrix (only where Salary and CompanySizeNumeric are not NaN):")
print(correlation_matrix)

# Preview the filtered data for review
print(valid_data[['CompanySize', 'CompanySizeNumeric', 'Salary']].head())


# #### Exploratory Data Analysis: Mapping Undergraduate Majors to American Community Survey (ACS) Categories
# - Objective: Align the MajorUndergrad field with the American Community Survey (ACS) categories for compatibility with existing datasets and for the purpose of STEM vs. non-STEM analysis.
# 
# - Steps Taken:
#     - Mapped MajorUndergrad values into ACS categories:
#         - For example: "Computer science or software engineering" = FoD-Computers_Math_Stats.
#         - Uncommon or ambiguous responses, such as "Something else," were mapped to broader categories (e.g., FoD-Other_EG_Criminal_Justice_or_Social_Work).
#     - Created a new column, ACS_Major, reflecting these mappings.
#     - Defined STEM fields based on ACS categories, including: FoD-Computers_Math_Stats, FoD-Engineering, FoD-Physical_Sciences, etc.
#     - Added a binary column, STEM Degree, indicating whether the major falls under STEM (Yes/No).

# In[73]:


# Map MajorUndergrad to match American Community Survey degree categories
major_to_acs_mapping = {
    "Computer science or software engineering": "FoD-Computers_Math_Stats",
    "Computer engineering or electrical/electronics engineering": "FoD-Engineering",
    "A natural science": "FoD-Physical_Sciences",
    "Computer programming or Web development": "FoD-Computers_Math_Stats",
    "Mathematics or statistics": "FoD-Computers_Math_Stats",
    "A non-computer-focused engineering discipline": "FoD-Engineering",
    "A humanities discipline": "FoD-Literature_and_Languages",
    "Information technology, networking, or system administration": "FoD-Computers_Math_Stats",
    "Fine arts or performing arts": "FoD-Visual_and_Performing_Arts",
    "Something else": "FoD-Other_EG_Criminal_Justice_or_Social_Work",
    "Management information systems": "FoD-Computers_Math_Stats",
    "A business discipline": "FoD-Business",
    "A social science": "FoD-Social_Sciences",
    "I never declared a major": "FoD-Multidiscipline",
    "Psychology": "FoD-Psychology",
    "A health science": "FoD-Biological_Environmental_Agricultural_Sciences"
}

# Create ACS_Major column
stack_overflow_results_df['ACS_Major'] = stack_overflow_results_df['MajorUndergrad'].map(major_to_acs_mapping)

# Define STEM vs. non-STEM categories
stem_fields = [
    "FoD-Computers_Math_Stats", "FoD-Engineering", "FoD-Physical_Sciences",
    "FoD-Biological_Environmental_Agricultural_Sciences", "FoD-Science_and_Engineering_Related"
]

# Add STEM Degree column (binary)
stack_overflow_results_df['STEM Degree'] = stack_overflow_results_df['ACS_Major'].apply(lambda x: 'Yes' if x in stem_fields else 'No')

# Preview results
print(stack_overflow_results_df[['MajorUndergrad', 'ACS_Major', 'STEM Degree']].head())


# #### Exploratory Data Analysis: Median Salaries of STEM vs. Non-STEM Degree Earners
# - Objective: Compare the median salaries of STEM and non-STEM degree earners to ensure consistency with trends observed in the NCSES and American Community Survey datasets.
# 
# - Methodology:
#     - Filter out rows where the Salary field is NaN.
#     - Group the remaining data by the STEM Degree field and calculate the median salary for each group.
# 
# - Findings: The median salary for respondents with a STEM degree is $95,000, and the median salary for respondents without a STEM degree is $90,000.
# 
# - Implications: This aligns with existing findings from other datasets that STEM degree holders typically earn higher salaries than non-STEM degree holders. However, the gap here appears relatively small, which could warrant further investigation into factors like occupation type and level of education.

# In[74]:


# Filter for rows where Salary is not NaN
valid_salary_data = stack_overflow_results_df[stack_overflow_results_df['Salary'].notna()]

# Group by STEM Degree and calculate median salary
stem_salary_comparison = valid_salary_data.groupby('STEM Degree')['Salary'].median()

# Display the comparison
print("Median Salary for STEM vs. Non-STEM Degrees:")
print(stem_salary_comparison)


# #### Clean Dataset: Update and Consolidate Education Categories
# - Removed records with education levels inconsistent with the project's focus:
#   - I never completed any formal education
#   - I prefer not to answer
#   - Primary/elementary school
#   - Secondary school
#   - Professional degree
#   - Some college/university study without earning a bachelor's degree
# - Remaining education levels were mapped into three categories: *Bachelor's*, *Graduate*, and *Doctorate*.
# - Combined these education levels with the STEM classification to create the following final categories:
#   - Bachelor's - STEM
#   - Bachelor's - Non-STEM
#   - Graduate - STEM
#   - Graduate - Non-STEM
#   - Doctorate - STEM
#   - Doctorate - Non-STEM

# In[75]:


# Define categories to remove
categories_to_remove = [
    "I never completed any formal education",
    "I prefer not to answer",
    "Primary/elementary school",
    "Secondary school",
    "Professional degree",
    "Some college/university study without earning a bachelor's degree"
]

# Filter the dataset to exclude unwanted education levels
stack_overflow_results_df = stack_overflow_results_df[
    ~stack_overflow_results_df['FormalEducation'].isin(categories_to_remove)
]

# Map the remaining education levels to "Bachelor's" or "Graduate"
education_mapping = {
    "Bachelor's degree": "Bachelor's",
    "Master's degree": "Graduate",
    "Doctoral degree": "Doctorate"
}

# Apply the mapping
stack_overflow_results_df['EducationLevel'] = stack_overflow_results_df['FormalEducation'].map(education_mapping)

# Combine EducationLevel with STEM Degree to create final education categories
stack_overflow_results_df['EducationCategory'] = stack_overflow_results_df.apply(
    lambda x: f"{x['EducationLevel']} - {'STEM' if x['STEM Degree'] == 'Yes' else 'Non-STEM'}",
    axis=1
)

# Display updated categories
print("Updated Education Categories:")
print(stack_overflow_results_df['EducationCategory'].value_counts())

# Display the updated dataset size
print("\nRemaining records:", len(stack_overflow_results_df))


# #### Exploratory Data Analysis: Map Education Categories into Buckets for Correlation with Salary
# - The *EducationCategory* column combines both education level and STEM/Non-STEM classification. 
# - A numerical mapping was created to reflect the hierarchy of education levels, with higher values indicating higher education levels and STEM fields ranked above Non-STEM fields.
# - The following mapping was applied:
#   - Bachelor's - Non-STEM: 1
#   - Bachelor's - STEM: 2
#   - Graduate - Non-STEM: 3
#   - Graduate - STEM: 4
#   - Doctorate - Non-STEM: 5
#   - Doctorate - STEM: 6
# - Records with missing salary values were excluded before calculating the correlation.
# 
# **Correlation Matrix:**
# - The correlation between *EducationCategoryNumeric* and *Salary* is 0.232, indicating a moderate positive relationship.
# - This suggests that education level and STEM/Non-STEM classification contribute to salary, but other factors, such as job experience or job title, may have a more significant influence.

# In[76]:


# Define a numerical mapping for EducationCategory
education_category_mapping = {
    "Bachelor's - STEM": 2,
    "Bachelor's - Non-STEM": 1,
    "Graduate - STEM": 4,
    "Graduate - Non-STEM": 3,
    "Doctorate - STEM": 6,
    "Doctorate - Non-STEM": 5
}

# Apply the mapping
stack_overflow_results_df['EducationCategoryNumeric'] = stack_overflow_results_df['EducationCategory'].map(education_category_mapping)

# Filter out rows where Salary is NaN
valid_salary_data = stack_overflow_results_df.dropna(subset=['Salary'])

# Calculate the correlation between EducationCategoryNumeric and Salary
correlation_matrix = valid_salary_data[['EducationCategoryNumeric', 'Salary']].corr()

# Display the correlation
print("Correlation matrix:")
print(correlation_matrix)

# Preview the filtered dataset
print(valid_salary_data[['EducationCategory', 'EducationCategoryNumeric', 'Salary']].head(10))


# #### Exploratory Data Analysis: Identify Inconsistencies Between DeveloperType, WebDeveloperType, and MobileDeveloperType
# - Objective: Verify consistency across developer-related columns (*DeveloperType*, *WebDeveloperType*, and *MobileDeveloperType*) to ensure data reliability.
# 
# - Steps:
#   1. Check for Missing Related Fields:
#       - Identified rows where *DeveloperType* has a value but both *WebDeveloperType* and *MobileDeveloperType* are NaN.
#         - Total Inconsistent Records: 3629.
#       - These inconsistencies suggest that specific developer roles (e.g., Web or Mobile) are not recorded in the corresponding fields.
#   2. Check for Reverse Inconsistencies: Identified rows where *DeveloperType* is NaN but either *WebDeveloperType* or *MobileDeveloperType* has a value.
#       - Total Reverse Inconsistent Records: 0.

# In[77]:


# Identify rows with DeveloperType having a value but related fields being NaN
inconsistent_records = stack_overflow_results_df[
    (stack_overflow_results_df['DeveloperType'].notna()) &
    (stack_overflow_results_df['WebDeveloperType'].isna()) &
    (stack_overflow_results_df['MobileDeveloperType'].isna())
]

# Display the inconsistent records
print("Inconsistent records (DeveloperType has value, others are NaN):")
print(inconsistent_records[['DeveloperType', 'WebDeveloperType', 'MobileDeveloperType']].head())

# Count the number of inconsistent records
print(f"Total inconsistent records: {len(inconsistent_records)}")


# In[78]:


# Display a random sample of inconsistent records
sample_inconsistent_records = inconsistent_records[['DeveloperType', 'WebDeveloperType', 'MobileDeveloperType', 'NonDeveloperType']].sample(10)
print(sample_inconsistent_records)


# In[79]:


# Filter for records where DeveloperType is NaN but either of the related columns have a value
reverse_inconsistent_records = stack_overflow_results_df[
    stack_overflow_results_df['DeveloperType'].isna() &
    (
        stack_overflow_results_df['WebDeveloperType'].notna() |
        stack_overflow_results_df['MobileDeveloperType'].notna() 
    )
]

# Count the records
print(f"Number of reverse inconsistent records: {len(reverse_inconsistent_records)}")


# #### Exploratory Data Analysis: Determine If There Is a Strong Correlation Between Mobile Developer Type and Salary to Decide Whether to Keep or Drop the Column
# 
# - Objective: Assess the correlation between *MobileDeveloperType* and *Salary* to determine if this feature is valuable for predictive modeling.
# 
# - Steps:
#   1. Mapped *MobileDeveloperType* into ordinal numeric values to represent the level of expertise or range of platforms:
#       - iOS; Android; Windows Phone: 3  
#       - iOS; Android: 2  
#       - iOS; Windows Phone: 2  
#       - iOS or Android: 1  
#       - NaN or missing values were retained as NaN.
#   2. Filtered the dataset to include only rows where both *Salary* and *MobileDeveloperLevel* are not NaN.
#   3. Calculated the correlation between *MobileDeveloperLevel* and *Salary*.
# 
# - Conclusion:
#     - The correlation coefficient between MobileDeveloperLevel and Salary is approximately -0.0566, indicating no meaningful relationship.
#     - Decision: Drop the MobileDeveloperType column from the dataset, as it does not contribute valuable predictive information.

# In[80]:


# Define the mapping
mobile_developer_mapping = {
    "iOS; Android; Windows Phone": 3,
    "iOS; Android": 2,
    "iOS; Windows Phone": 2,
    "iOS": 1,
    "Android": 1,
    None: None  # Handle NaN values
}

# Apply the mapping to create a numerical column
stack_overflow_results_df['MobileDeveloperLevel'] = stack_overflow_results_df['MobileDeveloperType'].map(mobile_developer_mapping)

# Filter the dataset for non-NaN Salary and MobileDeveloperLevel
valid_data = stack_overflow_results_df[
    stack_overflow_results_df['Salary'].notna() & stack_overflow_results_df['MobileDeveloperLevel'].notna()
]

# Calculate correlation with salary
correlation_matrix = valid_data[['MobileDeveloperLevel', 'Salary']].corr()

# Display the correlation
print("Correlation matrix (only where Salary and MobileDeveloperLevel are not NaN):")
print(correlation_matrix)

# Preview the updated dataset for valid rows
print(valid_data[['MobileDeveloperType', 'MobileDeveloperLevel', 'Salary']].head())


# #### Exploratory Data Analysis: Determine If There Is a Strong Correlation Between Web Developer Type and Salary to Decide Whether to Keep or Drop the Column
# 
# - Objective: Assess the correlation between *WebDeveloperType* and *Salary* to determine if this feature is valuable for predictive modeling.
# 
# - Steps:
#   1. Mapped *WebDeveloperType* into ordinal numeric values to represent specialization levels:
#       - Full stack Web developer: 3  
#       - Back-end Web developer: 2  
#       - Front-end Web developer: 1  
#       - NaN or missing values were retained as NaN.
#   2. Filtered the dataset to include only rows where both *Salary* and *WebDeveloperLevel* are not NaN.
#   3. Calculated the correlation between *WebDeveloperLevel* and *Salary*.
# 
# - Conclusion:
#     - The correlation coefficient between WebDeveloperLevel and Salary is approximately 0.0244, indicating no meaningful relationship.
#     - Decision: Drop the WebDeveloperType column from the dataset, as it does not contribute valuable predictive information.

# In[81]:


# Map WebDeveloperType to numerical values
web_developer_mapping = {
    "Full stack Web developer": 3,
    "Back-end Web developer": 2,
    "Front-end Web developer": 1,
    None: None  # Retain NaNs for now
}

# Apply the mapping
stack_overflow_results_df['WebDeveloperLevel'] = stack_overflow_results_df['WebDeveloperType'].map(web_developer_mapping)

# Filter out rows where Salary or WebDeveloperLevel is NaN
valid_web_dev_data = stack_overflow_results_df.dropna(subset=['WebDeveloperLevel', 'Salary'])

# Calculate the correlation matrix on filtered data
correlation_matrix = valid_web_dev_data[['WebDeveloperLevel', 'Salary']].corr()

# Display the correlation matrix
print("Correlation matrix (only where Salary and WebDeveloperLevel are not NaN):")
print(correlation_matrix)

# Display sample records for review
print(valid_web_dev_data[['WebDeveloperType', 'WebDeveloperLevel', 'Salary']].head(10))


# #### Exploratory Data Analysis: Determine If There Is a Correlation Between Remote Work and Salary
# 
# - Objective: Assess the correlation between *HomeRemote* and *Salary* to determine if remote work arrangements influence salary levels.
# 
# - Steps:
#   1. Mapped *HomeRemote* values into ordinal numeric levels to represent the degree of remote work:
#       - All or almost all the time (I'm full-time remote): 6  
#       - More than half, but not all, the time: 5  
#       - About half the time: 4  
#       - Less than half the time, but at least one day each week: 3  
#       - A few days each month: 2  
#       - Never: 1  
#       - "It's complicated" and NaN values were retained as missing.
#   2. Filtered the dataset to include only rows where both *Salary* and *HomeRemoteLevel* are not NaN.
#   3. Calculated the correlation between *HomeRemoteLevel* and *Salary*.
# 
# - Conclusion:
#     - The correlation coefficient between *HomeRemoteLevel* and *Salary* is approximately 0.1719, indicating a weak positive relationship.
#     - Decision: Retain the *HomeRemote* column, as it may provide some predictive value in modeling.

# In[82]:


# Map HomeRemote to numerical values based on the provided ranking
home_remote_mapping = {
    "All or almost all the time (I'm full-time remote)": 6,
    "More than half, but not all, the time": 5,
    "About half the time": 4,
    "Less than half the time, but at least one day each week": 3,
    "A few days each month": 2,
    "Never": 1,
    "It's complicated": None,  # Treat as missing for now
    None: None  # Leave NaNs as missing
}

# Apply the mapping
stack_overflow_results_df['HomeRemoteLevel'] = stack_overflow_results_df['HomeRemote'].map(home_remote_mapping)

# Filter out rows where Salary or HomeRemoteLevel is NaN
valid_home_remote_data = stack_overflow_results_df.dropna(subset=['HomeRemoteLevel', 'Salary'])

# Calculate the correlation matrix on filtered data
correlation_matrix = valid_home_remote_data[['HomeRemoteLevel', 'Salary']].corr()

# Display the correlation matrix
print("Correlation matrix (only where Salary and HomeRemoteLevel are not NaN):")
print(correlation_matrix)

# Display sample records for review
print(valid_home_remote_data[['HomeRemote', 'HomeRemoteLevel', 'Salary']].head(10))


# #### Exploratory Data Analysis: Determine If There Is a Correlation Between Gender and Salary
# 
# - Objective: Assess the correlation between *Gender* and *Salary* to identify any potential gender-based salary disparities.
# 
# - Steps:
#   1. Mapped *Gender* values into numeric levels:
#       - Male: 1  
#       - Female: 0  
#   2. Filtered the dataset to include only rows where *Salary* is not NaN.
#   3. Calculated the correlation between *GenderNumeric* and *Salary*.
# 
# - Conclusion:
#     - The correlation coefficient between *GenderNumeric* and *Salary* is approximately 0.0667, indicating a negligible positive relationship.
#     - Decision: While the correlation is weak, it is appropriate to retain *Gender* in the dataset for further analysis, as gender may interact with other features to influence salary.

# In[83]:


# Map Gender to numerical values
gender_mapping = {
    "Male": 1,
    "Female": 0
}

# Apply the mapping
stack_overflow_results_df['GenderNumeric'] = stack_overflow_results_df['Gender'].map(gender_mapping)

# Filter out rows where Salary is NaN
valid_salary_data = stack_overflow_results_df[stack_overflow_results_df['Salary'].notna()]

# Calculate correlation with Salary
correlation_matrix = valid_salary_data[['GenderNumeric', 'Salary']].corr()

# Display the correlation matrix
print("Correlation matrix:")
print(correlation_matrix)

# Display sample records for review
print(valid_salary_data[['Gender', 'GenderNumeric', 'Salary']].head(10))


# #### Exploratory Data Analysis: Determine If There Is a Correlation Between University Enrollment and Salary
# - Objective: Evaluate whether current university enrollment status is correlated with salary.
# 
# - Steps:
#     1. Mapped the *University* column to numerical values:
#         - Yes, full-time or Yes, part-time: 1 (Enrolled)
#         - No or I prefer not to say: 0 (Not enrolled).
#     2. Filtered the dataset to include only rows where *Salary* is not NaN.
#     3. Calculated the correlation between *UniversityEnrolled* and *Salary*.
# 
# - Conclusion:
#     - The correlation coefficient between *University Enrollment* and *Salary* is approximately -0.134, indicating a weak negative relationship.
#     - Decision: The weak correlation suggests that *University Enrollment* may not significantly contribute to salary prediction. The feature will likely be dropped from the dataset to simplify the model.   

# In[84]:


# Define the mapping for University column
university_mapping = {
    "Yes, full-time": 1,
    "Yes, part-time": 1,
    "No": 0,
    "I prefer not to say": 0
}

# Apply the mapping
stack_overflow_results_df['UniversityEnrolled'] = stack_overflow_results_df['University'].map(university_mapping)

# Filter out rows where Salary is NaN
valid_salary_data = stack_overflow_results_df[stack_overflow_results_df['Salary'].notna()]

# Calculate the correlation matrix with salary
correlation_matrix = valid_salary_data[['UniversityEnrolled', 'Salary']].corr()

# Display the correlation matrix
print("Correlation matrix:")
print(correlation_matrix)

# Preview the dataset to confirm changes
print(stack_overflow_results_df[['University', 'UniversityEnrolled', 'Salary']].head(10))


# #### Exploratory Data Analysis: Determine if There Is a Correlation Between Years Coding Professionally and Salary
# - Objective: Assess the correlation between *YearsCodedJob* and *Salary* to evaluate its predictive value for salary modeling.
# 
# - Steps:
#     1. Mapped the *YearsCodedJob* values into ordinal numeric values representing the number of years coded professionally.
#     2. Filtered the dataset to include only rows where both *YearsCodedJobNumeric* and *Salary* are not NaN.
#     3. Calculated the correlation between *YearsCodedJobNumeric* and *Salary*.
# 
# - Findings::
#     - Correlation coefficient: 0.5176, indicating a moderate-to-strong positive relationship.
#     - This is the strongest correlation observed so far in this analysis and will be retained in the dataset. 

# In[85]:


# Define the mapping for YearsCodedJob
years_coded_job_mapping = {
    "Less than a year": 0,
    "1 to 2 years": 1,
    "2 to 3 years": 2,
    "3 to 4 years": 3,
    "4 to 5 years": 4,
    "5 to 6 years": 5,
    "6 to 7 years": 6,
    "7 to 8 years": 7,
    "8 to 9 years": 8,
    "9 to 10 years": 9,
    "10 to 11 years": 10,
    "11 to 12 years": 11,
    "12 to 13 years": 12,
    "13 to 14 years": 13,
    "14 to 15 years": 14,
    "15 to 16 years": 15,
    "16 to 17 years": 16,
    "17 to 18 years": 17,
    "18 to 19 years": 18,
    "19 to 20 years": 19,
    "20 or more years": 20,
    None: None  # Handle NaN values
}

# Apply the mapping to create a numerical column
stack_overflow_results_df['YearsCodedJobNumeric'] = stack_overflow_results_df['YearsCodedJob'].map(years_coded_job_mapping)

# Filter the dataset for rows where both YearsCodedJobNumeric and Salary are not NaN
valid_years_coded_job_data = stack_overflow_results_df.dropna(subset=['YearsCodedJobNumeric', 'Salary'])

# Calculate the correlation matrix on filtered data
correlation_matrix = valid_years_coded_job_data[['YearsCodedJobNumeric', 'Salary']].corr()

# Display the correlation matrix
print("Correlation matrix (only where Salary and YearsCodedJobNumeric are not NaN):")
print(correlation_matrix)

# Display sample records for review
print(valid_years_coded_job_data[['YearsCodedJob', 'YearsCodedJobNumeric', 'Salary']].head(10))


# #### Exploratory Data Analysis: Determine If There Is a Correlation Between Company Type and Salary
# - Objective: Assess the correlation between *CompanyType* and *Salary* to determine if *CompanyType* is a meaningful predictor.
# 
# - Steps:
#     1. Mapped CompanyType to ordinal numeric values reflecting the potential salary hierarchy:
#         - Publicly-traded corporation: 8
#         - Privately-held limited company: 7
#         - Venture-funded startup: 6
#         - Government agency or public school/university: 5
#         - State-owned company: 4
#         - Non-profit organization or private school/university: 3
#         - Sole proprietorship or partnership: 2
#         - Pre-series A startup: 1
#         - Other or unknown values (e.g., "I don't know"): excluded.
#     2. Filtered for rows where both *Salary* and *CompanyTypeNumeric* are not NaN.
#     3. Calculated the correlation between *CompanyTypeNumeric* and *Salary*.
# 
# - Findings::
#     - Correlation coefficient: 0.1996, indicating a moderate positive relationship.

# In[86]:


# Display unique values and counts in the CompanyType column
company_type_summary = stack_overflow_results_df['CompanyType'].value_counts(dropna=False)
print(company_type_summary)


# In[87]:


# Define mapping for CompanyType
company_type_mapping = {
    "Publicly-traded corporation": 8,
    "Privately-held limited company, not in startup mode": 7,
    "Venture-funded startup": 6,
    "Government agency or public school/university": 5,
    "State-owned company": 4,
    "Non-profit/non-governmental organization or private school/university": 3,
    "Sole proprietorship or partnership, not in startup mode": 2,
    "Pre-series A startup": 1,
    "I don't know": None,
    "I prefer not to answer": None,
    "Something else": None,
    None: None
}

# Apply the mapping
stack_overflow_results_df['CompanyTypeNumeric'] = stack_overflow_results_df['CompanyType'].map(company_type_mapping)

# Filter out rows where Salary or CompanyTypeNumeric is NaN
valid_company_data = stack_overflow_results_df.dropna(subset=['Salary', 'CompanyTypeNumeric'])

# Calculate correlation matrix
correlation_matrix = valid_company_data[['CompanyTypeNumeric', 'Salary']].corr()

# Display correlation matrix
print("Correlation matrix (only where Salary and CompanyTypeNumeric are not NaN):")
print(correlation_matrix)

# Display sample records for review
print(valid_company_data[['CompanyType', 'CompanyTypeNumeric', 'Salary']].head(10))


# #### Exploratory Data Analysis: Evaluate NonDeveloperType Column for Usefulness in Salary Prediction
# - Objective: Assess whether *NonDeveloperType* provides meaningful data for salary prediction.
# 
# - Steps:
#     1. Attempted to calculate a correlation between *NonDeveloperType* (mapped to numeric values) and *Salary*. The correlation matrix consistently returned NaN values.
#     2. Investigated whether any records in *NonDeveloperType* also had corresponding values in *Salary*:
#         - Number of *NonDeveloperType* records with a valid *Salary*: 0
#         - Without salary data, the *NonDeveloperType* field provides no basis for predictive modeling.
#     3. Checked for overlap between *NonDeveloperType* and *DeveloperType*:
#         - Confirmed no records exist in both columns, verifying they represent distinct groups.
# 
# - Findings:
#     - The lack of salary data makes the *NonDeveloperType* column irrelevant for modeling.
#     - Action: Drop the *NonDeveloperType* column from the dataset.

# In[88]:


# Check for NonDeveloperType records with a valid Salary
nondeveloper_with_salary = stack_overflow_results_df[
    stack_overflow_results_df['NonDeveloperType'].notna() & stack_overflow_results_df['Salary'].notna()
]

# Count the number of records
print(f"Number of NonDeveloperType records with a valid Salary: {len(nondeveloper_with_salary)}")

# Display sample records with NonDeveloperType and Salary
print("Sample records with NonDeveloperType and Salary:")
print(nondeveloper_with_salary[['NonDeveloperType', 'Salary']].head(10))


# #### Exploratory Data Analysis: Determine If There Is a Correlation Between DeveloperType and Salary
# - Objective: Assess the correlation between *DeveloperType* and *Salary* to understand whether this field has predictive value in its current form.
# 
# - Steps:
#     1. Assigned ranks to each developer role based on industry salary trends:
#         - Machine learning specialist: 14  
#         - Data scientist: 13  
#         - Embedded applications/devices developer: 12  
#         - Graphics programmer: 11  
#         - DevOps specialist: 10  
#         - Web developer: 9  
#         - Systems administrator: 8  
#         - Database administrator: 7  
#         - Desktop applications developer: 6  
#         - Mobile developer: 5  
#         - Graphic designer: 4  
#         - Quality assurance engineer: 3  
#         - Something else: 1  
#     2. Mapped the highest-ranking role for each respondent in the *DeveloperType* field to a numeric value.
#     3. Calculated the correlation between the primary role rank and salary.
# 
# - Findings:
#     - Correlation coefficient: 0.0563, indicating an extremely weak relationship in the current implementation.
#     - Despite this weak correlation, real-world evidence demonstrates that salaries differ significantly across roles. Therefore, this field has potential value when combined with other features in later modeling steps.
#     - Next Steps: Explore more sophisticated methods, such as weighted combinations of roles or one-hot encoding, to better capture the impact of developer roles on salary.

# In[89]:


# Define the ranking for DeveloperType roles
developer_role_ranking = {
    "Machine learning specialist": 14,
    "Data scientist": 13,
    "Embedded applications/devices developer": 12,
    "Graphics programmer": 11,
    "DevOps specialist": 10,
    "Web developer": 9,
    "Systems administrator": 8,
    "Database administrator": 7,
    "Desktop applications developer": 6,
    "Mobile developer": 5,
    "Graphic designer": 4,
    "Quality assurance engineer": 3,
    "Something else": 1
}

# Extract the primary role from DeveloperType based on the highest rank
def get_primary_role(dev_type):
    for role, rank in sorted(developer_role_ranking.items(), key=lambda x: x[1], reverse=True):
        if role in dev_type:
            return role
    return None

# Apply the primary role extraction
stack_overflow_results_df['PrimaryRole'] = stack_overflow_results_df['DeveloperType'].apply(
    lambda x: get_primary_role(str(x)) if pd.notna(x) else None
)

# Map the primary role to numeric values
stack_overflow_results_df['PrimaryRoleNumeric'] = stack_overflow_results_df['PrimaryRole'].map(developer_role_ranking)

# Filter for rows where Salary and PrimaryRoleNumeric are not NaN
valid_primary_role_data = stack_overflow_results_df.dropna(subset=['PrimaryRoleNumeric', 'Salary'])

# Calculate correlation between PrimaryRoleNumeric and Salary
correlation_matrix = valid_primary_role_data[['PrimaryRoleNumeric', 'Salary']].corr()

# Display the correlation matrix
print("Correlation matrix (only where Salary and PrimaryRoleNumeric are not NaN):")
print(correlation_matrix)

# Preview the valid data
print(valid_primary_role_data[['DeveloperType', 'PrimaryRole', 'PrimaryRoleNumeric', 'Salary']].head())


# ### Narrowing Down the Dataset
# 
# - Objective: Reduce the dataset to include only the most relevant columns for machine learning modeling. 
# 
# - Steps:
#     1. Displayed a random sample of the dataset and the column names to review the current structure.
#     2. Defined a list of columns to retain based on the findings from exploratory data analysis:
#         - *Gender*: Retained for potential influence on salary.
#         - *Professional*: Retained to distinguish between developer and non-developer roles.
#         - *HomeRemote*: Retained despite a weak correlation, as remote work trends may still influence salary predictions.
#         - *CompanySize*: Retained despite a weak correlation, as company size trends may still influence salary predictions.
#         - *CompanyType*: Retained to evaluate its moderate correlation with salary.
#         - *YearsCodedJobNumeric*: Retained as it showed the strongest correlation with salary.
#         - *EducationCategoryNumeric*: Retained for its hierarchical structure related to salary.
#         - *DeveloperType*: Retained for potential predictive power once properly encoded.
#         - *Salary*: Target variable for prediction.
#     3. Filtered the dataset to include only these columns.
# 
# - Next Steps:
#     - Review the refined dataset structure to confirm accuracy.
#     - Proceed with preprocessing, including encoding and scaling, to prepare the dataset for machine learning modeling.

# In[90]:


display(stack_overflow_results_df.sample(10))

display(list(stack_overflow_results_df.columns))


# In[91]:


# Define columns to keep
columns_to_keep = [
    'Gender',
    'Professional',
    'HomeRemote',
    'CompanySize',
    'CompanyType',
    'YearsCodedJobNumeric',
    'EducationCategoryNumeric',
    'DeveloperType',
    'Salary'
]

# Narrow down the dataset
stack_overflow_results_df = stack_overflow_results_df[columns_to_keep]

# Display the new dataframe's structure
print("Updated dataframe columns:")
display(list(stack_overflow_results_df.columns))
print("\nRemaining records:", len(stack_overflow_results_df))
display(stack_overflow_results_df.head())


# #### Refining the Dataset for Salary Prediction
# - Objective: Ensure that the dataset is properly filtered and structured for accurate feature analysis and predictive modeling.
# 
# - Steps:
#    1. Filtered Out Rows Without Developer Roles
#       - Removed rows where *DeveloperType* was *NaN* to ensure only respondents with relevant roles are included.
#       - This step prevents irrelevant or missing roles from skewing analysis or predictions.
#    2. Reassessed Data Distribution
#       - Reviewed the unique values and distributions for key features
#       - Identified imbalances (e.g., overwhelming representation of males in *Gender*), which may need to be addressed during modeling.
#    3. Dropped Irrelevant Column:
#       - Removed the *Professional* column as it was no longer needed after filtering for relevant developer roles.
#    4. Evaluated Multiple Roles in *DeveloperType*
#       - Calculated the number of roles listed for each respondent in the DeveloperType field.
#       - Filtered rows where respondents listed multiple roles (e.g., "Web developer; Mobile developer").
#       - Checked for rows with multiple roles that also had valid salary values to evaluate their relevance:
#          - Identified 1,355 rows where multiple roles were listed along with valid salary data.
#       
# Findings:
# - The dataset now focuses exclusively on respondents with valid developer roles.
# - *DeveloperType* requires further refinement to handle multiple roles effectively in the next steps.
# 
# Next Steps:
# - Explore options for handling multiple roles in DeveloperType (e.g., selecting a primary role, creating a composite score, or applying one-hot encoding).
# - Recalculate correlations for refined and encoded features.
# - Prepare the dataset for machine learning by encoding categorical fields and applying normalization or scaling.

# In[92]:


# Filter out rows where DeveloperType is NaN
stack_overflow_results_df = stack_overflow_results_df[
    stack_overflow_results_df['DeveloperType'].notna()
]

# Check the updated record count
print(f"Remaining records after filtering out rows with no DeveloperType: {len(stack_overflow_results_df)}")


# In[93]:


display(stack_overflow_results_df['Gender'].value_counts())
display(stack_overflow_results_df['Professional'].value_counts())
display(stack_overflow_results_df['HomeRemote'].value_counts())
display(stack_overflow_results_df['CompanySize'].value_counts())
display(stack_overflow_results_df['CompanyType'].value_counts())
display(stack_overflow_results_df['YearsCodedJobNumeric'].value_counts())
display(stack_overflow_results_df['EducationCategoryNumeric'].value_counts())
display(stack_overflow_results_df['DeveloperType'].value_counts())


# In[94]:


# Drop the Professional column
stack_overflow_results_df.drop(columns=['Professional'], inplace=True)
display(list(stack_overflow_results_df.columns))


# In[95]:


# Identify rows with multiple roles in DeveloperType
stack_overflow_results_df['RoleCount'] = stack_overflow_results_df['DeveloperType'].apply(
    lambda x: len(str(x).split("; ")) if pd.notna(x) else 0
)

# Filter rows with multiple roles
multiple_roles = stack_overflow_results_df[stack_overflow_results_df['RoleCount'] > 1]

# Check if any rows with multiple roles have a valid Salary
multiple_roles_with_salary = multiple_roles[multiple_roles['Salary'].notna()]
print(f"Number of rows with multiple DeveloperType roles and a valid Salary: {len(multiple_roles_with_salary)}")

# Display a sample of rows with multiple roles and Salary
if not multiple_roles_with_salary.empty:
    display(multiple_roles_with_salary[['DeveloperType', 'Salary']].head())
else:
    print("No rows with multiple DeveloperType roles have a valid Salary.")


# ### Encode Categorical Features

# In[96]:


# Extract the Gender column
gender_data = stack_overflow_results_df[['Gender']]

# Initialize OneHotEncoder
one_hot_encoder = OneHotEncoder()

# Fit and transform the Gender column
gender_encoded = one_hot_encoder.fit_transform(gender_data)

# Convert the encoded data to a DataFrame
gender_encoded_df = pd.DataFrame(
    gender_encoded.toarray(), 
    columns=one_hot_encoder.get_feature_names_out(['Gender'])
)

# Concatenate the encoded Gender columns with the main DataFrame
encoded_stack_overflow_df = pd.concat(
    [stack_overflow_results_df.reset_index(drop=True), gender_encoded_df], 
    axis=1
)

# Drop the original Gender column
encoded_stack_overflow_df.drop(columns=['Gender'], inplace=True)

# Display the updated DataFrame
print("Updated DataFrame columns after encoding Gender:")
print(list(encoded_stack_overflow_df.columns))

# Preview the updated DataFrame
display(encoded_stack_overflow_df.sample(10))


# **Calculate mean, median, and count of *HomeRemote* to determine how to encode the values**

# In[97]:


# Replace NaN with 'Unknown' in HomeRemote
encoded_stack_overflow_df['HomeRemote'] = encoded_stack_overflow_df['HomeRemote'].fillna('Unknown')

# Calculate mean, median, and count for each HomeRemote category
home_remote_salary = encoded_stack_overflow_df.groupby('HomeRemote')['Salary'].agg(['mean', 'median', 'count']).sort_index()

# Display results
display(home_remote_salary)


# In[98]:


# Combine "About half the time," "Less than half the time, but at least one day each week," 
# and "More than half, but not all, the time" into "Partial Remote"
encoded_stack_overflow_df['HomeRemote'] = encoded_stack_overflow_df['HomeRemote'].replace(
    {
        "About half the time": "Partial Remote",
        "Less than half the time, but at least one day each week": "Partial Remote",
        "More than half, but not all, the time": "Partial Remote"
    }
)

# Drop the 'Unknown' category
encoded_stack_overflow_df = encoded_stack_overflow_df[encoded_stack_overflow_df['HomeRemote'] != "Unknown"]

# Define the categories in the order of their desired ranking
home_remote_categories = [
    "Never",
    "It's complicated",
    "A few days each month",
    "Partial Remote",
    "All or almost all the time (I'm full-time remote)"
]

# Initialize the OrdinalEncoder
ordinal_encoder = OrdinalEncoder(categories=[home_remote_categories])

# Apply ordinal encoding directly to HomeRemote column
encoded_stack_overflow_df['HomeRemoteEncoded'] = ordinal_encoder.fit_transform(
    encoded_stack_overflow_df[['HomeRemote']]
)

# Verify the encoding and display the results
display(encoded_stack_overflow_df[['HomeRemote', 'HomeRemoteEncoded']].head())


# **Calculate mean, median, and count of *CompanySize* to determine how to encode the values**

# In[99]:


# Replace NaN with 'Unknown' in CompanySize
encoded_stack_overflow_df['CompanySize'] = encoded_stack_overflow_df['CompanySize'].fillna('Unknown')

# Calculate mean, median, and count for each CompanySize category
company_size_salary = encoded_stack_overflow_df.groupby('CompanySize')['Salary'].agg(['mean', 'median', 'count']).sort_index()

# Display results
display(company_size_salary)


# In[100]:


# Combine "100 to 499 employees" and "500 to 999 employees" into "Mid-Sized"
encoded_stack_overflow_df['CompanySize'] = encoded_stack_overflow_df['CompanySize'].replace(
    {
        "100 to 499 employees": "Mid-Sized",
        "500 to 999 employees": "Mid-Sized"
    }
)

# Drop ambiguous categories: "Unknown," "I prefer not to answer"
encoded_stack_overflow_df = encoded_stack_overflow_df[
    ~encoded_stack_overflow_df['CompanySize'].isin(["Unknown", "I prefer not to answer"])
]

# Rank "I don't know" lowest to retain those records
encoded_stack_overflow_df['CompanySize'] = encoded_stack_overflow_df['CompanySize'].replace(
    {"I don't know": "Lowest Rank"}
)

# Define the categories in the order of their desired ranking
company_size_categories = [
    "Lowest Rank",  
    "Fewer than 10 employees",
    "10 to 19 employees",
    "20 to 99 employees",
    "Mid-Sized",
    "1,000 to 4,999 employees",
    "5,000 to 9,999 employees",
    "10,000 or more employees"
]

# Initialize the OrdinalEncoder
ordinal_encoder = OrdinalEncoder(categories=[company_size_categories])

# Apply ordinal encoding directly to CompanySize column
encoded_stack_overflow_df['CompanySizeEncoded'] = ordinal_encoder.fit_transform(
    encoded_stack_overflow_df[['CompanySize']]
)

# Verify the encoding and display the results
display(encoded_stack_overflow_df[['CompanySize', 'CompanySizeEncoded']].head())


# **Calculate mean, median, and count of *CompanyType* to determine how to encode the values**

# In[101]:


# Replace NaN with 'Unknown' in CompanyType
encoded_stack_overflow_df['CompanyType'] = encoded_stack_overflow_df['CompanyType'].fillna('Unknown')

# Calculate mean, median, and count for each CompanyType category
company_type_salary = encoded_stack_overflow_df.groupby('CompanyType')['Salary'].agg(['mean', 'median', 'count']).sort_index()

# Display results
display(company_type_salary)


# In[102]:


# Combine "Pre-series A startup" and "Privately-held limited company" into "Privately-held or Pre-series A"
encoded_stack_overflow_df['CompanyType'] = encoded_stack_overflow_df['CompanyType'].replace(
    {
        "Pre-series A startup": "Privately-held or Pre-series A",
        "Privately-held limited company, not in startup mode": "Privately-held or Pre-series A",
    }
)

# Combine "Non-profit" and "Sole Proprietorship" into "Non-profit or Sole Proprietorship"
encoded_stack_overflow_df['CompanyType'] = encoded_stack_overflow_df['CompanyType'].replace(
    {
        "Non-profit/non-governmental organization or private school/university": "Non-profit or Sole Proprietorship",
        "Sole proprietorship or partnership, not in startup mode": "Non-profit or Sole Proprietorship",
    }
)

# Drop ambiguous categories: "Unknown," "I prefer not to answer", and "Something else"
encoded_stack_overflow_df = encoded_stack_overflow_df[
    ~encoded_stack_overflow_df['CompanyType'].isin(["Unknown", "I prefer not to answer", "Something else"])
]

# Rank "I don't know" lowest to retain those records
encoded_stack_overflow_df['CompanyType'] = encoded_stack_overflow_df['CompanyType'].replace(
    {"I don't know": "Lowest Rank"}
)

# Define the categories in the order of their desired ranking
company_type_categories = [
    "Lowest Rank",  
    "State-owned company",
    "Government agency or public school/university",
    "Non-profit or Sole Proprietorship",
    "Privately-held or Pre-series A",
    "Venture-funded startup",
    "Publicly-traded corporation"
]

# Initialize the OrdinalEncoder
ordinal_encoder = OrdinalEncoder(categories=[company_type_categories])

# Apply ordinal encoding directly to CompanyType column
encoded_stack_overflow_df['CompanyTypeEncoded'] = ordinal_encoder.fit_transform(
    encoded_stack_overflow_df[['CompanyType']]
)

# Verify the encoding and display the results
display(encoded_stack_overflow_df[['CompanyType', 'CompanyTypeEncoded']].head())


# **Calculate mean, median, and count of *EducationCategory* to verify numeric mapping is accurately ordered**
# - Although the mean and median of category 5 (Master's Degree: Non-STEM) is slightly lower than that of category 4 (Bachelor's Degree: STEM), the sample size is very small. Maintaining the existing order is consistent with the trends observed in other datasets with larger sample sizes.

# In[103]:


# Replace NaN with 'Unknown' in EducationCategory
encoded_stack_overflow_df['EducationCategoryNumeric'] = encoded_stack_overflow_df['EducationCategoryNumeric'].fillna('Unknown')

# Calculate mean, median, and count for each EducationCategoryNumeric category
education_salary = encoded_stack_overflow_df.groupby('EducationCategoryNumeric')['Salary'].agg(['mean', 'median', 'count']).sort_index()

# Display results
display(education_salary)


# #### Processing the *DeveloperType* Column
# - Objective: Understand the complexity of the *DeveloperType* column to ensure it is appropriately represented in the dataset for analysis and modeling.
# 
# - Initial Exploration:
#    - The *DeveloperType* column contains 517 unique combinations of roles, indicating significant diversity in the dataset.
#    - With such a large number of combinations, initial analysis focused on identifying those with a sufficient sample size to draw meaningful conclusions for the type of encoding to use.
#    - Filtering for combinations with more than 10 responses reduced the number of groups to 47.
# 
# - Alphabetization and Re-Evaluation:
#    - Roles within each response were alphabetized to standardize combinations and eliminate duplicates caused by order differences.
#    - Example: *"Mobile developer; Web developer"* and *"Web developer; Mobile developer"* were treated as the same group after sorting.
#    - After alphabetization, the same 47 unique combinations with more than 10 responses were identified, confirming the output was processed correctly and that the role order did not affect grouping.
# 
# - Distribution Analysis:
#    - The most common single role is *Web developer* (1,410 responses), followed by multi-role combinations involving *Desktop applications developer* and *Web developer*.
#    - Many combinations occur infrequently, with a long-tail distribution of low-sample groups.
# 
# - Findings:
#    - Filtering and alphabetization revealed that the role order in the *DeveloperType* column does not represent importance.
#    - A focused subset of 47 combinations with sufficient sample sizes can be used for further analysis.
# 
# - Next Steps:
#    - Explore strategies for encoding multi-role combinations effectively.
#    - Consider methods for handling low-frequency combinations to prevent overfitting while preserving representation in the dataset.

# In[104]:


# Sort roles alphabetically within each response
encoded_stack_overflow_df['SortedDeveloperType'] = (
    encoded_stack_overflow_df['DeveloperType']
    .dropna()  # Handle NaN values
    .apply(lambda x: "; ".join(sorted(x.split("; "))))
)

# Count occurrences of the standardized combinations
sorted_developer_type_counts = (
    encoded_stack_overflow_df['SortedDeveloperType']
    .value_counts()
    .reset_index()
    .rename(columns={'index': 'SortedDeveloperType', 'SortedDeveloperType': 'Count'})
)

display(sorted_developer_type_counts)


# In[105]:


# Count occurrences of each unique DeveloperType combination
developer_type_counts = (
    encoded_stack_overflow_df['DeveloperType']
    .value_counts()
    .reset_index()
    .rename(columns={0: 'count', 'index': 'DeveloperType'})
)

# Verify the structure of the resulting DataFrame
print(developer_type_counts.head())

# Filter combinations with more than 10 responses
more_than_ten = developer_type_counts[developer_type_counts['count'] > 10]

# Display the number of combinations with more than 10 responses
print(f"Number of unique DeveloperType combinations with more than 10 responses: {len(more_than_ten)}")

# Display the filtered DataFrame
display(more_than_ten)


# In[106]:


# Alphabetize roles within each DeveloperType response
encoded_stack_overflow_df['DeveloperTypeAlphabetized'] = (
    encoded_stack_overflow_df['DeveloperType']
    .apply(lambda x: "; ".join(sorted(x.split("; "))) if pd.notna(x) else x)
)

# Count occurrences of each unique DeveloperType combination
alphabetized_counts = (
    encoded_stack_overflow_df['DeveloperTypeAlphabetized']
    .value_counts()
    .reset_index()
    .rename(columns={0: 'count', 'index': 'DeveloperTypeAlphabetized'})
)

# Filter combinations with more than 10 responses
more_than_ten_alphabetized = alphabetized_counts[alphabetized_counts['count'] > 10]

# Display the results
print(f"Number of unique alphabetized DeveloperType combinations with more than 10 responses: {len(more_than_ten_alphabetized)}")
display(more_than_ten_alphabetized)


# **Exploring Single Roles**
# - Analyzed respondents with only one role in the *DeveloperType* column to identify base salary trends for individual roles.
# - Findings: Salaries varied significantly by role, with *Web developer* being the most common single role (843 responses).

# In[107]:


# Identify rows where only a single job role is chosen
encoded_stack_overflow_df['RoleCount'] = encoded_stack_overflow_df['DeveloperType'].apply(
    lambda x: len(str(x).split("; ")) if pd.notna(x) else 0
)

# Filter rows where only one role is present
single_role_df = encoded_stack_overflow_df[encoded_stack_overflow_df['RoleCount'] == 1]

# Step 3: Calculate mean, median, and count for each single job role
single_role_salary = single_role_df.groupby('DeveloperType')['Salary'].agg(['mean', 'median', 'count']).sort_index()

# Display results
display(single_role_salary)


# **Exploring Two Roles**
# - Examined combinations of exactly two roles to understand salary trends and sample sizes for multi-role respondents.
# - Findings: 13 combinations had at least 10 responses, with salaries typically close to the corresponding single-role salaries.

# In[108]:


# Filter for rows with exactly two roles
two_role_rows = encoded_stack_overflow_df[
    encoded_stack_overflow_df['DeveloperType'].str.contains("; ") & 
    (encoded_stack_overflow_df['DeveloperType'].str.count("; ") == 1)
].copy()  # Create a copy to avoid the warning

# Sort roles alphabetically within each row
two_role_rows.loc[:, 'DeveloperTypeSorted'] = two_role_rows['DeveloperType'].apply(
    lambda x: "; ".join(sorted(x.split("; "))) if pd.notna(x) else x
)

# Calculate mean, median, and count for each two-role combination
two_role_stats = two_role_rows.groupby('DeveloperTypeSorted')['Salary'].agg(['mean', 'median', 'count']).sort_values('count', ascending=False)

# Filter for combinations with at least 10 occurrences
filtered_two_role_stats = two_role_stats[two_role_stats['count'] >= 10]
display(filtered_two_role_stats)


# **Exploring Three Roles**
# - Analyzed respondents with exactly three roles to determine salary trends and evaluate representation in the dataset.
# - Findings: Nine combinations had at least 10 responses, with trends similar to those seen in two-role combinations.

# In[109]:


# Filter for rows with exactly three roles
three_role_rows = encoded_stack_overflow_df[
    encoded_stack_overflow_df['DeveloperType'].str.contains("; ") & 
    (encoded_stack_overflow_df['DeveloperType'].str.count("; ") == 2)
].copy()  # Create a copy to avoid the warning

# Sort roles alphabetically within each row
three_role_rows.loc[:, 'DeveloperTypeSorted'] = three_role_rows['DeveloperType'].apply(
    lambda x: "; ".join(sorted(x.split("; "))) if pd.notna(x) else x
)

# Calculate mean, median, and count for each two-role combination
three_role_stats = three_role_rows.groupby('DeveloperTypeSorted')['Salary'].agg(['mean', 'median', 'count']).sort_values('count', ascending=False)

# Filter for combinations with at least 10 occurrences
filtered_three_role_stats = three_role_stats[three_role_stats['count'] >= 10]
display(filtered_three_role_stats)


# **Comparing Multi-Role Salaries to Single Roles**
# - Compared the average salary of respondents with multi-role combinations to the average salary of the highest-paying single role included in the combination.
# - Findings: Multi-role salaries often showed decreases compared to the average salary of the most lucrative single role in the combination.
# - Summary statistics of percentage change: 
#   - Mean: -7.54%
#   - Median: -9.01%
#   - Standard deviation: 27.76%.

# In[110]:


# Filter single-role statistics to exclude rows with NaN mean salary
single_role_salary_filtered = single_role_salary.dropna(subset=['mean'])

# Extract single-role statistics as a dictionary
single_role_stats_dict = single_role_salary_filtered['mean'].to_dict()

# Combine two-role and three-role datasets, excluding rows with NaN mean salary
two_role_stats_filtered = two_role_stats.dropna(subset=['mean'])
three_role_stats_filtered = three_role_stats.dropna(subset=['mean'])

multi_role_stats_combined = pd.concat([two_role_stats_filtered.reset_index(), three_role_stats_filtered.reset_index()])
multi_role_stats_combined = multi_role_stats_combined.rename(columns={"mean": "ActualMeanSalary"})

# Calculate the maximum base salary for each multi-role combination
def calculate_max_single_role_salary(roles):
    role_list = roles.split("; ")
    return max([single_role_stats_dict.get(role, 0) for role in role_list])

multi_role_stats_combined['MaxBaseSalary'] = multi_role_stats_combined['DeveloperTypeSorted'].apply(calculate_max_single_role_salary)

# Exclude rows where MaxBaseSalary is 0 (indicating no matching single-role stats)
multi_role_stats_combined = multi_role_stats_combined[multi_role_stats_combined['MaxBaseSalary'] > 0]

# Calculate the percentage increase
multi_role_stats_combined['PercentageChange'] = (
    (multi_role_stats_combined['ActualMeanSalary'] - multi_role_stats_combined['MaxBaseSalary']) / 
    multi_role_stats_combined['MaxBaseSalary']
) * 100

# Aggregate results
percentage_change_summary = multi_role_stats_combined['PercentageChange'].agg(['mean', 'median', 'std'])

# Display results
print("Percentage Change Summary:")
print(percentage_change_summary)

# Display the combined dataset for review
display(multi_role_stats_combined[['DeveloperTypeSorted', 'ActualMeanSalary', 'MaxBaseSalary', 'PercentageChange']])


# **Verifying Columns Before Encoding**
# - Verified the structure of the dataset, including all relevant columns, before proceeding to encode the *DeveloperType* column.

# In[111]:


display(list(encoded_stack_overflow_df.columns))


# #### Encoding the *DeveloperType* Column
# - Objective: Encode the *DeveloperType* column to represent multi-role combinations effectively in the dataset, ensuring compatibility with machine learning models.
# 
# - Rationale for Chosen Encoding Method:
#    - Ordinal encoding is not suitable due to the lack of a meaningful ranking or hierarchy among roles.
#    - One-hot encoding is an effective method for capturing the presence or absence of each role without implying a ranking.
#    - Encoding approach:
#       - Each role is represented by a binary field (e.g., *Web Developer*, *Mobile Developer*).
#       - A value of 1 indicates that the respondent has listed the role, and 0 indicates otherwise.
#       - For example:
#          - A respondent with the role *Web Developer* will have *Web Developer* = 1 and all other fields = 0.
#          - A respondent with the roles *Web Developer* and *Mobile Developer* will have both *Web Developer* and *Mobile Developer* = 1, with all other fields = 0.

# In[112]:


# Split the DeveloperType column into individual roles
all_roles = (
    encoded_stack_overflow_df['DeveloperType']
    .dropna()  # Exclude NaN values
    .apply(lambda x: x.split("; "))  # Split roles
)

# Flatten the list of all roles and find unique values
unique_roles = set(role for sublist in all_roles for role in sublist)

# Create binary columns for each unique role
for role in unique_roles:
    encoded_stack_overflow_df[role] = encoded_stack_overflow_df['DeveloperType'].apply(
        lambda x: 1 if pd.notna(x) and role in x.split("; ") else 0
    )

# Verify the new columns
print("Columns added for individual roles:")
display(list(unique_roles))

# Check the resulting DataFrame shape to ensure the expected number of columns
print(f"DataFrame shape after encoding DeveloperType: {encoded_stack_overflow_df.shape}")


# In[113]:


# Display a few rows with the original DeveloperType and the new columns
display(encoded_stack_overflow_df[['DeveloperType'] + list(unique_roles)].head(10))


# #### Preparing the Dataset for Machine Learning
# - Objective: Finalize the dataset for machine learning by removing unnecessary columns, retaining only numeric features, and exploring correlations to identify relationships among variables.

# **Dropping the *DeveloperType* Column**
# - The *DeveloperType* column was removed after encoding its values into binary columns representing individual roles.
# - A new DataFrame, *numeric_stack_overflow_df*, was created to retain the updated structure without modifying the dataset used for encoding.

# In[114]:


# Create a new DataFrame with the DeveloperType column removed
numeric_stack_overflow_df = encoded_stack_overflow_df.drop(columns=['DeveloperType'])

# Verify the new DataFrame structure
print("Columns in numeric_stack_overflow_df:")
display(list(numeric_stack_overflow_df.columns))

# Check the shape of the new DataFrame
print(f"Shape of numeric_stack_overflow_df: {numeric_stack_overflow_df.shape}")


# In[115]:


print(numeric_stack_overflow_df.dtypes)


# **Filtering for Numeric Columns**
# - The dataset was filtered to retain only numeric columns using the *select_dtypes* method.
# - This ensures compatibility with machine learning models, which typically require numerical inputs.

# In[116]:


# Filter only numeric columns from the DataFrame
numeric_stack_overflow_df = numeric_stack_overflow_df.select_dtypes(include=['number'])

# Verify the resulting columns and DataFrame shape
print("Columns in numeric_stack_overflow_df after selecting only numeric types:")
display(list(numeric_stack_overflow_df.columns))

print(f"Shape of numeric_stack_overflow_df after selecting numeric types: {numeric_stack_overflow_df.shape}")


# ### Prepare the Data
# 
# #### Split the Dataset into Training and Test Sets
# - Objective: Split the dataset into training and test sets to ensure that the model can be evaluated effectively while maintaining proportional representation of key variables.
# 
# - Methodology:
#     - The dataset was split into training (80%) and test (20%) sets.
#     - Stratification was applied based on the *Gender_Female* column to preserve the proportional representation of men and women in both sets.
# 
# Justification for Stratified Sampling:
#     - Gender is a critical variable in this analysis, as the project investigates salary disparities between men and women in the tech industry.
#     - The dataset contains significantly more men than women. Without stratification, women may be underrepresented in the test set, leading to biased model evaluation.
#     - Stratification ensures that the gender distribution in both the training and test sets reflects the original dataset.

# In[181]:


# Define features and target
X = numeric_stack_overflow_df.drop(columns=['Salary'])  # Features
y = numeric_stack_overflow_df['Salary']                # Target variable

# Perform train/test split, stratified by Gender_Female
X_train, X_test, y_train, y_test = train_test_split(
    X, y, test_size=0.2, random_state=42, stratify=X['Gender_Female']
)

# Verify distributions
print("Training set Gender distribution:")
print(X_train['Gender_Female'].value_counts(normalize=True))
print("\nTest set Gender distribution:")
print(X_test['Gender_Female'].value_counts(normalize=True))


# #### Validating the Stratified Train/Test Split
# - Objective: Verify that the stratified split preserved the gender distribution from the original dataset.
# - Methodology:
#   - Compared gender proportions in the overall dataset, stratified test set, and a random sample for evaluation.
#   - Calculated percentage errors to assess differences between the splits.
# - Results:
#   - The stratified split perfectly maintained the original gender proportions (0% error).
#   - The random split introduced a 15% overrepresentation of women and a 1.91% underrepresentation of men, confirming the need for stratification.

# In[182]:


# Function to calculate gender proportions
def gender_proportions(data):
    return data['Gender_Female'].value_counts(normalize=True)

# Compare proportions
compare_props = pd.DataFrame({
    "Overall": gender_proportions(numeric_stack_overflow_df),
    "Stratified": gender_proportions(X_test),
    "Random": gender_proportions(X.sample(frac=0.2, random_state=42))  # Random sample for comparison
}).sort_index()

# Calculate percentage errors
compare_props["Rand. %error"] = 100 * (compare_props["Random"] / compare_props["Overall"] - 1)
compare_props["Strat. %error"] = 100 * (compare_props["Stratified"] / compare_props["Overall"] - 1)

print(compare_props)


# #### Exploring Feature Correlations
# To better understand the relationships between numerical features and the target variable (*Salary*), several analyses were conducted to guide feature selection for scaling, normalization, and inclusion in the regression model:
# 
# 1. **Correlation Matrix**:
#    - A correlation matrix was generated to quantify the relationships between features and *Salary*.
#    - Key Observations:
#      - *Salary* shows moderate positive correlations with:
#        - *YearsCodedJobNumeric* (0.517).
#        - *CompanyTypeEncoded* (0.312).
#        - *EducationCategoryNumeric* (0.237).
#      - Weak or negligible correlations were observed between *Salary* and most role-based features (e.g., *Web Developer*: -0.102, *Mobile Developer*: 0.030).
#      - *Gender_Female* displayed a weak negative correlation with *Salary* (-0.066), indicating a slight inverse relationship.
# 
# 2. **Heatmap Visualization**:
#    - A heatmap of the correlation matrix was created to provide a visual summary of the relationships between features.
#    - This visualization highlighted moderate correlations between *Salary* and select features, as well as inter-feature correlations that may require consideration during modeling.
# 
# 3. **Scatter Matrix**:
#    - A scatter matrix was generated to visualize pairwise relationships between *Salary* and key features, including:
#      - *YearsCodedJobNumeric*
#      - *EducationCategoryNumeric*
#      - *CompanyTypeEncoded*
#      - *CompanySizeEncoded*
#      - *HomeRemoteEncoded*
#    - Observations:
#      - The scatter plots revealed trends between *Salary* and *YearsCodedJobNumeric*, with a general positive relationship.
#      - Histograms on the diagonal provided insights into feature distributions, such as the skewed distribution of *Salary*.
#      - While the scatter matrix offers a broad overview, targeted scatter plots may be more effective for interpreting specific relationships.

# In[183]:


stack_overflow_corr = numeric_stack_overflow_df.corr()

stack_overflow_corr['Salary'].sort_values(ascending=False)


# In[184]:


sns.heatmap(stack_overflow_corr, annot=False, cmap="coolwarm")

plt.show()


# In[185]:


# Select key numerical features for the scatter matrix
selected_features = [
    'Salary', 
    'YearsCodedJobNumeric', 
    'EducationCategoryNumeric', 
    'CompanyTypeEncoded', 
    'CompanySizeEncoded',
    'HomeRemoteEncoded'
]

# Generate the scatter matrix
scatter_matrix(numeric_stack_overflow_df[selected_features], figsize=(20, 16), alpha=0.2, diagonal='hist')
plt.show()


# **Remove Rows with Missing Key Values**
# - Rows where both *YearsCodedJobNumeric* and *Salary* were missing were removed to ensure that the dataset only includes records with sufficient information for meaningful analysis and salary prediction. This step minimizes noise and preserves data integrity.

# In[186]:


# Remove rows where both YearsCodedJobNumeric and Salary are missing
cleaned_stack_overflow_df = numeric_stack_overflow_df.dropna(subset=['YearsCodedJobNumeric', 'Salary'], how='all').copy()

# Verify the shape of the cleaned dataset
print(f"Shape of dataset after removing rows with both YearsCodedJobNumeric and Salary missing: {cleaned_stack_overflow_df.shape}")

# Count rows removed
rows_removed = numeric_stack_overflow_df.shape[0] - cleaned_stack_overflow_df.shape[0]
print(f"Number of rows removed: {rows_removed}")


# #### Investigating Developer Roles and Their Impact on Salary
# The relationship between developer roles and salary was explored to determine whether roles should be included as features in the salary prediction model. Specifically, this analysis evaluated whether combining roles with years of coding experience (*YearsCodedJobNumeric*) provided meaningful insights or predictive power.
# 
# **Why This Was Done**:
# - Developer roles might influence salary, but the relationship is not straightforward due to respondents often listing multiple roles.
# - Combining developer roles with *YearsCodedJobNumeric* could highlight salary trends that are specific to particular roles and levels of experience.
# - The analysis was conducted to determine whether role-based features add enough value to justify their inclusion in the final dataset.
# 
# **Steps Taken**:
# 1. Interaction terms were generated by combining developer roles with *YearsCodedJobNumeric* to assess whether the combination strengthened their correlation with salary.
# 2. A baseline analysis was performed using *YearsCodedJobNumeric* alone to compare its predictive power against the interaction model.
# 
# **Key Findings**:
# - *YearsCodedJobNumeric* as a standalone feature explained 26.78% of the variance in salary (R² = 0.2678).
# - Adding interaction terms for developer roles increased the R² to 29.63%, representing a minimal improvement of 2.85%.
# - The small improvement suggests that roles, even when combined with years of coding, contribute limited additional predictive value relative to experience alone.

# In[187]:


# Create interaction terms between developer roles and YearsCodedJobNumeric
role_columns = [
    'Web developer', 'Mobile developer', 'DevOps specialist', 
    'Database administrator', 'Embedded applications/devices developer', 
    'Quality assurance engineer', 'Desktop applications developer', 
    'Machine learning specialist', 'Developer with a statistics or mathematics background',
    'Other', 'Graphic designer', 'Systems administrator', 'Data scientist', 'Graphics programming'
]

# Create a temporary dataset meeting the specified criteria
temp_df = cleaned_stack_overflow_df[
    (cleaned_stack_overflow_df[role_columns].sum(axis=1) > 0) &  # At least one role is present
    cleaned_stack_overflow_df['Salary'].notna() &               # Salary is not NaN
    cleaned_stack_overflow_df['YearsCodedJobNumeric'].notna()   # YearsCodedJobNumeric is not NaN
].copy()

# Verify the shape of the temporary dataset
print(f"Shape of temporary dataset: {temp_df.shape}")


# In[188]:


# Create interaction terms between developer roles and YearsCodedJobNumeric
role_columns = [
    'Web developer', 'Mobile developer', 'DevOps specialist', 
    'Database administrator', 'Embedded applications/devices developer', 
    'Quality assurance engineer', 'Desktop applications developer', 
    'Machine learning specialist', 'Developer with a statistics or mathematics background',
    'Other', 'Graphic designer', 'Systems administrator', 'Data scientist', 'Graphics programming'
]

# Generate interaction terms
for role in role_columns:
    temp_df[f'{role}_YearsCoded'] = (
        temp_df[role] * temp_df['YearsCodedJobNumeric']
    )

# Calculate correlations between interaction terms and Salary
interaction_corr = temp_df.corr()['Salary'].sort_values(ascending=False)

# Display correlations for interaction terms
print(interaction_corr[interaction_corr.index.str.contains('_YearsCoded')])


# In[189]:


# Define the target variable
y = temp_df['Salary']

# Baseline Model: YearsCodedJobNumeric only
X_baseline = temp_df[['YearsCodedJobNumeric']]
baseline_model = LinearRegression()
baseline_model.fit(X_baseline, y)
baseline_r2 = r2_score(y, baseline_model.predict(X_baseline))
print(f"Baseline Model R² (YearsCodedJobNumeric only): {baseline_r2:.4f}")

# Interaction Model: YearsCodedJobNumeric + Role Interaction Terms
role_interaction_columns = [f'{role}_YearsCoded' for role in role_columns]
X_interaction = temp_df[['YearsCodedJobNumeric'] + role_interaction_columns]
interaction_model = LinearRegression()
interaction_model.fit(X_interaction, y)
interaction_r2 = r2_score(y, interaction_model.predict(X_interaction))
print(f"Interaction Model R² (YearsCodedJobNumeric + Role Interactions): {interaction_r2:.4f}")

# Compare R² values
r2_difference = interaction_r2 - baseline_r2
print(f"Increase in R² from adding role interactions: {r2_difference:.4f}")


# ### Feature Reduction: Dropping Developer Role Features
# - Objective: Simplify the dataset by removing developer role features that added complexity without significantly improving the model's performance.
# 
# - Method:
#   - A pipeline was used to drop all developer role columns, including interaction terms between roles and *YearsCodedJobNumeric*.
#   - This ensures a consistent and scalable approach for both the training and test datasets.

# In[190]:


# Define the columns to drop
drop_columns = [
    'Web developer', 'Mobile developer', 'DevOps specialist',
    'Database administrator', 'Embedded applications/devices developer',
    'Quality assurance engineer', 'Desktop applications developer',
    'Machine learning specialist', 'Developer with a statistics or mathematics background',
    'Other', 'Graphic designer', 'Systems administrator', 'Data scientist', 'Graphics programming',
    'Web developer_YearsCoded', 'Mobile developer_YearsCoded', 'DevOps specialist_YearsCoded',
    'Database administrator_YearsCoded', 'Embedded applications/devices developer_YearsCoded',
    'Quality assurance engineer_YearsCoded', 'Desktop applications developer_YearsCoded',
    'Machine learning specialist_YearsCoded', 'Developer with a statistics or mathematics background_YearsCoded',
    'Other_YearsCoded', 'Graphic designer_YearsCoded', 'Systems administrator_YearsCoded',
    'Data scientist_YearsCoded', 'Graphics programming_YearsCoded', 'RoleCount'
]

# Dynamically filter the columns to drop based on the current DataFrame columns
valid_drop_columns = [col for col in drop_columns if col in X_train.columns]

# Define a ColumnTransformer to drop the columns
drop_transformer = ColumnTransformer(
    transformers=[
        ('drop_columns', 'drop', valid_drop_columns)  # Drop only valid columns
    ],
    remainder='passthrough'  # Pass through remaining columns
)

# Create the pipeline
drop_pipeline = Pipeline([
    ('column_transform', drop_transformer)  # Use the ColumnTransformer
])

# Apply the pipeline to X_train and X_test
X_train_transformed = drop_pipeline.fit_transform(X_train)
X_test_transformed = drop_pipeline.transform(X_test)

# Extract the remaining column names
remaining_columns = [col for col in X_train.columns if col not in drop_columns]

# Convert back to DataFrame, matching transformed data to original indices and remaining columns
X_train_transformed = pd.DataFrame(X_train_transformed, columns=remaining_columns, index=X_train.index)
X_test_transformed = pd.DataFrame(X_test_transformed, columns=remaining_columns, index=X_test.index)

# Verify the results
print(f"Shape of training set after pipeline transformation: {X_train_transformed.shape}")
print(f"Shape of test set after pipeline transformation: {X_test_transformed.shape}")
display(list(X_train_transformed.columns))


# **Check for features with missing values**

# In[191]:


# Analyze missing values in the transformed training dataset
train_feature_summary = X_train_transformed.isna().sum().reset_index()
train_feature_summary.columns = ['Feature', 'NaN_Count']
train_feature_summary['Non-NaN_Count'] = X_train_transformed.shape[0] - train_feature_summary['NaN_Count']
train_feature_summary['Total_Count'] = X_train_transformed.shape[0]
train_feature_summary['NaN_Percentage'] = (train_feature_summary['NaN_Count'] / train_feature_summary['Total_Count']) * 100

# Analyze missing values in the transformed testing dataset
test_feature_summary = X_test_transformed.isna().sum().reset_index()
test_feature_summary.columns = ['Feature', 'NaN_Count']
test_feature_summary['Non-NaN_Count'] = X_test_transformed.shape[0] - test_feature_summary['NaN_Count']
test_feature_summary['Total_Count'] = X_test_transformed.shape[0]
test_feature_summary['NaN_Percentage'] = (test_feature_summary['NaN_Count'] / test_feature_summary['Total_Count']) * 100

# Sort by the number of NaN values for clarity
train_feature_summary = train_feature_summary.sort_values(by='NaN_Count', ascending=False).reset_index(drop=True)
test_feature_summary = test_feature_summary.sort_values(by='NaN_Count', ascending=False).reset_index(drop=True)

# Display the results
print("Training Feature Summary:")
print(train_feature_summary)

print("\nTesting Feature Summary:")
print(test_feature_summary)


# In[ ]:


# Define features requiring imputation
num_features = ['YearsCodedJobNumeric']

cat_features = ['EducationCategoryNumeric', 'Gender_Female', 'Gender_Male', 'HomeRemoteEncoded', 'CompanySizeEncoded', 'CompanyTypeEncoded']

# Define the pipeline for numerical features
num_pipeline = Pipeline([
    ('imputer', SimpleImputer(strategy='median'))  # Impute missing values with median
])

# Define the pipeline for categorical features
cat_pipeline = Pipeline([
  ('imputer', SimpleImputer(strategy='most_frequent')), # This would impute missing values with the most frequent, but there are no missing values
])

# Combine the numerical and categorica pipelines with passthrough for all remaining columns
full_pipeline = ColumnTransformer([
    ('num', num_pipeline, num_features),  # Apply numerical pipeline to numerical features
    ('cat', cat_pipeline, cat_features)
]) 

# Apply the transformation to the pre-transformed datasets
X_train_transformed = full_pipeline.fit_transform(X_train)
X_test_transformed = full_pipeline.transform(X_test)

# Retrieve updated column names
transformed_columns = full_pipeline.get_feature_names_out()

# Transform the numpy n-dimensional array into a pandas dataframe
X_train_transformed = pd.DataFrame(X_train_transformed, columns=transformed_columns, index=X_train.index)
X_test_transformed = pd.DataFrame(X_test_transformed, columns=transformed_columns, index=X_test.index)

# Verify the transformed datasets
print("Transformed Training Set:")
display(X_train_transformed.head())

print("\nTransformed Testing Set:")
display(X_test_transformed.head())


# In[193]:


full_pipeline.get_feature_names_out()


# In[194]:


column_names = [ 
  feature.replace('num__', '').replace('cat__', '') 
  for feature in full_pipeline.get_feature_names_out()
]
column_names


# In[195]:


X_train_transformed.columns = column_names
X_test_transformed.columns = column_names

# Verify the updated column names
print("Updated column names in the training set:")
print(X_train_transformed.columns)

print("\nUpdated column names in the testing set:")
print(X_test_transformed.columns)


# In[196]:


print("Missing values in transformed training set:")
display(X_train_transformed.isnull().sum())

print("\nMissing values in transformed testing set:")
display(X_test_transformed.isnull().sum())


# ### Standardization of Numerical Features
# - To ensure that all numerical features are on the same scale, the *YearsCodedJobNumeric* column was standardized using *StandardScaler*. Standardization is necessary because it prevents features with larger scales from disproportionately influencing the model.
# 
# - Encoding pipelines were also established for nominal and ordinal categorical features (*Gender_Female*, *Gender_Male*, *HomeRemoteEncoded*, *CompanySizeEncoded*, *CompanyTypeEncoded*, and *EducationCategoryNumeric*). However, these pipelines were not applied because these features did not contain any missing values requiring imputation.
# 
# - Verification was performed by checking the mean and standard deviation of the standardized column:
#     - *Mean:* 0.0000000000  
#     - *Standard Deviation:* 1.0000000000  
# 
# - These results confirm that the feature was scaled appropriately for use in machine learning models.

# In[197]:


# Define feature groups
num_features = ['YearsCodedJobNumeric']
ordinal_features = ['HomeRemoteEncoded', 'CompanySizeEncoded', 'CompanyTypeEncoded', 'EducationCategoryNumeric']
nominal_features = ['Gender_Female', 'Gender_Male']

# Numerical pipeline: Imputation + Scaling
num_pipeline = Pipeline([
    ('scaler', StandardScaler())  # Scaling only since no missing values remain
])

# Ordinal pipeline: Ordinal Encoding
ordinal_pipeline = Pipeline([
    ('encoder', OrdinalEncoder())
])

# Nominal pipeline: One-Hot Encoding
nominal_pipeline = Pipeline([
    ('encoder', OneHotEncoder(drop='first'))  # Drop one dummy variable to avoid redundancy
])

# Combine all pipelines
full_pipeline = ColumnTransformer([
    ('num', num_pipeline, num_features),
    ('ordinal', ordinal_pipeline, ordinal_features),
    ('nominal', nominal_pipeline, nominal_features)
], remainder='passthrough')  # Ensure other columns are passed through unchanged

# Apply the pipeline to pre-transformed datasets
X_train_transformed = full_pipeline.fit_transform(X_train_transformed)
X_test_transformed = full_pipeline.transform(X_test_transformed)

# Get updated column names after transformation
transformed_columns = full_pipeline.get_feature_names_out()

# Clean up the column names for consistency
cleaned_column_names = [
    col.replace('num__', '').replace('ordinal__', '').replace('nominal__', '')
    for col in transformed_columns
]

# Convert back to DataFrame with cleaned names
X_train_transformed = pd.DataFrame(X_train_transformed, columns=cleaned_column_names, index=X_train.index)
X_test_transformed = pd.DataFrame(X_test_transformed, columns=cleaned_column_names, index=X_test.index)

# Verify the transformed datasets
print("Transformed Training Set:")
display(X_train_transformed.head())

print("\nTransformed Testing Set:")
display(X_test_transformed.head())


# In[198]:


# Get the mean and standard deviation of the transformed dataset
mean = np.mean(X_train_transformed['YearsCodedJobNumeric'])
std_dev = np.std(X_train_transformed['YearsCodedJobNumeric'])

# Print the results with formatting
print(f"Mean: {mean:.10f}")
print(f"Standard Deviation: {std_dev:.10f}")


# ### Handling Missing Salary Values
# 
# #### Rationale
# Rather than imputing missing values in the *Salary* column with static measures such as the median or mean, I aim to build a predictive model to estimate these values. This approach ensures that relationships between features (e.g., *YearsCodedJobNumeric*, *EducationCategoryNumeric*) and the *Salary* target are preserved. By predicting missing values based on these relationships, the imputation process becomes both realistic and data-driven.
# 
# #### Benefits of This Approach
# 1. **Preserving Relationships**: Static imputations can ignore interactions between variables, potentially introducing bias. A predictive model captures these interactions, leading to more accurate imputations.
# 2. **Alignment with Assignment Goals**: This approach tests and evaluates multiple models, a requirement of the machine learning portion of the assignment, ensuring the best algorithm is selected for imputation and future predictions.
# 3. **Avoiding Arbitrary Assumptions**: Static imputations assume uniformity in the data (e.g., all respondents with missing values have a similar salary), which is often unrealistic.
# 
# #### Process
# 1. **Separate Data**: Rows with missing *Salary* values are separated from those with non-missing values.
# 2. **Train Models**: Models are trained on the non-missing data to predict *Salary*.
# 3. **Evaluate Models**: Multiple algorithms are tested to find the best-performing one based on metrics such as RMSE or MAE.
# 4. **Predict Missing Values**: The best model can be used to estimate missing *Salary* values.
# 
# This approach ensures that the imputation process is both methodologically sound and aligned with the machine learning section's overall objective of accurately predicting salaries.

# In[199]:


display(list(cleaned_stack_overflow_df.columns))


# In[200]:


column_drop = [
    'Web developer', 'Mobile developer', 'DevOps specialist',
    'Database administrator', 'Embedded applications/devices developer',
    'Quality assurance engineer', 'Desktop applications developer',
    'Machine learning specialist', 'Developer with a statistics or mathematics background',
    'Other', 'Graphic designer', 'Systems administrator', 'Data scientist', 'Graphics programming',
    'Web developer_YearsCoded', 'Mobile developer_YearsCoded', 'DevOps specialist_YearsCoded',
    'Database administrator_YearsCoded', 'Embedded applications/devices developer_YearsCoded',
    'Quality assurance engineer_YearsCoded', 'Desktop applications developer_YearsCoded',
    'Machine learning specialist_YearsCoded', 'Developer with a statistics or mathematics background_YearsCoded',
    'Other_YearsCoded', 'Graphic designer_YearsCoded', 'Systems administrator_YearsCoded',
    'Data scientist_YearsCoded', 'Graphics programming_YearsCoded', 'RoleCount'
]

revised_stack_overflow_df = cleaned_stack_overflow_df.drop(columns=column_drop, errors='ignore')

display(list(revised_stack_overflow_df.columns))


# In[201]:


# Separate rows with missing and non-missing Salary values
non_missing_data = revised_stack_overflow_df.dropna(subset=['Salary'])
missing_data = revised_stack_overflow_df[cleaned_stack_overflow_df['Salary'].isna()]

# Define features and target for non-missing data
X_imputation = non_missing_data.drop(columns=['Salary'])
y_imputation = non_missing_data['Salary']

X_train_imputation, X_test_imputation, y_train_imputation, y_test_imputation = train_test_split(
    X_imputation, y_imputation, 
    test_size=0.2, 
    random_state=42, 
    stratify=non_missing_data['Gender_Female']  # Maintain balance across gender
)

# Verify the distribution of the new training and testing sets
print("Training set shape:", X_train_imputation.shape)
print("Testing set shape:", X_test_imputation.shape)


# In[202]:


# Verify the gender distribution in training and testing sets
print("Training set Gender distribution:")
print(y_train_imputation.groupby(X_train_imputation['Gender_Female']).count())

print("\nTesting set Gender distribution:")
print(y_test_imputation.groupby(X_test_imputation['Gender_Female']).count())


# **Check for missing feature values**

# In[203]:


# Analyze missing values in the transformed training dataset
train_feature_summary = X_train_imputation.isna().sum().reset_index()
train_feature_summary.columns = ['Feature', 'NaN_Count']
train_feature_summary['Non-NaN_Count'] = X_train_imputation.shape[0] - train_feature_summary['NaN_Count']
train_feature_summary['Total_Count'] = X_train_imputation.shape[0]
train_feature_summary['NaN_Percentage'] = (train_feature_summary['NaN_Count'] / train_feature_summary['Total_Count']) * 100

# Analyze missing values in the transformed testing dataset
test_feature_summary = X_test_imputation.isna().sum().reset_index()
test_feature_summary.columns = ['Feature', 'NaN_Count']
test_feature_summary['Non-NaN_Count'] = X_test_imputation.shape[0] - test_feature_summary['NaN_Count']
test_feature_summary['Total_Count'] = X_test_imputation.shape[0]
test_feature_summary['NaN_Percentage'] = (test_feature_summary['NaN_Count'] / test_feature_summary['Total_Count']) * 100

# Sort by the number of NaN values for clarity
train_feature_summary = train_feature_summary.sort_values(by='NaN_Count', ascending=False).reset_index(drop=True)
test_feature_summary = test_feature_summary.sort_values(by='NaN_Count', ascending=False).reset_index(drop=True)

# Display the results
print("Training Feature Summary:")
print(train_feature_summary)

print("\nTesting Feature Summary:")
print(test_feature_summary)


# In[204]:


# Define features requiring imputation
num_features = ['YearsCodedJobNumeric']

cat_features = ['EducationCategoryNumeric', 'Gender_Female', 'Gender_Male', 'HomeRemoteEncoded', 'CompanySizeEncoded', 'CompanyTypeEncoded']

# Define the pipeline for numerical features
num_pipeline = Pipeline([
    ('imputer', SimpleImputer(strategy='median'))  # Impute missing values with median
])

# Define the pipeline for categorical features
cat_pipeline = Pipeline([
  ('imputer', SimpleImputer(strategy='most_frequent')), # This would impute missing values with the most frequent, but there are no missing values
])

# Combine the numerical and categorica pipelines with passthrough for all remaining columns
full_pipeline = ColumnTransformer([
    ('num', num_pipeline, num_features),  # Apply numerical pipeline to numerical features
    ('cat', cat_pipeline, cat_features)
]) 

# Apply the transformation to the imputation training and testing sets
X_train_imputed_transformed = full_pipeline.fit_transform(X_train_imputation)
X_test_imputed_transformed = full_pipeline.transform(X_test_imputation)

# Retrieve updated column names
imputed_transformed_columns = full_pipeline.get_feature_names_out()

# Transform the numpy n-dimensional array into a pandas dataframe
X_train_imputed_transformed = pd.DataFrame(X_train_imputed_transformed, columns=imputed_transformed_columns, index=X_train_imputation.index)
X_test_imputed_transformed = pd.DataFrame(X_test_imputed_transformed, columns=imputed_transformed_columns, index=X_test_imputation.index)

# Verify the transformed datasets
print("Transformed Training Set (Imputed):")
print(X_train_imputed_transformed.head())

print("\nTransformed Testing Set (Imputed):")
print(X_test_imputed_transformed.head())


# In[205]:


full_pipeline.get_feature_names_out()


# In[206]:


imputed_column_names = [ 
  feature.replace('num__', '').replace('cat__', '') 
  for feature in full_pipeline.get_feature_names_out()
]
imputed_column_names


# In[207]:


X_train_imputed_transformed.columns = imputed_column_names
X_test_imputed_transformed.columns = imputed_column_names

# Verify the updated column names
print("Updated column names in the training set:")
print(X_train_imputed_transformed.columns)

print("\nUpdated column names in the testing set:")
print(X_test_imputed_transformed.columns)


# In[208]:


print("Missing values in transformed training set:")
display(X_train_imputed_transformed.isnull().sum())

print("\nMissing values in transformed testing set:")
display(X_test_imputed_transformed.isnull().sum())


# ### Standardization of Numerical Features

# In[209]:


# Define feature groups
num_features = ['YearsCodedJobNumeric']
ordinal_features = ['HomeRemoteEncoded', 'CompanySizeEncoded', 'CompanyTypeEncoded', 'EducationCategoryNumeric']
nominal_features = ['Gender_Female', 'Gender_Male']

# Numerical pipeline: Imputation + Scaling
num_pipeline = Pipeline([
    ('scaler', StandardScaler())
])

# Ordinal pipeline: Ordinal Encoding
ordinal_pipeline = Pipeline([
    ('encoder', OrdinalEncoder())
])

# Nominal pipeline: One-Hot Encoding
nominal_pipeline = Pipeline([
    ('encoder', OneHotEncoder(drop='first'))  # Drop one dummy variable to avoid redundancy
])

# Combine all pipelines
full_pipeline = ColumnTransformer([
    ('num', num_pipeline, num_features),
    ('ordinal', ordinal_pipeline, ordinal_features),
    ('nominal', nominal_pipeline, nominal_features)
], remainder='passthrough')  # Ensure other columns are passed through unchanged

# Apply the pipeline to pre-transformed datasets
X_train_imputed_transformed = full_pipeline.fit_transform(X_train_imputed_transformed)
X_test_imputed_transformed = full_pipeline.transform(X_test_imputed_transformed)

# Get updated column names after transformation
imputed_transformed_columns = full_pipeline.get_feature_names_out()

# Clean up the column names for consistency
cleaned_column_names = [
    col.replace('num__', '').replace('ordinal__', '').replace('nominal__', '')
    for col in imputed_transformed_columns
]

# Convert back to DataFrame with cleaned names
X_train_imputed_transformed = pd.DataFrame(X_train_imputed_transformed, columns=cleaned_column_names, index=X_train_imputation.index)
X_test_imputed_transformed = pd.DataFrame(X_test_imputed_transformed, columns=cleaned_column_names, index=X_test_imputation.index)

# Verify the transformed datasets
print("Transformed Training Set:")
display(X_train_imputed_transformed.head())

print("\nTransformed Testing Set:")
display(X_test_imputed_transformed.head())


# In[210]:


# Get the mean and standard deviation of the transformed dataset
mean = np.mean(X_train_imputed_transformed['YearsCodedJobNumeric'])
std_dev = np.std(X_train_imputed_transformed['YearsCodedJobNumeric'])

# Print the results with formatting
print(f"Mean: {mean:.10f}")
print(f"Standard Deviation: {std_dev:.10f}")


# ### Analyze the Data
# 
# #### Implement and Evaluate Models
# - Objective: Begin the process of analyzing the dataset by implementing multiple models to predict *Salary* and evaluating their performance.
#   
# - Methodology:
#     - A baseline model using the mean of the training target variable (*Salary*) has been implemented to establish a performance benchmark.
#     - Advanced models will be tested and evaluated next to compare their effectiveness against the baseline.
# 
# - Justification for Baseline Model:
#     - Provides a simple benchmark to assess whether more complex models significantly improve predictive performance.
#     - Allows for comparison of RMSE, MAE, and R² metrics to evaluate the added value of advanced modeling techniques.
# 
# - Next Steps:
#     - Test multiple algorithms, including linear regression, tree-based models, and regularized regression techniques.
#     - Compare model performance on both the training and testing sets using appropriate metrics.

# #### Baseline Model (Mean Prediction)
# The Baseline Model predicts the mean *Salary* from the training set for all observations. 
# 
# Model Fit:
# - The model performs similarly on the training and testing sets, with an RMSE of ~33,400 on the training set and ~34,000 on the testing set. 
# - The near-zero \( R^2 \) on both sets indicates the model explains no variance in the data, consistent with its simplicity.
# 
# Purpose:
# - This model serves as a reference point. Any improvement in RMSE, MAE, or \( R^2 \) by subsequent models reflects their ability to capture relationships between features and the target variable.

# In[212]:


# Calculate baseline predictions using the mean of the training target
y_train_baseline_pred = [y_train_imputation.mean()] * len(y_train_imputation)
y_test_baseline_pred = [y_train_imputation.mean()] * len(y_test_imputation)  # Use train mean for test predictions

# Evaluate the baseline model using RMSE
baseline_rmse_train = np.sqrt(mean_squared_error(y_train_imputation, y_train_baseline_pred))
baseline_rmse_test = np.sqrt(mean_squared_error(y_test_imputation, y_test_baseline_pred))

# Evaluate the baseline model using MAE
baseline_mae_train = mean_absolute_error(y_train_imputation, y_train_baseline_pred)
baseline_mae_test = mean_absolute_error(y_test_imputation, y_test_baseline_pred)

# Calculate R² for baseline model
baseline_r2_train = r2_score(y_train_imputation, y_train_baseline_pred)
baseline_r2_test = r2_score(y_test_imputation, y_test_baseline_pred)

# Display the results
print(f"Baseline Model (Mean Prediction) - RMSE on Training Set: {baseline_rmse_train:.2f}")
print(f"Baseline Model (Mean Prediction) - RMSE on Testing Set: {baseline_rmse_test:.2f}")
print(f"Baseline Model (Mean Prediction) - MAE on Training Set: {baseline_mae_train:.2f}")
print(f"Baseline Model (Mean Prediction) - MAE on Testing Set: {baseline_mae_test:.2f}")
print(f"Baseline Model (Mean Prediction) - R² on Training Set: {baseline_r2_train:.4f}")
print(f"Baseline Model (Mean Prediction) - R² on Testing Set: {baseline_r2_test:.4f}")


# #### Linear Regression Model
# The Linear Regression Model assumes linear relationships between the features and the *Salary*.
# 
# Model Fit:
# - The Linear Regression Model performs slightly better on the training set than on the testing set. While the RMSE and MAE are lower on the training set, the testing set results are reasonably close, suggesting that the model is not severely overfit but could be slightly underfit, as it does not perfectly generalize to unseen data.
# 
# Comparison with the Baseline:
# - The RMSE decreases from ~34,000 (Baseline) to ~28,000 on the testing set, indicating that the Linear Regression Model provides more accurate predictions compared to predicting the mean.
# - The MAE drops significantly, from ~27,900 (Baseline) to ~21,600 on the testing set, showing that the model predicts salaries closer to their true values on average.
# - The \( R^2 \) improves from ~0.0 (Baseline) to 0.3208, meaning that the model explains approximately 32% of the variance in *Salary*, a significant improvement over the Baseline, which explains none of the variance.
# 
# Insights:
# - A higher \( R^2 \) value shows the Linear Regression Model captures meaningful relationships between features and the target variable, *Salary*.
# - The reduction in RMSE and MAE demonstrates that the model consistently provides more accurate predictions than the Baseline Model.

# In[213]:


# Initialize and fit the linear regression model
lin_reg = LinearRegression()
lin_reg.fit(X_train_imputed_transformed, y_train_imputation)

# Predict on the training and test sets
y_train_pred = lin_reg.predict(X_train_imputed_transformed)
y_test_pred = lin_reg.predict(X_test_imputed_transformed)

# Evaluate the model
lin_rmse_train = np.sqrt(mean_squared_error(y_train_imputation, y_train_pred))
lin_rmse_test = np.sqrt(mean_squared_error(y_test_imputation, y_test_pred))

lin_mae_train = mean_absolute_error(y_train_imputation, y_train_pred)
lin_mae_test = mean_absolute_error(y_test_imputation, y_test_pred)

lin_r2_train = r2_score(y_train_imputation, y_train_pred)
lin_r2_test = r2_score(y_test_imputation, y_test_pred)

# Display results
print(f"Linear Regression Model - RMSE on Training Set: {lin_rmse_train:.2f}")
print(f"Linear Regression Model - RMSE on Testing Set: {lin_rmse_test:.2f}")
print(f"Linear Regression Model - MAE on Training Set: {lin_mae_train:.2f}")
print(f"Linear Regression Model - MAE on Testing Set: {lin_mae_test:.2f}")
print(f"Linear Regression Model - R² on Training Set: {lin_r2_train:.4f}")
print(f"Linear Regression Model - R² on Testing Set: {lin_r2_test:.4f}")


# #### Polynomial Regression Model (Degree 2)
# The Polynomial Regression Model (Degree 2) introduces non-linear relationships between the features and the *Salary*, allowing it to capture more complex patterns in the data compared to Linear Regression.
# 
# Model Fit:
# - The model performs better on the training set (RMSE = ~25,600, MAE = ~19,700) than on the testing set (RMSE = ~27,500, MAE = ~21,200). This suggests the model generalizes reasonably well but may slightly overfit the training data.
# 
# Comparison with the Baseline:
# - RMSE improves from ~34,000 (Baseline) to ~27,500 on the testing set, a notable reduction in prediction error.
# - MAE decreases from ~27,900 (Baseline) to ~21,200, showing the model predicts salaries closer to their true values on average.
# - \( R^2 \) improves from ~0.0 (Baseline) to 0.3487, indicating the model explains approximately 35% of the variance in *Salary*, compared to the Baseline’s complete lack of explanatory power.
# 
# Insights:
# - The higher \( R^2 \) demonstrates that the model captures non-linear relationships between features and *Salary* that Linear Regression may miss.
# - The reduction in RMSE and MAE confirms that the model makes more accurate predictions than both the Baseline and Linear Regression models.
# - The relatively small gap between training and testing performance suggests that the model balances complexity with generalizability.

# In[215]:


# Define the polynomial regression pipeline
poly_reg = Pipeline([
    ('poly_features', PolynomialFeatures(degree=2, include_bias=False)),  # Generate polynomial features
    ('lin_reg', LinearRegression())  # Apply linear regression on the transformed features
])

# Train the model
poly_reg.fit(X_train_imputed_transformed, y_train_imputation)

# Predict on training and testing sets
y_train_poly_pred = poly_reg.predict(X_train_imputed_transformed)
y_test_poly_pred = poly_reg.predict(X_test_imputed_transformed)

# Evaluate the Polynomial Regression model
poly_rmse_train = np.sqrt(mean_squared_error(y_train_imputation, y_train_poly_pred))
poly_rmse_test = np.sqrt(mean_squared_error(y_test_imputation, y_test_poly_pred))
poly_mae_train = mean_absolute_error(y_train_imputation, y_train_poly_pred)
poly_mae_test = mean_absolute_error(y_test_imputation, y_test_poly_pred)
poly_r2_train = r2_score(y_train_imputation, y_train_poly_pred)
poly_r2_test = r2_score(y_test_imputation, y_test_poly_pred)

# Display the results
print(f"Polynomial Regression Model (Degree 2) - RMSE on Training Set: {poly_rmse_train:.2f}")
print(f"Polynomial Regression Model (Degree 2) - RMSE on Testing Set: {poly_rmse_test:.2f}")
print(f"Polynomial Regression Model (Degree 2) - MAE on Training Set: {poly_mae_train:.2f}")
print(f"Polynomial Regression Model (Degree 2) - MAE on Testing Set: {poly_mae_test:.2f}")
print(f"Polynomial Regression Model (Degree 2) - R² on Training Set: {poly_r2_train:.4f}")
print(f"Polynomial Regression Model (Degree 2) - R² on Testing Set: {poly_r2_test:.4f}")


# #### Polynomial Regression Model (Degree 3)
# The Polynomial Regression Model (Degree 3) builds on the previous models by capturing more detailed patterns between the features and *Salary*.
# 
# Model Fit:
# - The model performs well on both the training set (RMSE = ~24,700, MAE = ~19,200) and the testing set (RMSE = ~27,600, MAE = ~20,700). The small difference between the two suggests the model handles new data effectively while slightly favoring the training data, indicating mild overfitting.
# 
# Comparison with the Baseline:
# - RMSE improves from ~34,000 (Baseline) to ~27,600 on the testing set, reducing average error by a large margin.
# - MAE decreases from ~27,900 (Baseline) to ~20,700, showing the model’s predictions are significantly closer to the actual salaries.
# - \( R^2 \) improves from ~0.0 (Baseline) to 0.3433, meaning the model explains around 34% of the variation in *Salary*. This is a meaningful improvement over the Baseline’s inability to explain any variance.
# 
# Insights:
# - The improved \( R^2 \) value and reduced RMSE and MAE suggest that the model captures more of the relationships in the data compared to the Baseline and simpler models.
# - While there is a small gap between training and testing performance, the testing results indicate that the model can make accurate predictions on unseen data.

# In[216]:


# Define the polynomial regression pipeline for degree 3
poly_reg_degree3 = Pipeline([
    ('poly_features', PolynomialFeatures(degree=3, include_bias=False)),  # Generate polynomial features
    ('lin_reg', LinearRegression())  # Apply linear regression on the transformed features
])

# Train the model
poly_reg_degree3.fit(X_train_imputed_transformed, y_train_imputation)

# Predict on training and testing sets
y_train_poly3_pred = poly_reg_degree3.predict(X_train_imputed_transformed)
y_test_poly3_pred = poly_reg_degree3.predict(X_test_imputed_transformed)

# Evaluate the Polynomial Regression model (Degree 3)
poly3_rmse_train = np.sqrt(mean_squared_error(y_train_imputation, y_train_poly3_pred))
poly3_rmse_test = np.sqrt(mean_squared_error(y_test_imputation, y_test_poly3_pred))
poly3_mae_train = mean_absolute_error(y_train_imputation, y_train_poly3_pred)
poly3_mae_test = mean_absolute_error(y_test_imputation, y_test_poly3_pred)
poly3_r2_train = r2_score(y_train_imputation, y_train_poly3_pred)
poly3_r2_test = r2_score(y_test_imputation, y_test_poly3_pred)

# Display the results
print(f"Polynomial Regression Model (Degree 3) - RMSE on Training Set: {poly3_rmse_train:.2f}")
print(f"Polynomial Regression Model (Degree 3) - RMSE on Testing Set: {poly3_rmse_test:.2f}")
print(f"Polynomial Regression Model (Degree 3) - MAE on Training Set: {poly3_mae_train:.2f}")
print(f"Polynomial Regression Model (Degree 3) - MAE on Testing Set: {poly3_mae_test:.2f}")
print(f"Polynomial Regression Model (Degree 3) - R² on Training Set: {poly3_r2_train:.4f}")
print(f"Polynomial Regression Model (Degree 3) - R² on Testing Set: {poly3_r2_test:.4f}")


# #### Polynomial Regression Model (Degree 4)
# The Polynomial Regression Model (Degree 4) aims to capture even more detailed relationships between the features and *Salary*.
# 
# Model Fit:
# - The model performs very well on the training set (RMSE = ~23,800, MAE = ~18,300), but its performance drops noticeably on the testing set (RMSE = ~30,000, MAE = ~21,600). This significant gap indicates that the model is overfitting, fitting the training data too closely and struggling to generalize to unseen data.
# 
# Comparison with the Baseline:
# - RMSE improves from ~34,000 (Baseline) to ~30,000 on the testing set, showing a reduction in average error, but the improvement is smaller compared to simpler models.
# - MAE decreases from ~27,900 (Baseline) to ~21,600, indicating predictions are closer to actual salaries than the Baseline.
# - \( R^2 \) improves from ~0.0 (Baseline) to 0.2216, meaning the model explains around 22% of the variation in *Salary*. However, this is a weaker performance compared to the Degree 2 and Degree 3 polynomial models.
# 
# Insights:
# - While the Degree 4 model captures more complexity, it overfits the training data, leading to poorer generalization and reduced testing performance compared to simpler polynomial models.
# - The decrease in \( R^2 \) and increase in RMSE on the testing set suggest that adding more complexity does not necessarily lead to better predictions.
# - This model demonstrates the trade-off between capturing detailed patterns in training data and maintaining generalizability for new data.

# In[217]:


# Define the polynomial regression pipeline for degree 4
poly_reg_degree4 = Pipeline([
    ('poly_features', PolynomialFeatures(degree=4, include_bias=False)),  # Generate polynomial features
    ('lin_reg', LinearRegression())  # Apply linear regression on the transformed features
])

# Train the model
poly_reg_degree4.fit(X_train_imputed_transformed, y_train_imputation)

# Predict on training and testing sets
y_train_poly4_pred = poly_reg_degree4.predict(X_train_imputed_transformed)
y_test_poly4_pred = poly_reg_degree4.predict(X_test_imputed_transformed)

# Evaluate the Polynomial Regression model (degree 4)
poly4_rmse_train = np.sqrt(mean_squared_error(y_train_imputation, y_train_poly4_pred))
poly4_rmse_test = np.sqrt(mean_squared_error(y_test_imputation, y_test_poly4_pred))
poly4_mae_train = mean_absolute_error(y_train_imputation, y_train_poly4_pred)
poly4_mae_test = mean_absolute_error(y_test_imputation, y_test_poly4_pred)
poly4_r2_train = r2_score(y_train_imputation, y_train_poly4_pred)
poly4_r2_test = r2_score(y_test_imputation, y_test_poly4_pred)

# Display the results
print(f"Polynomial Regression Model (Degree 4) - RMSE on Training Set: {poly4_rmse_train:.2f}")
print(f"Polynomial Regression Model (Degree 4) - RMSE on Testing Set: {poly4_rmse_test:.2f}")
print(f"Polynomial Regression Model (Degree 4) - MAE on Training Set: {poly4_mae_train:.2f}")
print(f"Polynomial Regression Model (Degree 4) - MAE on Testing Set: {poly4_mae_test:.2f}")
print(f"Polynomial Regression Model (Degree 4) - R² on Training Set: {poly4_r2_train:.4f}")
print(f"Polynomial Regression Model (Degree 4) - R² on Testing Set: {poly4_r2_test:.4f}")


# #### Decision Tree Model
# The Decision Tree Model attempts to create a series of decision rules based on the features to predict *Salary*.
# 
# Model Fit:
# - The model performs exceptionally well on the training set (RMSE = ~13,350, MAE = ~6,800), but its performance on the testing set (RMSE = ~34,800, MAE = ~26,200) is significantly worse. This large gap indicates severe overfitting, where the model learns the training data too rigidly and fails to generalize to new data.
# - Cross-validation results (mean RMSE = ~34,400, standard deviation = ~1,800) confirm the poor generalization, showing consistent testing errors across folds.
# 
# Comparison with the Baseline:
# - RMSE on the testing set (34,800) is only marginally better than the Baseline (34,042), suggesting limited improvement in overall prediction accuracy.
# - MAE improves slightly compared to the Baseline (27,900 to 26,200), meaning predictions are slightly closer to actual salaries on average.
# - \( R^2 \) on the testing set drops to -0.0477, indicating that the model performs worse than simply predicting the mean salary for all data points.
# 
# Insights:
# - The Decision Tree Model heavily overfits the training data, capturing too many specific details that do not generalize to new data.
# - While the training metrics show excellent performance, the testing set and cross-validation results highlight the model's inability to generalize.
# - This model demonstrates the importance of balancing training performance with testing set generalizability and underscores the need for hyperparameter tuning, such as setting a maximum depth or limiting the number of leaves.

# In[222]:


# Initialize and fit the decision tree regressor
tree_reg = DecisionTreeRegressor(random_state=42)
tree_reg.fit(X_train_imputed_transformed, y_train_imputation)

# Predict on training and test sets
y_train_tree_pred = tree_reg.predict(X_train_imputed_transformed)
y_test_tree_pred = tree_reg.predict(X_test_imputed_transformed)

# Evaluate the Decision Tree model
tree_rmse_train = np.sqrt(mean_squared_error(y_train_imputation, y_train_tree_pred))
tree_rmse_test = np.sqrt(mean_squared_error(y_test_imputation, y_test_tree_pred))

tree_mae_train = mean_absolute_error(y_train_imputation, y_train_tree_pred)
tree_mae_test = mean_absolute_error(y_test_imputation, y_test_tree_pred)

tree_r2_train = r2_score(y_train_imputation, y_train_tree_pred)
tree_r2_test = r2_score(y_test_imputation, y_test_tree_pred)

# Display the results
print(f"Decision Tree Model - RMSE on Training Set: {tree_rmse_train:.2f}")
print(f"Decision Tree Model - RMSE on Testing Set: {tree_rmse_test:.2f}")
print(f"Decision Tree Model - MAE on Training Set: {tree_mae_train:.2f}")
print(f"Decision Tree Model - MAE on Testing Set: {tree_mae_test:.2f}")
print(f"Decision Tree Model - R² on Training Set: {tree_r2_train:.4f}")
print(f"Decision Tree Model - R² on Testing Set: {tree_r2_test:.4f}")

# Perform cross-validation
tree_cv_scores = cross_val_score(tree_reg, X_train_imputed_transformed, y_train_imputation, 
                                 cv=10, scoring="neg_mean_squared_error")
tree_cv_rmse = np.sqrt(-tree_cv_scores)

print(f"\nDecision Tree Model - Cross-Validation RMSE (Mean): {tree_cv_rmse.mean():.2f}")
print(f"Decision Tree Model - Cross-Validation RMSE (Std Dev): {tree_cv_rmse.std():.2f}")


# #### Decision Tree Model with Max Depth (Depth 5)
# This Decision Tree Model applies a maximum depth of 5 to limit the complexity of the decision tree and reduce overfitting.
# 
# Model Fit:
# - The model demonstrates improved generalization compared to the unrestricted Decision Tree, with a smaller gap between training and testing RMSE values (~25,070 vs. ~27,730).
# - Cross-validation results (mean RMSE = ~26,618, standard deviation = ~1,729) confirm more stable and reliable performance across folds.
# 
# Comparison with the Baseline:
# - RMSE on the testing set (27,726) is significantly lower than the Baseline (34,042), showing an improvement in prediction accuracy.
# - MAE decreases from ~27,900 (Baseline) to ~21,385, meaning predictions are closer to actual salaries on average.
# - \( R^2 \) on the testing set increases to 0.3364, indicating that the model explains ~34% of the variance in *Salary*, compared to the Baseline's near-zero \( R^2 \).
# 
# Insights:
# - Introducing a maximum depth effectively reduces overfitting, balancing training and testing performance.
# - The improvement in metrics compared to the unrestricted Decision Tree Model demonstrates the importance of constraining model complexity to enhance generalizability.
# - While this model does not outperform polynomial regression in testing \( R^2 \), it achieves relatively competitive performance with less risk of overfitting.

# In[223]:


# Initialize a Decision Tree Regressor with a max_depth constraint
tree_reg_depth = DecisionTreeRegressor(max_depth=5, random_state=42)
tree_reg_depth.fit(X_train_imputed_transformed, y_train_imputation)

# Predict on training and testing sets
y_train_tree_depth_pred = tree_reg_depth.predict(X_train_imputed_transformed)
y_test_tree_depth_pred = tree_reg_depth.predict(X_test_imputed_transformed)

# Evaluate the Decision Tree with max_depth on Training and Testing sets
tree_depth_rmse_train = np.sqrt(mean_squared_error(y_train_imputation, y_train_tree_depth_pred))
tree_depth_rmse_test = np.sqrt(mean_squared_error(y_test_imputation, y_test_tree_depth_pred))
tree_depth_mae_train = mean_absolute_error(y_train_imputation, y_train_tree_depth_pred)
tree_depth_mae_test = mean_absolute_error(y_test_imputation, y_test_tree_depth_pred)
tree_depth_r2_train = r2_score(y_train_imputation, y_train_tree_depth_pred)
tree_depth_r2_test = r2_score(y_test_imputation, y_test_tree_depth_pred)

# Perform cross-validation
tree_depth_cv_scores = cross_val_score(tree_reg_depth, X_train_imputed_transformed, y_train_imputation, 
                                       scoring="neg_mean_squared_error", cv=10)
tree_depth_cv_rmse_scores = np.sqrt(-tree_depth_cv_scores)

# Display the results
print(f"Decision Tree Model with Max Depth (Depth 5) - RMSE on Training Set: {tree_depth_rmse_train:.2f}")
print(f"Decision Tree Model with Max Depth (Depth 5) - RMSE on Testing Set: {tree_depth_rmse_test:.2f}")
print(f"Decision Tree Model with Max Depth (Depth 5) - MAE on Training Set: {tree_depth_mae_train:.2f}")
print(f"Decision Tree Model with Max Depth (Depth 5) - MAE on Testing Set: {tree_depth_mae_test:.2f}")
print(f"Decision Tree Model with Max Depth (Depth 5) - R² on Training Set: {tree_depth_r2_train:.4f}")
print(f"Decision Tree Model with Max Depth (Depth 5) - R² on Testing Set: {tree_depth_r2_test:.4f}")
print(f"Decision Tree Model with Max Depth (Depth 5) - Cross-Validation RMSE (Mean): {tree_depth_cv_rmse_scores.mean():.2f}")
print(f"Decision Tree Model with Max Depth (Depth 5) - Cross-Validation RMSE (Std Dev): {tree_depth_cv_rmse_scores.std():.2f}")


# ## Summary of Model Testing
# Throughout the model testing process, various approaches were evaluated to predict *Salary*, each offering unique strengths and challenges:
# 
# 1. Baseline Model: The simplest approach, using the mean salary for predictions, provided a benchmark for comparison. Its inability to explain variance (\( R^2 \approx 0 \)) and high RMSE (~34,000) highlighted the need for more sophisticated models.
# 
# 2. Linear Regression: Showed a significant improvement over the Baseline, reducing RMSE and explaining ~32% of the variance in *Salary*. However, the relatively high testing RMSE (~28,000) indicates room for improvement.
# 
# 3. Polynomial Regression:
#    - Degree 2 and Degree 3 models further reduced RMSE on the testing set, with Degree 3 providing a balanced improvement in \( R^2 \) (~34%) while maintaining generalizability.
#    - Degree 4 introduced overfitting, with a sharp decline in testing \( R^2 \) (~22%) despite better training performance.
# 
# 4. Decision Tree Models:
#    - An unrestricted Decision Tree severely overfit the training data (\( R^2 = 0.84 \)) and performed poorly on the testing set (\( R^2 = -0.05 \)).
#    - Adding a maximum depth (Depth 5) significantly improved testing performance, matching Polynomial Regression (Degree 3) with an \( R^2 \) of ~34% and achieving stable results across cross-validation.
# 
# #### Best Performing Model: 
# The Polynomial Regression Model (Degree 3) achieved the best balance of training and testing performance, with a testing \( R^2 \) of 0.3433 and an RMSE of ~27,580. It captures the complexity of the data without significant overfitting, making it the most effective model in this analysis.

# ## Relating the Best-Performing Model to the Hypothesis
# 
# To connect our modeling process to the project hypothesis, we tested the best-performing model, the Polynomial Regression Model (Degree 3), to predict salaries for two hypothetical individuals. Both individuals were assigned identical values for all features (e.g., years of experience, company size, remote work arrangement, and education), except for gender. 
# 
# #### Results:
# - **Predicted Salary (Female):** \$95,431.68  
# - **Predicted Salary (Male):** \$100,799.36  
# - **Salary Difference (Female - Male):** -\$5,367.68  
# 
# These results align with the hypothesis, suggesting that women in the tech industry may experience lower salaries compared to men with similar qualifications and job characteristics. While this analysis is based on modeled data and does not account for all potential factors, it provides evidence that supports the exploration of gender-based disparities in tech salaries.

# In[247]:


# Suppress specific UserWarning
warnings.filterwarnings("ignore", category=UserWarning, module="sklearn")

# Define the feature values for prediction
person_features = {
    'YearsCodedJobNumeric': 7,  
    'EducationCategoryNumeric': 4,  
    'Gender_Female': 1,  # Female
    'Gender_Male': 0,    # Not Male
    'HomeRemoteEncoded': 2,  
    'CompanySizeEncoded': 3,  
    'CompanyTypeEncoded': 3   
}

# Create a DataFrame for the features
person_female_df = pd.DataFrame([person_features])

# Modify for male prediction
person_features['Gender_Female'] = 0
person_features['Gender_Male'] = 1
person_male_df = pd.DataFrame([person_features])

# Apply the pipeline to transform the data
person_female_transformed = full_pipeline.transform(person_female_df)
person_male_transformed = full_pipeline.transform(person_male_df)

# Predict salaries using the Polynomial Regression Model (Degree 3)
predicted_salary_female = poly_reg.predict(person_female_transformed)[0]
predicted_salary_male = poly_reg.predict(person_male_transformed)[0]

# Calculate the difference
salary_difference = predicted_salary_female - predicted_salary_male

# Display results
print(f"Predicted Salary (Female): {predicted_salary_female:.2f}")
print(f"Predicted Salary (Male): {predicted_salary_male:.2f}")
print(f"Salary Difference (Female - Male): {salary_difference:.2f}")


# ## Resources and References
# *What resources and references have you used for this project?*
# 📝 <!-- Answer Below -->
# 
# - https://it4063c.github.io/course-notes/working-with-data/data-sources for methods to import the various data types  
# - https://www.kaggle.com/datasets/hackerrank/developer-survey-2018/data for the HackerRank Survey Kaggle dataset 
# - https://www.pewresearch.org/social-trends/2018/01/09/women-and-men-in-stem-often-at-odds-over-workplace-equity/  for the link to the Pew Research Survey
# - https://ncses.nsf.gov/pubs/nsb20212/participation-of-demographic-groups-in-stem for the NCSES html scraped dataset
# - https://ncses.nsf.gov/pubs/nsb20212/downloads for additional NCSES tables that were not part of the html scrape
# - https://www.census.gov/library/stories/2021/06/does-majoring-in-stem-lead-to-stem-job-after-graduation.html for the links to the American Community Survey 2019
# - https://survey.stackoverflow.co/ for the Stack Overflow Annual Developer Survey
# - IT4075 Applied Machine Learning zyBooks for data classification logic and code
# - https://medium.com/@acceldia/python-101-reading-excel-and-spss-files-with-pandas-eed6d0441c0b to learn how to work with .sav files
# - https://python-docx.readthedocs.io/en/latest/user/documents.html to learn how to work with .docx files inside Python
# - https://seaborn.pydata.org/generated/seaborn.FacetGrid.html to learn how to create and edit a seaborn FacetGrid
# - ChatGPT to troubleshoot visualizations (such as legends not showing correctly, labels partially hidden, and correcting sort order of categorical data)
# - https://stackoverflow.com/questions/1388450/giving-graphs-a-subtitle to learn how to add titles and subtitles to matplotlib visualizations
# - https://matplotlib.org/stable/gallery/color/named_colors.html to choose consistent color palette for visualizations
# - https://scikit-learn.org/dev/modules/generated/sklearn.compose.ColumnTransformer.html for issues with passthrough in pipeline

# In[ ]:


# ⚠️ Make sure you run this cell at the end of your notebook before every submission!
get_ipython().system('jupyter nbconvert --to python source.ipynb')

